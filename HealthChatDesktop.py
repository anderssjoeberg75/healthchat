"""
HealthChat - Standalone Desktop Application
A local desktop chatbot for querying Garmin Connect data.
"""

# Application version
APP_VERSION = "4.0.4"

import sys
from typing import Dict, Any, List, Optional
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox
import threading
import json
from pathlib import Path
from garmin_handler import GarminDataHandler
from fitbit_handler import FitbitHandler
from withings_handler import WithingsDataHandler
from strava_handler import StravaHandler
from ai_client import AIClient
from charts_view import HealthChartsView
from garmin_db import GarminDatabase
import logging
from datetime import datetime

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class SplashScreen(tk.Toplevel):
    """Splash screen shown during app startup"""
    
    def __init__(self, parent):
        super().__init__(parent)
        self.title("HealthChat")
        
        # Remove window decorations
        self.overrideredirect(True)
        
        # Set size (increased height to prevent text cutoff)
        width = 400
        height = 350
        
        # Center on screen
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2
        self.geometry(f"{width}x{height}+{x}+{y}")
        
        # Set background color
        self.configure(bg='#0078D4')
        
        # Create frame with shadow effect
        frame = tk.Frame(self, bg='#0078D4')
        frame.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        # Try to load and display logo
        try:
            if getattr(sys, 'frozen', False):
                base_path = Path(sys._MEIPASS)
            else:
                base_path = Path(__file__).parent
            
            logo_path = base_path / "logo.png"
            if logo_path.exists():
                from PIL import Image, ImageTk
                img = Image.open(logo_path)
                img = img.resize((100, 100), Image.Resampling.LANCZOS)
                photo = ImageTk.PhotoImage(img)
                logo_label = tk.Label(frame, image=photo, bg='#0078D4')
                logo_label.image = photo  # Keep reference
                logo_label.pack(pady=(40, 10))
        except Exception as e:
            logger.debug(f"Could not load splash logo: {e}")
        
        # App name
        title_label = tk.Label(frame,
                              text="HealthChat",
                              font=('Segoe UI', 24, 'bold'),
                              bg='#0078D4',
                              fg='white')
        title_label.pack(pady=(10, 5))
        
        # Version
        version_label = tk.Label(frame,
                                text=f"Version {APP_VERSION}",
                                font=('Segoe UI', 10),
                                bg='#0078D4',
                                fg='white')
        version_label.pack(pady=(0, 30))
        
        # Loading message
        self.status_label = tk.Label(frame,
                                     text="Loading...",
                                     font=('Segoe UI', 10),
                                     bg='#0078D4',
                                     fg='white')
        self.status_label.pack(pady=10)
        
        # Progress bar
        self.progress = ttk.Progressbar(frame,
                                       mode='indeterminate',
                                       length=300)
        self.progress.pack(pady=10)
        self.progress.start(10)
        
        # Update to show window
        self.update()
    
    def update_status(self, message):
        """Update the status message"""
        self.status_label.config(text=message)
        self.update()
    
    def close(self):
        """Close the splash screen"""
        self.progress.stop()
        self.destroy()


class SettingsDialog(tk.Toplevel):
    """Dialog for managing application settings including AI provider selection"""
    
    def __init__(self, parent, current_config=None, colors=None):
        super().__init__(parent)
        
        # Make modal and transient FIRST (before withdraw)
        self.transient(parent)
        
        self.title("Settings")
        
        # Store colors (use parent's colors or defaults)
        self.colors = colors or {
            'bg': '#F3F3F3',
            'card_bg': '#FFFFFF',
            'text': '#1F1F1F',
            'text_secondary': '#605E5C',
            'border': '#EDEBE9',
            'accent': '#0078D4'
        }
        
        # Calculate centered position BEFORE setting geometry
        width = 700
        height = 700
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (width // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (height // 2)
        
        # Set geometry with position in one call (prevents flashing)
        self.geometry(f"{width}x{height}+{x}+{y}")
        self.resizable(False, False)
        
        # Withdraw to prevent flash during setup
        self.withdraw()
        
        # Apply Fluent Design theme to dialog
        self.configure(bg=self.colors['bg'])
        # Set window icon (same as main window)
        try:
            # Get the correct base path for PyInstaller exe
            if getattr(sys, 'frozen', False):
                # Running as compiled executable
                base_path = Path(sys._MEIPASS)
            else:
                # Running as script
                base_path = Path(__file__).parent
            
            icon_path = base_path / "logo.ico"
            if icon_path.exists():
                self.iconbitmap(str(icon_path))
        except Exception as e:
            logger.debug(f"Could not load Settings dialog icon: {e}")
        
        self.result = None
        self.current_config = current_config or {}
        
        # Get AI provider info
        from ai_client import AIClient
        self.providers = AIClient.get_available_providers()
        
        # Create StringVars for ALL providers upfront (so they persist when switching)
        # This ensures all keys are saved, not just the currently selected provider
        for provider_id in self.providers.keys():
            if provider_id == 'azure':
                # Azure needs special handling
                if not hasattr(self, 'azure_endpoint_var'):
                    self.azure_endpoint_var = tk.StringVar(value=current_config.get('azure_endpoint', ''))
                if not hasattr(self, 'azure_deployment_var'):
                    self.azure_deployment_var = tk.StringVar(value=current_config.get('azure_deployment', ''))
                if not hasattr(self, 'azure_key_var'):
                    self.azure_key_var = tk.StringVar(value=current_config.get('azure_api_key', ''))
            elif provider_id == 'ollama':
                # Ollama uses URL + model, no API key
                if not hasattr(self, 'ollama_base_url_var'):
                    self.ollama_base_url_var = tk.StringVar(value=current_config.get('ollama_base_url', 'http://localhost:11434/v1'))
                if not hasattr(self, 'ollama_model_var'):
                    self.ollama_model_var = tk.StringVar(value=current_config.get('ollama_model', 'llama3.2'))
            else:
                # Create key and model vars for other providers
                key_var_name = f'{provider_id}_key_var'
                if not hasattr(self, key_var_name):
                    setattr(self, key_var_name, tk.StringVar(value=current_config.get(f'{provider_id}_api_key', '')))
                
                model_var_name = f'{provider_id}_model_var'
                if not hasattr(self, model_var_name):
                    default_model = current_config.get(f'{provider_id}_model', self.providers[provider_id]['default_model'])
                    setattr(self, model_var_name, tk.StringVar(value=default_model))
        
        self.create_widgets()
        
        # Show window and grab focus after everything is set up
        self.deiconify()
        self.grab_set()
        
    def create_widgets(self):
        """Create settings dialog widgets"""
        # Configure ttk styles for this dialog with current theme
        style = ttk.Style()
        
        style.configure('Settings.TFrame',
                       background=self.colors['bg'])
        
        style.configure('Settings.TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'])
        
        style.configure('Settings.Header.TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 11, 'bold'))
        
        style.configure('Settings.CardHeader.TLabel',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 10, 'bold'))
        
        style.configure('Settings.CardText.TLabel',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 9))
        
        style.configure('Settings.Title.TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 14, 'bold'))
        
        style.configure('Settings.Help.TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text_secondary'],
                       font=('Segoe UI', 9))
        
        style.configure('Settings.TEntry',
                       fieldbackground=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       bordercolor=self.colors['border'])
        
        style.configure('Settings.TCheckbutton',
                       background=self.colors['bg'],
                       foreground=self.colors['text'])
        
        # Map state-specific colors for checkbuttons (fixes hover in dark mode)
        style.map('Settings.TCheckbutton',
                 background=[('active', self.colors['bg']), ('!active', self.colors['bg'])],
                 foreground=[('active', self.colors['text']), ('!active', self.colors['text'])])
        
        style.configure('Settings.TButton',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'])
        
        style.configure('Settings.TRadiobutton',
                       background=self.colors['bg'],
                       foreground=self.colors['text'])
        
        # Map state-specific colors for radiobuttons (fixes hover in dark mode)
        style.map('Settings.TRadiobutton',
                 background=[('active', self.colors['bg']), ('!active', self.colors['bg'])],
                 foreground=[('active', self.colors['text']), ('!active', self.colors['text'])])
        
        # Card style for sections
        style.configure('Settings.Card.TFrame',
                       background=self.colors['card_bg'],
                       relief='flat',
                       borderwidth=1)
        
        # Create scrollable frame for settings
        canvas = tk.Canvas(self, bg=self.colors['bg'], highlightthickness=0)
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style='Settings.TFrame')
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        main_frame = ttk.Frame(scrollable_frame, padding="25", style='Settings.TFrame')
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        main_frame.columnconfigure(1, weight=1)
        
        # Title
        title_label = ttk.Label(main_frame,
                               text="Application Settings",
                               style='Settings.Title.TLabel')
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 25))
        
        current_row = 1
        
        # AI Provider Selection
        provider_header = ttk.Label(main_frame,
                                   text="AI Provider Selection",
                                   style='Settings.Header.TLabel')
        provider_header.grid(row=current_row, column=0, columnspan=2, sticky=tk.W, pady=(0, 10))
        current_row += 1
        
        ttk.Label(main_frame, 
                 text="Choose your preferred AI provider:",
                 style='Settings.TLabel').grid(row=current_row, column=0, columnspan=2, sticky=tk.W, pady=5)
        current_row += 1
        
        # Provider radio buttons
        self.provider_var = tk.StringVar(value=self.current_config.get('ai_provider', 'xai'))
        
        for provider_id, provider_info in self.providers.items():
            rb = ttk.Radiobutton(main_frame,
                                text=provider_info['name'],
                                variable=self.provider_var,
                                value=provider_id,
                                style='Settings.TRadiobutton',
                                command=self.on_provider_change)
            rb.grid(row=current_row, column=0, columnspan=2, sticky=tk.W, pady=2, padx=20)
            current_row += 1
        
        current_row += 1
        
        # API Keys Section - Dynamic based on selected provider
        self.api_keys_frame = ttk.Frame(main_frame, style='Settings.Card.TFrame', padding="15")
        self.api_keys_frame.grid(row=current_row, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(10, 0))
        self.api_keys_frame.columnconfigure(1, weight=1)
        current_row += 1
        
        # Storage for API key entries
        self.api_key_entries = {}
        
        # Create API key fields for all providers
        self.create_api_key_fields()
        
        # Garmin Credentials section
        garmin_header = ttk.Label(main_frame,
                                 text="Garmin Connect Credentials",
                                 style='Settings.Header.TLabel')
        garmin_header.grid(row=current_row, column=0, columnspan=2, sticky=tk.W, pady=(20, 10))
        current_row += 1
        
        ttk.Label(main_frame, text="Email:", style='Settings.TLabel').grid(row=current_row, column=0, sticky=tk.W, pady=8)
        
        self.email_var = tk.StringVar(value=self.current_config.get('garmin_email', ''))
        email_entry = ttk.Entry(main_frame,
                               textvariable=self.email_var,
                               width=50,
                               style='Settings.TEntry')
        email_entry.grid(row=current_row, column=1, sticky=(tk.W, tk.E), pady=8)
        current_row += 1
        
        ttk.Label(main_frame, text="Password:", style='Settings.TLabel').grid(row=current_row, column=0, sticky=tk.W, pady=8)
        
        self.password_var = tk.StringVar(value=self.current_config.get('garmin_password', ''))
        password_entry = ttk.Entry(main_frame,
                                  textvariable=self.password_var,
                                  width=50,
                                  show="*",
                                  style='Settings.TEntry')
        password_entry.grid(row=current_row, column=1, sticky=(tk.W, tk.E), pady=8)
        current_row += 1
        
        # Withings API Credentials Section
        withings_header = ttk.Label(main_frame,
                                   text="Withings Health Mate API Credentials",
                                   style='Settings.Header.TLabel')
        withings_header.grid(row=current_row, column=0, columnspan=2, sticky=tk.W, pady=(20, 10))
        current_row += 1
        
        ttk.Label(main_frame, text="Client ID:", style='Settings.TLabel').grid(row=current_row, column=0, sticky=tk.W, pady=8)
        self.withings_client_id_var = tk.StringVar(value=self.current_config.get('withings_client_id', ''))
        ttk.Entry(main_frame, textvariable=self.withings_client_id_var, width=50, style='Settings.TEntry').grid(row=current_row, column=1, sticky=(tk.W, tk.E), pady=8)
        current_row += 1

        ttk.Label(main_frame, text="Client Secret:", style='Settings.TLabel').grid(row=current_row, column=0, sticky=tk.W, pady=8)
        self.withings_client_secret_var = tk.StringVar(value=self.current_config.get('withings_client_secret', ''))
        ttk.Entry(main_frame, textvariable=self.withings_client_secret_var, width=50, show="*", style='Settings.TEntry').grid(row=current_row, column=1, sticky=(tk.W, tk.E), pady=8)
        current_row += 1

        ttk.Label(main_frame, text="Refresh Token:", style='Settings.TLabel').grid(row=current_row, column=0, sticky=tk.W, pady=8)
        self.withings_refresh_token_var = tk.StringVar(value=self.current_config.get('withings_refresh_token', ''))
        ttk.Entry(main_frame, textvariable=self.withings_refresh_token_var, width=50, show="*", style='Settings.TEntry').grid(row=current_row, column=1, sticky=(tk.W, tk.E), pady=8)
        current_row += 1

        # Strava API Credentials Section
        strava_header = ttk.Label(main_frame,
                                   text="Strava API Credentials",
                                   style='Settings.Header.TLabel')
        strava_header.grid(row=current_row, column=0, columnspan=2, sticky=tk.W, pady=(20, 10))
        current_row += 1
        
        ttk.Label(main_frame, text="Client ID:", style='Settings.TLabel').grid(row=current_row, column=0, sticky=tk.W, pady=8)
        self.strava_client_id_var = tk.StringVar(value=self.current_config.get('strava_client_id', ''))
        ttk.Entry(main_frame, textvariable=self.strava_client_id_var, width=50, style='Settings.TEntry').grid(row=current_row, column=1, sticky=(tk.W, tk.E), pady=8)
        current_row += 1

        ttk.Label(main_frame, text="Client Secret:", style='Settings.TLabel').grid(row=current_row, column=0, sticky=tk.W, pady=8)
        self.strava_client_secret_var = tk.StringVar(value=self.current_config.get('strava_client_secret', ''))
        ttk.Entry(main_frame, textvariable=self.strava_client_secret_var, width=50, show="*", style='Settings.TEntry').grid(row=current_row, column=1, sticky=(tk.W, tk.E), pady=8)
        current_row += 1

        ttk.Label(main_frame, text="Refresh Token:", style='Settings.TLabel').grid(row=current_row, column=0, sticky=tk.W, pady=8)
        self.strava_refresh_token_var = tk.StringVar(value=self.current_config.get('strava_refresh_token', ''))
        ttk.Entry(main_frame, textvariable=self.strava_refresh_token_var, width=50, show="*", style='Settings.TEntry').grid(row=current_row, column=1, sticky=(tk.W, tk.E), pady=8)
        current_row += 1

        # Buttons
        button_frame = ttk.Frame(main_frame, style='Settings.TFrame')
        button_frame.grid(row=current_row, column=0, columnspan=2, pady=(30, 0))
        
        ttk.Button(button_frame,
                  text="Save Settings",
                  command=self.save_settings,
                  style='Settings.TButton').grid(row=0, column=0, padx=5)
        
        ttk.Button(button_frame,
                  text="Cancel",
                  command=self.cancel,
                  style='Settings.TButton').grid(row=0, column=1, padx=5)
        
        # Pack canvas and scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Show initial provider fields
        self.on_provider_change()
    
    def create_api_key_fields(self):
        """Create API key entry fields for all providers"""
        # Clear existing
        for widget in self.api_keys_frame.winfo_children():
            widget.destroy()
        
        row = 0
        
        # Header
        header = ttk.Label(self.api_keys_frame,
                          text="API Configuration",
                          style='Settings.CardHeader.TLabel')
        header.grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=(0, 10))
        row += 1
        
        selected_provider = self.provider_var.get()
        
        # Show fields based on selected provider
        if selected_provider == 'azure':
            # Azure needs endpoint, deployment, and API key
            ttk.Label(self.api_keys_frame, text="Azure Endpoint:", style='Settings.CardText.TLabel').grid(row=row, column=0, sticky=tk.W, pady=5)
            self.azure_endpoint_var = tk.StringVar(value=self.current_config.get('azure_endpoint', ''))
            ttk.Entry(self.api_keys_frame, textvariable=self.azure_endpoint_var, width=50, style='Settings.TEntry').grid(row=row, column=1, sticky=(tk.W, tk.E), pady=5)
            row += 1
            
            ttk.Label(self.api_keys_frame, text="Deployment Name:", style='Settings.CardText.TLabel').grid(row=row, column=0, sticky=tk.W, pady=5)
            self.azure_deployment_var = tk.StringVar(value=self.current_config.get('azure_deployment', ''))
            ttk.Entry(self.api_keys_frame, textvariable=self.azure_deployment_var, width=50, style='Settings.TEntry').grid(row=row, column=1, sticky=(tk.W, tk.E), pady=5)
            row += 1
            
            ttk.Label(self.api_keys_frame, text="API Key:", style='Settings.CardText.TLabel').grid(row=row, column=0, sticky=tk.W, pady=5)
            self.azure_key_var = tk.StringVar(value=self.current_config.get('azure_api_key', ''))
            ttk.Entry(self.api_keys_frame, textvariable=self.azure_key_var, width=50, show="*", style='Settings.TEntry').grid(row=row, column=1, sticky=(tk.W, tk.E), pady=5)
            row += 1
        elif selected_provider == 'ollama':
            # Ollama: no API key, but allow custom base URL and dynamic model fetching
            ttk.Label(self.api_keys_frame, text="Ollama URL:", style='Settings.CardText.TLabel').grid(row=row, column=0, sticky=tk.W, pady=5)
            if not hasattr(self, 'ollama_base_url_var'):
                self.ollama_base_url_var = tk.StringVar(value=self.current_config.get('ollama_base_url', 'http://localhost:11434/v1'))
            ttk.Entry(self.api_keys_frame, textvariable=self.ollama_base_url_var, width=50, style='Settings.TEntry').grid(row=row, column=1, sticky=(tk.W, tk.E), pady=5)
            row += 1
            
            ttk.Label(self.api_keys_frame, text="Model:", style='Settings.CardText.TLabel').grid(row=row, column=0, sticky=tk.W, pady=5)
            models = self.providers['ollama']['models']
            if not hasattr(self, 'ollama_model_var'):
                self.ollama_model_var = tk.StringVar(value=self.current_config.get('ollama_model', 'llama3.2'))
            
            model_frame = ttk.Frame(self.api_keys_frame, style='Settings.Card.TFrame')
            model_frame.grid(row=row, column=1, sticky=(tk.W, tk.E), pady=5)
            model_frame.columnconfigure(0, weight=1)
            
            model_combo = ttk.Combobox(model_frame, textvariable=self.ollama_model_var, values=models, width=35)
            model_combo.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))
            
            status_lbl = ttk.Label(self.api_keys_frame, text="", style='Settings.Help.TLabel')
            
            def _do_fetch_ollama_models():
                status_lbl.config(text="🔄 Querying Ollama server...")
                base_url = self.ollama_base_url_var.get().strip()
                
                def _thread_target():
                    fetched = AIClient.fetch_models('ollama', base_url=base_url)
                    def _update():
                        if fetched:
                            model_combo['values'] = fetched
                            status_lbl.config(text=f"✅ Loaded {len(fetched)} local Ollama models")
                        else:
                            status_lbl.config(text="⚠️ Using default models (check if Ollama is running)")
                    try:
                        self.after(0, _update)
                    except Exception:
                        pass
                
                threading.Thread(target=_thread_target, daemon=True).start()
            
            fetch_btn = ttk.Button(model_frame, text="🔄 Fetch Models", command=_do_fetch_ollama_models, style='Settings.TButton')
            fetch_btn.grid(row=0, column=1, sticky=tk.E)
            row += 1
            
            status_lbl.grid(row=row, column=1, sticky=tk.W, pady=(0, 5))
            row += 1
            
            # Auto-fetch local Ollama models on view
            _do_fetch_ollama_models()
        else:
            # Standard API key for other providers
            provider_name = self.providers[selected_provider]['name']
            ttk.Label(self.api_keys_frame, text=f"{provider_name} API Key:", style='Settings.CardText.TLabel').grid(row=row, column=0, sticky=tk.W, pady=5)
            
            # Use existing var (already created in __init__)
            key_var = getattr(self, f'{selected_provider}_key_var')
            ttk.Entry(self.api_keys_frame, textvariable=key_var, width=50, show="*", style='Settings.TEntry').grid(row=row, column=1, sticky=(tk.W, tk.E), pady=5)
            row += 1
            
            # Model selection
            models = self.providers[selected_provider]['models']
            if models:
                model_var = getattr(self, f'{selected_provider}_model_var')
                ttk.Label(self.api_keys_frame, text="Model:", style='Settings.CardText.TLabel').grid(row=row, column=0, sticky=tk.W, pady=5)
                
                model_frame = ttk.Frame(self.api_keys_frame, style='Settings.Card.TFrame')
                model_frame.grid(row=row, column=1, sticky=(tk.W, tk.E), pady=5)
                model_frame.columnconfigure(0, weight=1)
                
                model_combo = ttk.Combobox(model_frame, textvariable=model_var, values=models, width=35)
                model_combo.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))
                
                status_lbl = ttk.Label(self.api_keys_frame, text="", style='Settings.Help.TLabel')
                
                def _do_fetch_models():
                    status_lbl.config(text="🔄 Fetching models from API...")
                    api_key = key_var.get().strip() if key_var else ''
                    
                    def _thread_target():
                        fetched = AIClient.fetch_models(selected_provider, api_key=api_key)
                        def _update():
                            if fetched:
                                model_combo['values'] = fetched
                                status_lbl.config(text=f"✅ Loaded {len(fetched)} models dynamically from API")
                            else:
                                status_lbl.config(text="⚠️ Using default models (check API key)")
                        try:
                            self.after(0, _update)
                        except Exception:
                            pass
                    
                    threading.Thread(target=_thread_target, daemon=True).start()
                
                fetch_btn = ttk.Button(model_frame, text="🔄 Fetch Models", command=_do_fetch_models, style='Settings.TButton')
                fetch_btn.grid(row=0, column=1, sticky=tk.E)
                row += 1
                
                status_lbl.grid(row=row, column=1, sticky=tk.W, pady=(0, 5))
                row += 1
                
                # Auto-fetch if key exists
                if key_var and key_var.get().strip():
                    _do_fetch_models()
        
        # Help text
        help_text = self.get_provider_help_text(selected_provider)
        if help_text:
            help_label = ttk.Label(self.api_keys_frame,
                                  text=help_text,
                                  style='Settings.Help.TLabel',
                                  wraplength=550)
            help_label.grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=(10, 0))
    
    def get_provider_help_text(self, provider):
        """Get help text for each provider"""
        help_texts = {
            'xai': "Get your xAI API key from: https://console.x.ai/",
            'openai': "Get your OpenAI API key from: https://platform.openai.com/api-keys",
            'azure': "Azure endpoint format: https://your-resource.openai.azure.com/",
            'gemini': "Get your Google AI API key from: https://makersuite.google.com/app/apikey",
            'anthropic': "Get your Anthropic API key from: https://console.anthropic.com/",
            'ollama': "Ollama runs locally - no API key needed! Install from https://ollama.com, then run 'ollama serve'"
        }
        return help_texts.get(provider, "")
    
    def on_provider_change(self):
        """Called when provider selection changes"""
        self.create_api_key_fields()
    
    def save_settings(self):
        """Save settings and close dialog"""
        selected_provider = self.provider_var.get()
        
        self.result = {
            'ai_provider': selected_provider,
            'garmin_email': self.email_var.get(),
            'garmin_password': self.password_var.get(),
            'withings_client_id': getattr(self, 'withings_client_id_var', tk.StringVar()).get(),
            'withings_client_secret': getattr(self, 'withings_client_secret_var', tk.StringVar()).get(),
            'withings_refresh_token': getattr(self, 'withings_refresh_token_var', tk.StringVar()).get(),
            'strava_client_id': getattr(self, 'strava_client_id_var', tk.StringVar()).get(),
            'strava_client_secret': getattr(self, 'strava_client_secret_var', tk.StringVar()).get(),
            'strava_refresh_token': getattr(self, 'strava_refresh_token_var', tk.StringVar()).get()
        }
        
        # Save ALL providers' keys (not just selected one)
        # This allows easy switching between providers
        for provider_id in self.providers.keys():
            if provider_id == 'azure':
                # Azure has special fields
                self.result['azure_endpoint'] = self.azure_endpoint_var.get()
                self.result['azure_deployment'] = self.azure_deployment_var.get()
                self.result['azure_api_key'] = self.azure_key_var.get()
            elif provider_id == 'ollama':
                # Ollama has URL instead of API key
                if hasattr(self, 'ollama_base_url_var'):
                    self.result['ollama_base_url'] = self.ollama_base_url_var.get()
                if hasattr(self, 'ollama_model_var'):
                    self.result['ollama_model'] = self.ollama_model_var.get()
            else:
                # Save API key and model for each provider
                key_var = getattr(self, f'{provider_id}_key_var', None)
                if key_var:
                    self.result[f'{provider_id}_api_key'] = key_var.get()
                
                model_var = getattr(self, f'{provider_id}_model_var', None)
                if model_var:
                    self.result[f'{provider_id}_model'] = model_var.get()
        
        self.destroy()
    
    def cancel(self):
        """Cancel and close dialog"""
        self.result = None
        self.destroy()


class FitbitConnectDialog(tk.Toplevel):
    """Dedicated modal dialog to connect Fitbit Developer API account with instructions and link."""

    def __init__(self, parent, client_id="", client_secret="", colors=None):
        super().__init__(parent)
        self.title("Anslut till Fitbit Developer Portal")
        self.geometry("630x580")
        self.resizable(False, False)
        self.result = None
        self.colors = colors or {'bg': '#F3F4F6', 'card_bg': '#FFFFFF', 'text': '#1F2937', 'accent': '#0078D4'}

        self.transient(parent)
        self.grab_set()
        self.configure(bg=self.colors.get('bg', '#F3F4F6'))

        main_frame = ttk.Frame(self, padding="20", style='Card.TFrame')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        main_frame.columnconfigure(1, weight=1)

        # Header
        ttk.Label(
            main_frame,
            text="⌚ Anslut till Fitbit Developer Portal",
            font=('Segoe UI', 14, 'bold'),
            foreground=self.colors.get('accent', '#0078D4')
        ).grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=(0, 8))

        # Instructions
        instruction_text = (
            "Hur du hämtar dina Fitbit API-uppgifter:\n"
            "1. Klicka på länkknappen nedan för att öppna dev.fitbit.com/apps i webbläsaren.\n"
            "2. Logga in med ditt Fitbit-konto och klicka på 'Register a New App'.\n"
            "3. Fyll i följande i formuläret:\n"
            "   • Application Name: HealthChat\n"
            "   • Application Type: Personal\n"
            "   • Callback URL: http://localhost:8080\n"
            "   • Default Access Type: Read-Only (eller Read-Write)\n"
            "4. Klicka på 'Save'. Kopiera 'OAuth 2.0 Client ID' och 'Client Secret' hit nedan.\n"
            "5. Klicka på '▶ Öppna Inloggning i Webbläsare' för att ansluta helt automatiskt!"
        )
        ttk.Label(
            main_frame,
            text=instruction_text,
            wraplength=560,
            font=('Segoe UI', 9)
        ).grid(row=1, column=0, columnspan=2, sticky=tk.W, pady=(0, 8))

        # Clickable Link Label
        def _open_fitbit_dev_url(event=None):
            import webbrowser
            webbrowser.open("https://dev.fitbit.com/apps")

        link_label = tk.Label(
            main_frame,
            text="🔗 Öppna Fitbit Developer Portal (dev.fitbit.com/apps)",
            font=('Segoe UI', 9, 'underline', 'bold'),
            fg='#0078D4',
            bg=self.colors.get('card_bg', '#FFFFFF'),
            cursor="hand2"
        )
        link_label.grid(row=2, column=0, columnspan=2, sticky=tk.W, pady=(0, 12))
        link_label.bind("<Button-1>", _open_fitbit_dev_url)

        # Fields
        ttk.Label(main_frame, text="Client ID:", font=('Segoe UI', 10, 'bold')).grid(row=3, column=0, sticky=tk.W, pady=5)
        self.client_id_var = tk.StringVar(value=client_id)
        ttk.Entry(main_frame, textvariable=self.client_id_var, width=42).grid(row=3, column=1, sticky=(tk.W, tk.E), pady=5)

        ttk.Label(main_frame, text="Client Secret:", font=('Segoe UI', 10, 'bold')).grid(row=4, column=0, sticky=tk.W, pady=5)
        self.client_secret_var = tk.StringVar(value=client_secret)
        ttk.Entry(main_frame, textvariable=self.client_secret_var, width=42, show="*").grid(row=4, column=1, sticky=(tk.W, tk.E), pady=5)

        ttk.Label(main_frame, text="Callback URL:", font=('Segoe UI', 10, 'bold')).grid(row=5, column=0, sticky=tk.W, pady=5)
        self.callback_url_var = tk.StringVar(value="http://localhost:8080")
        ttk.Entry(main_frame, textvariable=self.callback_url_var, width=42).grid(row=5, column=1, sticky=(tk.W, tk.E), pady=5)

        ttk.Label(main_frame, text="Auth Code / Token:", font=('Segoe UI', 10, 'bold')).grid(row=6, column=0, sticky=tk.W, pady=5)
        self.code_var = tk.StringVar()
        ttk.Entry(main_frame, textvariable=self.code_var, width=42).grid(row=6, column=1, sticky=(tk.W, tk.E), pady=5)

        # Buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=7, column=0, columnspan=2, pady=(15, 0))

        ttk.Button(btn_frame, text="▶ Öppna Inloggning i Webbläsare", command=self._open_browser_auth, style='Modern.TButton').pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_frame, text="💾 Spara & Anslut", command=self.save, style='Accent.TButton').pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_frame, text="❌ Avbryt", command=self.destroy, style='Modern.TButton').pack(side=tk.LEFT, padx=4)

    def _open_browser_auth(self):
        cid = self.client_id_var.get().strip()

        if not cid:
            messagebox.showwarning("Client ID Saknas", "Fyll i ditt Client ID först innan du klickar på inloggning.", parent=self)
            return

        import http.server
        import urllib.parse
        import threading
        import webbrowser

        port = 8080
        dialog_ref = self

        class FitbitCallbackHandler(http.server.BaseHTTPRequestHandler):
            def do_GET(self):
                parsed_path = urllib.parse.urlparse(self.path)
                query_params = urllib.parse.parse_qs(parsed_path.query)
                code = query_params.get('code', [None])[0]

                self.send_response(200)
                self.send_header('Content-type', 'text/html; charset=utf-8')
                self.end_headers()

                if code:
                    html = """
                    <html>
                    <body style="font-family: Segoe UI, sans-serif; text-align: center; padding-top: 50px; background: #F3F4F6;">
                        <div style="background: white; max-width: 500px; margin: 0 auto; padding: 40px; border-radius: 12px; box-shadow: 0 4px 12px rgba(0,0,0,0.1);">
                            <h2 style="color: #10B981; margin-bottom: 10px;">✅ Fitbit Ansluten!</h2>
                            <p style="color: #4B5563; font-size: 16px;">Ditt Fitbit-konto har anslutits framgångsrikt till HealthChat.</p>
                            <p style="color: #6B7280; font-size: 14px;">Du kan nu stänga den här fliken i webbläsaren!</p>
                        </div>
                    </body>
                    </html>
                    """
                    self.wfile.write(html.encode('utf-8'))

                    def _update_dialog():
                        try:
                            if dialog_ref.winfo_exists():
                                dialog_ref.code_var.set(code)
                                dialog_ref.save()
                        except Exception:
                            pass
                    dialog_ref.after(100, _update_dialog)
                else:
                    self.wfile.write("<html><body><h2>Ingen kod mottogs</h2></body></html>".encode('utf-8'))

            def log_message(self, format, *args):
                pass

        try:
            server = http.server.HTTPServer(('127.0.0.1', port), FitbitCallbackHandler)
            server.timeout = 120
            threading.Thread(target=lambda: (server.handle_request(), server.server_close()), daemon=True).start()
        except Exception as e:
            logger.warning(f"Could not start local callback server on port {port}: {e}")

        auth_url = f"https://www.fitbit.com/oauth2/authorize?response_type=code&client_id={cid}&redirect_uri=http%3A%2F%2Flocalhost%3A8080&scope=activity%20heartrate%20location%20nutrition%20profile%20settings%20sleep%20social%20weight"
        webbrowser.open(auth_url)
        messagebox.showinfo(
            "Fitbit Inloggning",
            "Öppnar Fitbit-auktoriseringssidan i webbläsaren.\n\n"
            "1. Logga in och godkänn behörigheterna.\n"
            "2. Appen tar emot koden och genomför anslutningen helt automatiskt!",
            parent=self
        )

    def save(self):
        self.result = {
            'client_id': self.client_id_var.get().strip(),
            'client_secret': self.client_secret_var.get().strip(),
            'code': self.code_var.get().strip()
        }
        self.destroy()


def _clean_oauth_code(raw_input: str) -> str:
    """Extract clean authorization code if user pasted a full URL or query string with &state."""
    if not raw_input:
        return ""
    val = raw_input.strip()
    if 'code=' in val:
        val = val.split('code=', 1)[1]
        if '&' in val:
            val = val.split('&', 1)[0]
        return val.strip()
    if '&' in val:
        val = val.split('&', 1)[0]
    if '?' in val:
        val = val.split('?', 1)[-1]
    return val.strip()


class StravaConnectDialog(tk.Toplevel):
    """Dedicated modal dialog to connect Strava API account with instructions and link."""

    def __init__(self, parent, client_id="", client_secret="", colors=None):
        super().__init__(parent)
        self.title("Anslut till Strava API")
        self.geometry("630x580")
        self.resizable(False, False)
        self.result = None
        self.colors = colors or {'bg': '#F3F4F6', 'card_bg': '#FFFFFF', 'text': '#1F2937', 'accent': '#FC4C02'}

        self.transient(parent)
        self.grab_set()
        self.configure(bg=self.colors.get('bg', '#F3F4F6'))

        main_frame = ttk.Frame(self, padding="20", style='Card.TFrame')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        main_frame.columnconfigure(1, weight=1)

        # Header
        ttk.Label(
            main_frame,
            text="🏃 Anslut till Strava API",
            font=('Segoe UI', 14, 'bold'),
            foreground='#FC4C02'
        ).grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=(0, 8))

        # Instructions
        instruction_text = (
            "Hur du hämtar dina Strava API-uppgifter:\n"
            "1. Klicka på länkknappen nedan för att öppna strava.com/settings/api i webbläsaren.\n"
            "2. Skapa eller visa din API Application.\n"
            "3. Ange följande i formuläret om du skapar en ny app:\n"
            "   • Application Name: HealthChat\n"
            "   • Category: Visualizer / Analytics\n"
            "   • Authorization Callback Domain: localhost\n"
            "4. Klicka på 'Save'. Kopiera 'Client ID' och 'Client Secret' hit nedan.\n"
            "5. Klicka på '▶ Öppna Inloggning i Webbläsare' för att ansluta automatiskt!"
        )
        ttk.Label(
            main_frame,
            text=instruction_text,
            wraplength=560,
            font=('Segoe UI', 9)
        ).grid(row=1, column=0, columnspan=2, sticky=tk.W, pady=(0, 8))

        # Clickable Link Label
        def _open_strava_dev_url(event=None):
            import webbrowser
            webbrowser.open("https://www.strava.com/settings/api")

        link_label = tk.Label(
            main_frame,
            text="🔗 Öppna Strava Developer Portal (strava.com/settings/api)",
            font=('Segoe UI', 9, 'underline', 'bold'),
            fg='#FC4C02',
            bg=self.colors.get('card_bg', '#FFFFFF'),
            cursor="hand2"
        )
        link_label.grid(row=2, column=0, columnspan=2, sticky=tk.W, pady=(0, 12))
        link_label.bind("<Button-1>", _open_strava_dev_url)

        # Fields
        ttk.Label(main_frame, text="Client ID:", font=('Segoe UI', 10, 'bold')).grid(row=3, column=0, sticky=tk.W, pady=5)
        self.client_id_var = tk.StringVar(value=client_id)
        ttk.Entry(main_frame, textvariable=self.client_id_var, width=42).grid(row=3, column=1, sticky=(tk.W, tk.E), pady=5)

        ttk.Label(main_frame, text="Client Secret:", font=('Segoe UI', 10, 'bold')).grid(row=4, column=0, sticky=tk.W, pady=5)
        self.client_secret_var = tk.StringVar(value=client_secret)
        ttk.Entry(main_frame, textvariable=self.client_secret_var, width=42, show="*").grid(row=4, column=1, sticky=(tk.W, tk.E), pady=5)

        ttk.Label(main_frame, text="Callback URL:", font=('Segoe UI', 10, 'bold')).grid(row=5, column=0, sticky=tk.W, pady=5)
        self.callback_url_var = tk.StringVar(value="http://localhost:8081/")
        ttk.Entry(main_frame, textvariable=self.callback_url_var, width=42).grid(row=5, column=1, sticky=(tk.W, tk.E), pady=5)

        ttk.Label(main_frame, text="Auth Code / Token:", font=('Segoe UI', 10, 'bold')).grid(row=6, column=0, sticky=tk.W, pady=5)
        self.code_var = tk.StringVar()
        ttk.Entry(main_frame, textvariable=self.code_var, width=42).grid(row=6, column=1, sticky=(tk.W, tk.E), pady=5)

        # Buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=7, column=0, columnspan=2, pady=(15, 0))

        ttk.Button(btn_frame, text="▶ Öppna Inloggning i Webbläsare", command=self._open_browser_auth, style='Modern.TButton').pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_frame, text="💾 Spara & Anslut", command=self.save, style='Accent.TButton').pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_frame, text="❌ Avbryt", command=self.destroy, style='Modern.TButton').pack(side=tk.LEFT, padx=4)

    def _open_browser_auth(self):
        cid = self.client_id_var.get().strip()

        if not cid:
            messagebox.showwarning("Client ID Saknas", "Fyll i ditt Client ID först innan du klickar på inloggning.", parent=self)
            return

        import http.server
        import urllib.parse
        import threading
        import webbrowser

        port = 8081
        dialog_ref = self

        class StravaCallbackHandler(http.server.BaseHTTPRequestHandler):
            def do_GET(self):
                parsed_path = urllib.parse.urlparse(self.path)
                query_params = urllib.parse.parse_qs(parsed_path.query)
                code = query_params.get('code', [None])[0]

                self.send_response(200)
                self.send_header('Content-type', 'text/html; charset=utf-8')
                self.end_headers()

                if code:
                    html = """
                    <html>
                    <body style="font-family: Segoe UI, sans-serif; text-align: center; padding-top: 50px; background: #F3F4F6;">
                        <div style="background: white; max-width: 500px; margin: 0 auto; padding: 40px; border-radius: 12px; box-shadow: 0 4px 12px rgba(0,0,0,0.1);">
                            <h2 style="color: #FC4C02; margin-bottom: 10px;">✅ Strava Ansluten!</h2>
                            <p style="color: #4B5563; font-size: 16px;">Ditt Strava-konto har anslutits framgångsrikt till HealthChat.</p>
                            <p style="color: #6B7280; font-size: 14px;">Du kan nu stänga den här fliken i webbläsaren!</p>
                        </div>
                    </body>
                    </html>
                    """
                    self.wfile.write(html.encode('utf-8'))

                    def _update_dialog():
                        try:
                            if dialog_ref.winfo_exists():
                                dialog_ref.code_var.set(code)
                                dialog_ref.save()
                        except Exception:
                            pass
                    dialog_ref.after(100, _update_dialog)
                else:
                    self.wfile.write("<html><body><h2>Ingen kod mottogs</h2></body></html>".encode('utf-8'))

            def log_message(self, format, *args):
                pass

        try:
            server = http.server.HTTPServer(('127.0.0.1', port), StravaCallbackHandler)
            server.timeout = 120
            threading.Thread(target=lambda: (server.handle_request(), server.server_close()), daemon=True).start()
        except Exception as e:
            logger.warning(f"Could not start local callback server on port {port}: {e}")

        auth_url = f"https://www.strava.com/oauth/authorize?client_id={cid}&response_type=code&redirect_uri=http%3A%2F%2Flocalhost%3A8081%2F&approval_prompt=auto&scope=read%2Cactivity%3Aread_all"
        webbrowser.open(auth_url)
        messagebox.showinfo(
            "Strava Inloggning",
            "Öppnar Strava-auktoriseringssidan i webbläsaren.\n\n"
            "1. Logga in och godkänn behörigheterna.\n"
            "2. Appen tar emot koden och genomför anslutningen helt automatiskt!",
            parent=self
        )

    def save(self):
        self.result = {
            'client_id': self.client_id_var.get().strip(),
            'client_secret': self.client_secret_var.get().strip(),
            'code': self.code_var.get().strip()
        }
        self.destroy()


class WithingsConnectDialog(tk.Toplevel):
    """Dedicated modal dialog to connect Withings Health Mate API account with instructions and link."""

    def __init__(self, parent, client_id="", client_secret="", refresh_token="", colors=None):
        super().__init__(parent)
        self.title("Anslut till Withings Health Mate")
        self.geometry("580x560")
        self.resizable(False, False)
        self.result = None
        self.colors = colors or {'bg': '#F3F4F6', 'card_bg': '#FFFFFF', 'text': '#1F2937', 'accent': '#0078D4'}

        self.transient(parent)
        self.grab_set()

        self.configure(bg=self.colors.get('bg', '#F3F4F6'))

        main_frame = ttk.Frame(self, padding="20", style='Card.TFrame')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        main_frame.columnconfigure(1, weight=1)

        # Header
        ttk.Label(
            main_frame,
            text="⚖️ Anslut till Withings Developer",
            font=('Segoe UI', 14, 'bold'),
            foreground=self.colors.get('accent', '#0078D4')
        ).grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=(0, 8))

        # Instructions
        info_text = (
            "Hur du ansluter Withings:\n"
            "1. Kopiera Client ID och Secret från utvecklarportalen (developer.withings.com).\n"
            "2. Ange Callback URL: http://localhost:8000 i din app på developer.withings.com.\n"
            "3. Klicka på '▶ Öppna Inloggning i Webbläsare' nedan för att godkänna appen.\n"
            "4. Klistra in koden (eller hela URL:en du kopierar från adressfältet) nedan.\n"
            "   (Appen rensar och extraherar den rena koden mellan code= och &state= helt automatiskt!)"
        )
        ttk.Label(
            main_frame,
            text=info_text,
            wraplength=520,
            font=('Segoe UI', 9)
        ).grid(row=1, column=0, columnspan=2, sticky=tk.W, pady=(0, 8))

        # Clickable Link
        def _open_withings_dev_url(event=None):
            import webbrowser
            webbrowser.open("https://developer.withings.com/dashboard/")

        link_label = tk.Label(
            main_frame,
            text="🔗 Öppna Withings Developer Portal (developer.withings.com/dashboard)",
            font=('Segoe UI', 9, 'underline'),
            fg='#0078D4',
            bg=self.colors.get('card_bg', '#FFFFFF'),
            cursor="hand2"
        )
        link_label.grid(row=2, column=0, columnspan=2, sticky=tk.W, pady=(0, 12))
        link_label.bind("<Button-1>", _open_withings_dev_url)

        # Fields
        ttk.Label(main_frame, text="Client ID:", font=('Segoe UI', 10, 'bold')).grid(row=3, column=0, sticky=tk.W, pady=5)
        self.client_id_var = tk.StringVar(value=client_id)
        ttk.Entry(main_frame, textvariable=self.client_id_var, width=42).grid(row=3, column=1, sticky=(tk.W, tk.E), pady=5)

        ttk.Label(main_frame, text="Client Secret:", font=('Segoe UI', 10, 'bold')).grid(row=4, column=0, sticky=tk.W, pady=5)
        self.client_secret_var = tk.StringVar(value=client_secret)
        ttk.Entry(main_frame, textvariable=self.client_secret_var, width=42, show="*").grid(row=4, column=1, sticky=(tk.W, tk.E), pady=5)

        ttk.Label(main_frame, text="Callback URL:", font=('Segoe UI', 10, 'bold')).grid(row=5, column=0, sticky=tk.W, pady=5)
        self.callback_url_var = tk.StringVar(value="http://localhost:8000")
        ttk.Entry(main_frame, textvariable=self.callback_url_var, width=42).grid(row=5, column=1, sticky=(tk.W, tk.E), pady=5)

        ttk.Label(main_frame, text="Auth Code / Token:", font=('Segoe UI', 10, 'bold')).grid(row=6, column=0, sticky=tk.W, pady=5)
        self.code_or_token_var = tk.StringVar(value=refresh_token)
        ttk.Entry(main_frame, textvariable=self.code_or_token_var, width=42).grid(row=6, column=1, sticky=(tk.W, tk.E), pady=5)

        # Buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=7, column=0, columnspan=2, pady=(15, 0))

        ttk.Button(btn_frame, text="▶ Öppna Inloggning i Webbläsare", command=self._open_browser_auth, style='Modern.TButton').pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_frame, text="💾 Spara & Anslut", command=self.save, style='Accent.TButton').pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_frame, text="❌ Avbryt", command=self.destroy, style='Modern.TButton').pack(side=tk.LEFT, padx=4)

    def _open_browser_auth(self):
        cid = self.client_id_var.get().strip()
        csecret = self.client_secret_var.get().strip()
        cb_url = self.callback_url_var.get().strip() or "http://localhost:8000"

        if not cid:
            messagebox.showwarning("Client ID Saknas", "Fyll i ditt Client ID först innan du klickar på inloggning.", parent=self)
            return

        import http.server
        import urllib.parse
        import threading
        import webbrowser

        port = 8000
        if ':8000' in cb_url:
            port = 8000
        elif ':8080' in cb_url:
            port = 8080

        dialog_ref = self

        class OAuthCallbackHandler(http.server.BaseHTTPRequestHandler):
            def do_GET(self):
                parsed_path = urllib.parse.urlparse(self.path)
                query_params = urllib.parse.parse_qs(parsed_path.query)
                code = query_params.get('code', [None])[0]

                self.send_response(200)
                self.send_header('Content-type', 'text/html; charset=utf-8')
                self.end_headers()

                if code:
                    try:
                        handler = WithingsDataHandler(client_id=cid, client_secret=csecret)
                        res = handler.exchange_code_for_token(code, cid, csecret, redirect_uri=cb_url)
                        ref_tok = res.get("refresh_token", "")
                        acc_tok = res.get("access_token", "")

                        html = """
                        <html>
                        <body style="font-family: Segoe UI, sans-serif; text-align: center; padding-top: 50px; background: #F3F4F6;">
                            <div style="background: white; max-width: 500px; margin: 0 auto; padding: 40px; border-radius: 12px; box-shadow: 0 4px 12px rgba(0,0,0,0.1);">
                                <h2 style="color: #10B981; margin-bottom: 10px;">✅ Withings Ansluten!</h2>
                                <p style="color: #4B5563; font-size: 16px;">Din Withings-våg har anslutits framgångsrikt till HealthChat.</p>
                                <p style="color: #6B7280; font-size: 14px;">Du kan nu stänga den här fliken i webbläsaren!</p>
                            </div>
                        </body>
                        </html>
                        """
                        self.wfile.write(html.encode('utf-8'))

                        def _update_dialog():
                            try:
                                if dialog_ref.winfo_exists():
                                    dialog_ref.code_or_token_var.set(ref_tok)
                                    dialog_ref.result = {
                                        'client_id': cid,
                                        'client_secret': csecret,
                                        'callback_url': cb_url,
                                        'code_or_token': ref_tok,
                                        'access_token': acc_tok,
                                        'refresh_token': ref_tok,
                                        'auto_exchanged': True
                                    }
                                    dialog_ref.destroy()
                            except Exception:
                                pass
                        dialog_ref.after(100, _update_dialog)
                    except Exception as ex:
                        err_html = f"""
                        <html>
                        <body style="font-family: Segoe UI, sans-serif; text-align: center; padding-top: 50px; background: #F3F4F6;">
                            <div style="background: white; max-width: 500px; margin: 0 auto; padding: 40px; border-radius: 12px; box-shadow: 0 4px 12px rgba(0,0,0,0.1);">
                                <h2 style="color: #EF4444; margin-bottom: 10px;">❌ Fel vid anslutning</h2>
                                <p style="color: #4B5563; font-size: 14px;">{ex}</p>
                            </div>
                        </body>
                        </html>
                        """
                        self.wfile.write(err_html.encode('utf-8'))
                else:
                    self.wfile.write("<html><body><h2>Ingen kod mottogs</h2></body></html>".encode('utf-8'))

            def log_message(self, format, *args):
                pass

        try:
            server = http.server.HTTPServer(('127.0.0.1', port), OAuthCallbackHandler)
            server.timeout = 120
            threading.Thread(target=lambda: (server.handle_request(), server.server_close()), daemon=True).start()
        except Exception as e:
            logger.warning(f"Could not start local callback server on port {port}: {e}")

        auth_url = WithingsDataHandler.get_auth_url(cid, redirect_uri=cb_url)
        webbrowser.open(auth_url)
        messagebox.showinfo(
            "Withings Inloggning",
            "Öppnar Withings auktoriseringssida i webbläsaren.\n\n"
            "1. Logga in och klicka på 'Allow' i webbläsaren.\n"
            "2. Appen tar emot koden och genomför anslutningen helt automatiskt!\n\n"
            "(Om du vill klistra in manuellt kan du kopiera hela länken eller koden mellan code= och &state= och klicka på 'Spara & Anslut').",
            parent=self
        )

    def save(self):
        self.result = {
            'client_id': self.client_id_var.get().strip(),
            'client_secret': self.client_secret_var.get().strip(),
            'callback_url': self.callback_url_var.get().strip() or "http://localhost:8000",
            'code_or_token': self.code_or_token_var.get().strip()
        }
        self.destroy()


class HealthChatApp:
    """Main application class for HealthChat desktop app"""
    
    def __init__(self, root):
        """Initialize the application"""
        self.root = root
        self.root.title("HealthChat")
        self.root.geometry("1650x950")  # Expanded for side-by-side charts & chat view
        self.db = GarminDatabase()
        
        # Set window icon (works in both script and exe)
        try:
            # Get the correct base path for PyInstaller exe
            if getattr(sys, 'frozen', False):
                # Running as compiled executable
                base_path = Path(sys._MEIPASS)
            else:
                # Running as script
                base_path = Path(__file__).parent
            
            icon_path = base_path / "logo.ico"
            if icon_path.exists():
                self.root.iconbitmap(str(icon_path))
        except Exception as e:
            logger.debug(f"Could not load icon: {e}")
        
        # Set minimum window size
        self.root.minsize(1200, 750)  # Expanded minimum size for side-by-side view
        
        # Configuration file path
        old_config_dir = Path.home() / ".garmin_chat"
        self.config_dir = Path.home() / ".healthchat"
        if old_config_dir.exists() and not self.config_dir.exists():
            import shutil
            try:
                shutil.copytree(old_config_dir, self.config_dir)
                logger.info("Migrated configuration from ~/.garmin_chat to ~/.healthchat")
            except Exception as e:
                logger.warning(f"Could not migrate config directory: {e}")
        self.config_dir.mkdir(exist_ok=True)
        self.config_file = self.config_dir / "config.json"
        self.saved_prompts_file = self.config_dir / "saved_prompts.json"
        self.chat_history_dir = self.config_dir / "chat_history"
        self.chat_history_dir.mkdir(exist_ok=True)
        
        # Chat history for current session
        self.current_chat_history = []
        
        # Conversation context memory (last 10 messages for AI context)
        self.conversation_context = []
        self.max_context_messages = 10
        
        # User preferences learned from conversations
        self.user_preferences = {
            'favorite_activities': [],
            'goals': [],
            'interests': [],
            'last_queries': []
        }
        
        # Load conversation history
        self.load_conversation_history()
        
        # Application state
        self.garmin_handler = None
        self.fitbit_handler = FitbitHandler(db=self.db, token_store_dir=self.config_dir)
        self.strava_handler = StravaHandler(db=self.db, token_store_dir=self.config_dir)
        self.ai_client = None  # Changed from xai_client to ai_client
        self.authenticated = False
        self.mfa_required = False
        
        # AI Provider settings (support multiple providers)
        self.ai_provider = 'xai'  # Default provider
        self.xai_api_key = None
        self.xai_model = 'grok-3'
        self.openai_api_key = None
        self.openai_model = 'gpt-4o'
        self.azure_api_key = None
        self.azure_endpoint = None
        self.azure_deployment = None
        self.gemini_api_key = None
        self.gemini_model = 'gemini-1.5-flash'
        self.anthropic_api_key = None
        self.anthropic_model = 'claude-sonnet-4-6'
        self.ollama_api_key = ''  # Not needed, kept for consistency
        self.ollama_model = 'llama3.2'
        self.ollama_base_url = 'http://localhost:11434/v1'
        
        # Garmin settings
        self.garmin_email = None
        self.garmin_password = None

        # Withings settings
        self.withings_client_id = ''
        self.withings_client_secret = ''
        self.withings_refresh_token = ''
        self.withings_access_token = ''

        # Strava settings
        self.strava_client_id = ''
        self.strava_client_secret = ''
        self.strava_refresh_token = ''
        self.strava_access_token = ''

        self.auto_login = True  # Default to auto-login enabled
        self.dark_mode = False  # Start in light mode
        self.window_state_restored = False  # Track if window position was restored
        
        # Load configuration
        self.load_config()
        
        # Configure style
        self.setup_styles()
        
        # Create UI
        self.create_widgets()
        
        # Center window on screen (only if no saved state was restored)
        if not self.window_state_restored:
            self.center_window()
        
        # Set up window close handler to save state
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        
        # Check if credentials are configured
        current_ai_key = self.get_current_ai_key()
        if not current_ai_key or not self.garmin_email or not self.garmin_password:
            self.root.after(100, self.prompt_for_credentials)
        elif self.auto_login:
            # Auto-connect if credentials are configured and auto-login is enabled
            self.root.after(500, self.auto_connect)
        
    def load_config(self):
        """Load configuration from file"""
        try:
            if self.config_file.exists():
                with open(self.config_file, 'r') as f:
                    config = json.load(f)
                    
                    # AI Provider settings
                    self.ai_provider = config.get('ai_provider', 'xai')
                    self.xai_api_key = config.get('xai_api_key', '')
                    self.xai_model = config.get('xai_model', 'grok-3')
                    self.openai_api_key = config.get('openai_api_key', '')
                    self.openai_model = config.get('openai_model', 'gpt-4o')
                    self.azure_api_key = config.get('azure_api_key', '')
                    self.azure_endpoint = config.get('azure_endpoint', '')
                    self.azure_deployment = config.get('azure_deployment', '')
                    self.gemini_api_key = config.get('gemini_api_key', '')
                    self.gemini_model = config.get('gemini_model', 'gemini-1.5-flash')
                    self.anthropic_api_key = config.get('anthropic_api_key', '')
                    self.anthropic_model = config.get('anthropic_model', 'claude-sonnet-4-6')
                    self.ollama_model = config.get('ollama_model', 'llama3.2')
                    self.ollama_base_url = config.get('ollama_base_url', 'http://localhost:11434/v1')
                    
                    model_migrations = {
                        # Gemini migrations (old alias cleanups)
                        'gemini-1.5-pro-latest': 'gemini-1.5-pro',
                        'gemini-1.5-flash-latest': 'gemini-1.5-flash',
                        # xAI migrations
                        'grok-beta': 'grok-3',
                        'grok-2-1212': 'grok-3',
                        # Add more migrations as models get deprecated
                    }

                    # Migrate Anthropic model to latest if using older default
                    anthropic_model_migrations = {
                        'claude-sonnet-4-5-20250929': 'claude-sonnet-4-6',
                        'claude-opus-4-5-20251101': 'claude-opus-4-6',
                        'claude-3-5-haiku-20241022': 'claude-haiku-4-5-20251001',
                    }
                    if self.anthropic_model in anthropic_model_migrations:
                        old_model = self.anthropic_model
                        self.anthropic_model = anthropic_model_migrations[old_model]
                        logger.info(f"Auto-migrated Anthropic model: {old_model} → {self.anthropic_model}")
                    
                    # Migrate Gemini model if deprecated
                    if self.gemini_model in model_migrations:
                        old_model = self.gemini_model
                        self.gemini_model = model_migrations[old_model]
                        logger.info(f"Auto-migrated Gemini model: {old_model} → {self.gemini_model}")
                    
                    # Migrate xAI model if deprecated
                    if self.xai_model in model_migrations:
                        old_model = self.xai_model
                        self.xai_model = model_migrations[old_model]
                        logger.info(f"Auto-migrated xAI model: {old_model} → {self.xai_model}")
                    
                    # Garmin settings
                    self.garmin_email = config.get('garmin_email', '')
                    self.garmin_password = config.get('garmin_password', '')
                    
                    # Withings settings
                    self.withings_client_id = config.get('withings_client_id', '')
                    self.withings_client_secret = config.get('withings_client_secret', '')
                    self.withings_refresh_token = config.get('withings_refresh_token', '')
                    self.withings_access_token = config.get('withings_access_token', '')

                    # Strava settings
                    self.strava_client_id = config.get('strava_client_id', '')
                    self.strava_client_secret = config.get('strava_client_secret', '')
                    self.strava_refresh_token = config.get('strava_refresh_token', '')
                    self.strava_access_token = config.get('strava_access_token', '')

                    self.auto_login = config.get('auto_login', True)
                    self.dark_mode = config.get('dark_mode', False)
                    
                    # Window state (position and size)
                    window_state = config.get('window_state', {})
                    if window_state:
                        try:
                            width = max(window_state.get('width', 1650), 1650)
                            height = max(window_state.get('height', 950), 900)
                            x = window_state.get('x')
                            y = window_state.get('y')
                            
                            # Apply saved size
                            if x is not None and y is not None:
                                # Validate position is on screen
                                screen_width = self.root.winfo_screenwidth()
                                screen_height = self.root.winfo_screenheight()
                                
                                # Ensure window is visible
                                if x + width > screen_width:
                                    x = screen_width - width - 10
                                if y + height > screen_height:
                                    y = screen_height - height - 60  # Account for taskbar
                                if x < 0:
                                    x = 10
                                if y < 0:
                                    y = 10
                                
                                self.root.geometry(f'{width}x{height}+{x}+{y}')
                                self.window_state_restored = True
                                logger.info(f"Restored window state: {width}x{height}+{x}+{y}")
                        except Exception as e:
                            logger.warning(f"Could not restore window state: {e}")
                            # Will use center_window() as fallback
                    
                    logger.info("Configuration loaded")
        except Exception as e:
            logger.error(f"Error loading config: {e}")
            self.ai_provider = 'xai'
            self.xai_api_key = None
            self.gemini_model = 'gemini-1.5-flash'
            self.xai_model = 'grok-3'
            self.garmin_email = None
            self.garmin_password = None
            self.auto_login = True
            self.dark_mode = False  # Default to light mode on error
            
    def save_config(self):
        """Save configuration to file safely from any thread."""
        try:
            import threading
            # Only query Tkinter widget on the main UI thread
            if threading.current_thread() is threading.main_thread():
                try:
                    geometry = self.root.geometry()
                    match = geometry.split('+')
                    size = match[0].split('x')
                    self.last_window_state = {
                        'width': int(size[0]),
                        'height': int(size[1]),
                        'x': int(match[1]) if len(match) > 1 else None,
                        'y': int(match[2]) if len(match) > 2 else None
                    }
                except Exception:
                    pass

            window_state = getattr(self, 'last_window_state', {})
            
            config = {
                # AI Provider settings
                'ai_provider': self.ai_provider,
                'xai_api_key': self.xai_api_key or '',
                'xai_model': self.xai_model,
                'openai_api_key': self.openai_api_key or '',
                'openai_model': self.openai_model,
                'azure_api_key': self.azure_api_key or '',
                'azure_endpoint': self.azure_endpoint or '',
                'azure_deployment': self.azure_deployment or '',
                'gemini_api_key': self.gemini_api_key or '',
                'gemini_model': self.gemini_model,
                'anthropic_api_key': self.anthropic_api_key or '',
                'anthropic_model': self.anthropic_model,
                'ollama_model': self.ollama_model,
                'ollama_base_url': self.ollama_base_url,
                # Garmin settings
                'garmin_email': self.garmin_email,
                'garmin_password': self.garmin_password,
                # Withings settings
                'withings_client_id': self.withings_client_id or '',
                'withings_client_secret': self.withings_client_secret or '',
                'withings_refresh_token': self.withings_refresh_token or '',
                'withings_access_token': self.withings_access_token or '',
                # Strava settings
                'strava_client_id': getattr(self, 'strava_client_id', '') or '',
                'strava_client_secret': getattr(self, 'strava_client_secret', '') or '',
                'strava_refresh_token': getattr(self, 'strava_refresh_token', '') or '',
                'strava_access_token': getattr(self, 'strava_access_token', '') or '',
                'auto_login': self.auto_login,
                'dark_mode': self.dark_mode,
                # Window state
                'window_state': window_state
            }
            with open(self.config_file, 'w') as f:
                json.dump(config, f, indent=2)
            logger.info("Configuration saved")
        except Exception as e:
            logger.error(f"Error saving config: {e}")
    
    def on_closing(self):
        """Handle window close event - save state and exit"""
        try:
            # Save current configuration including window state
            self.save_config()
        except Exception as e:
            logger.error(f"Error saving config on close: {e}")
        finally:
            # Close the application
            self.root.destroy()
            
    def prompt_for_credentials(self):
        """Prompt user to enter credentials on first run"""
        messagebox.showinfo(
            "Welcome to HealthChat",
            "Welcome! Before you can start chatting, you need to configure your credentials.\n\n"
            "You'll need:\n"
            "1. An AI API key (xAI, OpenAI, Gemini, Anthropic, or Azure)\n"
            "2. Your Garmin Connect email and password\n\n"
            "Click OK to open the settings dialog.",
            parent=self.root
        )
        self.open_settings()
        
    def auto_connect(self):
        """Automatically connect to Garmin on startup if credentials are configured"""
        self.add_message("System", "Auto-connecting to Garmin Connect...", 'system')
        self.update_status("Connecting to Garmin...", False)
        self.connect_to_garmin()
        
    def setup_styles(self):
        """Configure ttk styles for modern Fluent Design look"""
        # Apply colors based on saved dark_mode preference
        if self.dark_mode:
            # Dark mode colors
            self.colors = {
                'bg': '#202020',
                'card_bg': '#2D2D30',
                'accent': '#60A5FA',
                'accent_hover': '#3B82F6',
                'accent_light': '#1E3A5F',
                'text': '#E5E5E5',
                'text_secondary': '#A0A0A0',
                'border': '#3E3E42',
                'success': '#10B981',
                'warning': '#F59E0B',
                'shadow': '#00000040',
            }
        else:
            # Light mode colors (default)
            self.colors = {
                'bg': '#F3F3F3',            # Light gray background
                'card_bg': '#FFFFFF',        # White cards
                'accent': '#0078D4',         # Windows 11 blue
                'accent_hover': '#106EBE',   # Darker blue on hover
                'accent_light': '#E6F2FA',   # Light blue background
                'text': '#1F1F1F',           # Almost black text
                'text_secondary': '#605E5C', # Gray text
                'border': '#EDEBE9',         # Light border
                'success': '#107C10',        # Green
                'warning': '#D83B01',        # Red/Orange
                'shadow': '#00000010',       # Subtle shadow
            }
        
        # Configure root window
        self.root.configure(bg=self.colors['bg'])
        
        # Modern TTK styles
        style = ttk.Style()
        style.theme_use('clam')
        
        # Base frame style
        style.configure('TFrame',
                       background=self.colors['bg'])
        
        # Card frame style (elevated white cards)
        style.configure('Card.TFrame',
                       background=self.colors['card_bg'],
                       relief='flat',
                       borderwidth=1)
        
        # Modern button styles
        style.configure('Modern.TButton',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       bordercolor=self.colors['border'],
                       borderwidth=1,
                       relief='flat',
                       padding=(12, 6),
                       font=('Segoe UI', 11))  # Larger font for emojis
        style.map('Modern.TButton',
                 background=[('active', self.colors['accent_light'] if self.dark_mode else self.colors['border']), 
                           ('pressed', self.colors['accent_light'] if self.dark_mode else self.colors['border'])],
                 foreground=[('active', self.colors['accent'] if self.dark_mode else self.colors['text']),
                           ('pressed', self.colors['accent'] if self.dark_mode else self.colors['text'])])
        
        # Accent button (primary action)
        style.configure('Accent.TButton',
                       background=self.colors['accent'],
                       foreground='white',
                       borderwidth=0,
                       relief='flat',
                       padding=(16, 8),
                       font=('Segoe UI', 10, 'bold'))
        style.map('Accent.TButton',
                 background=[('active', self.colors['accent_hover']), 
                           ('pressed', self.colors['accent_hover'])],
                 foreground=[('active', 'white'), ('pressed', 'white')])
        
        # Quick Question button (rounded hover effect)
        style.configure('QuickQuestion.TButton',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       bordercolor=self.colors['border'],
                       borderwidth=1,
                       relief='flat',
                       padding=(12, 8),
                       font=('Segoe UI', 10))
        style.map('QuickQuestion.TButton',
                 background=[('active', self.colors['accent_light']), 
                           ('pressed', self.colors['accent'])],
                 foreground=[('active', self.colors['accent']),
                           ('pressed', 'white')],
                 relief=[('active', 'raised'),  # Gives a subtle rounded effect
                        ('pressed', 'sunken')])
        
        # Label styles
        style.configure('Title.TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 18, 'bold'))
        
        style.configure('Heading.TLabel',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 11, 'bold'))
        
        style.configure('TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 10))
        
        style.configure('Status.TLabel',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text_secondary'],
                       font=('Segoe UI', 9))
        
        # Entry style
        style.configure('Modern.TEntry',
                       fieldbackground=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       bordercolor=self.colors['border'],
                       borderwidth=1,
                       font=('Segoe UI', 10))
        
        # LabelFrame style (card with title)
        style.configure('Card.TLabelframe',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       bordercolor=self.colors['border'],
                       borderwidth=1,
                       relief='flat')
        style.configure('Card.TLabelframe.Label',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 10, 'bold'))
        
    def center_window(self):
        """Center the window on screen, accounting for taskbar"""
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        
        # Get screen dimensions
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        
        # Calculate center position
        x = (screen_width // 2) - (width // 2)
        y = (screen_height // 2) - (height // 2)
        
        # Ensure window doesn't go off-screen or under taskbar
        # Account for typical taskbar height (40-50 pixels)
        taskbar_height = 50
        if y + height > screen_height - taskbar_height:
            y = screen_height - height - taskbar_height - 10  # 10px margin
        
        # Ensure window doesn't go above screen
        if y < 0:
            y = 10
        
        # Ensure window doesn't go off sides
        if x < 0:
            x = 10
        if x + width > screen_width:
            x = screen_width - width - 10
        
        self.root.geometry(f'{width}x{height}+{x}+{y}')

    def create_menu_bar(self):
        """Create Windows drop-down menu bar (Arkiv, Garmin, Verktyg, Hjälp)."""
        menubar = tk.Menu(self.root)
        
        # 1. Arkiv (File) Menu
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="📥 Kör Check-in (Alla Källor)", command=self.perform_unified_checkin)
        file_menu.add_command(label="🔄 Synka Full Historik (Alla Källor)", command=self.perform_full_historical_sync)
        file_menu.add_separator()
        file_menu.add_command(label="💾 Spara Chat...", command=self.save_chat_history)
        file_menu.add_command(label="📄 Exportera Rapport...", command=self.export_conversation_report)
        file_menu.add_separator()
        file_menu.add_command(label="⚙️ Inställningar", command=self.open_settings)
        file_menu.add_separator()
        file_menu.add_command(label="🚪 Avsluta", command=self.root.quit)
        menubar.add_cascade(label="Arkiv", menu=file_menu)
        
        # 2. Garmin Menu
        garmin_menu = tk.Menu(menubar, tearoff=0)
        garmin_menu.add_command(label="▶ Anslut till Garmin Connect", command=self.connect_to_garmin)
        garmin_menu.add_command(label="📥 Kör Check-in (Garmin)", command=self.perform_garmin_checkin)
        garmin_menu.add_command(label="🔄 Full Synk (365 Dagar)", command=self.perform_full_historical_sync)
        menubar.add_cascade(label="Garmin", menu=garmin_menu)
        
        # 3. Fitbit Menu
        fitbit_menu = tk.Menu(menubar, tearoff=0)
        fitbit_menu.add_command(label="▶ Anslut till Fitbit...", command=self.connect_to_fitbit)
        fitbit_menu.add_command(label="📥 Fitbit Check-in", command=self.perform_fitbit_checkin)
        fitbit_menu.add_command(label="📁 Importera Fitbit Export-fil (CSV/JSON)...", command=self.import_fitbit_file)
        menubar.add_cascade(label="Fitbit", menu=fitbit_menu)
        
        # 4. Withings Menu
        withings_menu = tk.Menu(menubar, tearoff=0)
        withings_menu.add_command(label="▶ Anslut till Withings...", command=self.connect_to_withings)
        withings_menu.add_command(label="📥 Withings Check-in", command=self.perform_withings_checkin)
        withings_menu.add_command(label="📁 Importera Withings Export-fil (CSV/JSON)...", command=self.import_withings_file)
        menubar.add_cascade(label="Withings", menu=withings_menu)
        
        # 5. Strava Menu
        strava_menu = tk.Menu(menubar, tearoff=0)
        strava_menu.add_command(label="▶ Anslut till Strava...", command=self.connect_to_strava)
        strava_menu.add_command(label="📥 Strava Check-in", command=self.perform_strava_checkin)
        strava_menu.add_command(label="📁 Importera Strava Export-fil (CSV/ZIP/GPX/FIT)...", command=self.import_strava_file)
        menubar.add_cascade(label="Strava", menu=strava_menu)
        
        # 3. Verktyg (Tools) Menu
        tools_menu = tk.Menu(menubar, tearoff=0)
        tools_menu.add_command(label="🔍 Sök i Chatthistorik", command=self.open_search)
        tools_menu.add_command(label="📝 Sparade Promptar", command=self.open_saved_prompts)
        tools_menu.add_command(label="📂 Chat-historik", command=self.open_chat_history_viewer)
        tools_menu.add_command(label="↺ Rensa Chat", command=self.reset_chat)
        tools_menu.add_separator()
        tools_menu.add_command(label="◐ Växla Mörkt/Ljust Tema", command=self.toggle_theme)
        menubar.add_cascade(label="Verktyg", menu=tools_menu)
        
        # 4. Hjälp (Help) Menu
        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(label="❓ Om HealthChat", command=self.show_about)
        menubar.add_cascade(label="Hjälp", menu=help_menu)
        
        self.root.config(menu=menubar)

    def show_about(self):
        """Show About HealthChat dialog"""
        messagebox.showinfo(
            "Om HealthChat Desktop",
            f"HealthChat Desktop v{APP_VERSION}\n\n"
            "En lokal AI-assistent för analys av Garmin Connect hälsodata.\n"
            "• Lokal SQLite-databas för hälsodata & historik\n"
            "• Interaktiva grafer för sömn, stress, Body Battery & träning\n"
            "• 100% GDPR-säker lokal lagring",
            parent=self.root
        )

    def perform_unified_checkin(self):
        """Perform Check-in across all configured health services (Garmin, Fitbit, etc.)."""
        configured_sources = []
        
        has_garmin = hasattr(self, 'garmin_handler') and self.garmin_handler and self.authenticated
        has_fitbit = hasattr(self, 'fitbit_handler') and self.fitbit_handler and self.fitbit_handler.is_authenticated()
        has_withings = bool(self.withings_client_id and self.withings_client_secret and self.withings_refresh_token)
        has_strava = hasattr(self, 'strava_handler') and self.strava_handler and self.strava_handler.is_authenticated()
        
        if has_garmin:
            configured_sources.append("Garmin")
        if has_fitbit:
            configured_sources.append("Fitbit")
        if has_withings:
            configured_sources.append("Withings")
        if has_strava:
            configured_sources.append("Strava")

        if not configured_sources:
            messagebox.showwarning(
                "Inga Källor Anslutna",
                "Du är inte ansluten till någon hälsokälla än.\n\nAnslut till Garmin Connect, Fitbit eller Withings via menyn längst upp först.",
                parent=self.root
            )
            return

        if hasattr(self, 'checkin_btn'):
            self.checkin_btn.config(state=tk.DISABLED, text="⏳ Synkar...")
        if hasattr(self, 'charts_view') and hasattr(self.charts_view, 'dashboard_checkin_btn'):
            self.charts_view.dashboard_checkin_btn.config(state=tk.DISABLED, text="⏳ Synkar...")

        source_str = " & ".join(configured_sources)
        self.update_status(f"📥 Kör Check-in för {source_str}...", False)
        if hasattr(self, 'charts_view') and self.charts_view:
            self.charts_view.set_sync_status(f"⏳ Synkroniserar {source_str}...")

        completed = {"count": 0, "total": len(configured_sources)}

        def _on_source_finished(source_name):
            completed["count"] += 1
            if completed["count"] >= completed["total"]:
                def _update_ui():
                    if hasattr(self, 'checkin_btn'):
                        self.checkin_btn.config(state=tk.NORMAL, text="📥 Check-in")
                    if hasattr(self, 'charts_view') and hasattr(self.charts_view, 'dashboard_checkin_btn'):
                        self.charts_view.dashboard_checkin_btn.config(state=tk.NORMAL, text="📥 Check-in")
                    self.update_status(f"✅ Check-in genomförd för {source_str}! Graferna har uppdaterats.", False)
                    if hasattr(self, 'charts_view') and self.charts_view:
                        self.charts_view.set_sync_status(f"✅ Synkning klar ({source_str})!", is_done=True)
                        self.charts_view.refresh_all_views()
                    messagebox.showinfo(
                        "Check-in Genomförd",
                        f"✅ Check-in klar för alla anslutna källor ({source_str})!\n\nSenaste hälsodata har sparats i din lokala databas och graferna har uppdaterats.",
                        parent=self.root
                    )
                self.root.after(0, _update_ui)

        # 1. Garmin sync if configured
        if has_garmin:
            sync_days = max(30, self.charts_view.days_range) if (hasattr(self, 'charts_view') and self.charts_view) else 30
            def _g_prog(curr, total, msg):
                def _update():
                    status_str_garmin = f"⏳ Synkar Garmin ({curr}/{total} d)..."
                    self.update_status(status_str_garmin, False)
                    if hasattr(self, 'charts_view') and self.charts_view:
                        self.charts_view.set_sync_status(status_str_garmin)
                self.root.after(0, _update)
            try:
                self.garmin_handler.sync_garmin_history(days=sync_days, on_progress=_g_prog, on_complete=lambda: _on_source_finished("Garmin"))
            except Exception as e:
                logger.error(f"Garmin check-in error: {e}")
                _on_source_finished("Garmin")

        # 2. Fitbit sync if configured
        if has_fitbit:
            def _f_prog(curr, total, msg):
                def _update():
                    status_str_fitbit = f"⏳ Synkar Fitbit ({curr}/{total} d)..."
                    self.update_status(status_str_fitbit, False)
                    if hasattr(self, 'charts_view') and self.charts_view:
                        self.charts_view.set_sync_status(status_str_fitbit)
                self.root.after(0, _update)
            try:
                self.fitbit_handler.sync_fitbit_history(days=7, on_progress=_f_prog, on_complete=lambda: _on_source_finished("Fitbit"))
            except Exception as e:
                logger.error(f"Fitbit check-in error: {e}")
                _on_source_finished("Fitbit")

        # 3. Withings sync if configured
        if has_withings:
            def _w_thread():
                def _w_update(text):
                    self.update_status(text, False)
                    if hasattr(self, 'charts_view') and self.charts_view:
                        self.charts_view.set_sync_status(text)
                self.root.after(0, lambda: _w_update("⏳ Synkar Withings (vikt & kroppssammansättning)..."))
                try:
                    res = self.sync_withings(days=365)
                    cnt = res.get("count", 0) if isinstance(res, dict) else 0
                    self.root.after(0, lambda: _w_update(f"✅ Withings klar ({cnt} mätvärden sparades)"))
                except Exception as e:
                    logger.error(f"Withings check-in error: {e}")
                finally:
                    _on_source_finished("Withings")
            threading.Thread(target=_w_thread, daemon=True).start()

        # 4. Strava sync if configured
        if has_strava:
            def _s_prog(curr, total, msg):
                def _update():
                    status_str_strava = f"⏳ Synkar Strava ({curr}/{total} pass)..."
                    self.update_status(status_str_strava, False)
                    if hasattr(self, 'charts_view') and self.charts_view:
                        self.charts_view.set_sync_status(status_str_strava)
                self.root.after(0, _update)
            try:
                self.strava_handler.sync_strava_history(days=30, on_progress=_s_prog, on_complete=lambda: _on_source_finished("Strava"))
            except Exception as e:
                logger.error(f"Strava check-in error: {e}")
                _on_source_finished("Strava")

    def perform_full_historical_sync(self):
        """Perform complete historical data sync across all connected services (Garmin, Fitbit, Withings) for up to 365 days / lifetime history and update local database."""
        configured_sources = []
        has_garmin = hasattr(self, 'garmin_handler') and self.garmin_handler and self.authenticated
        has_fitbit = hasattr(self, 'fitbit_handler') and self.fitbit_handler and self.fitbit_handler.is_authenticated()
        has_withings = bool(self.withings_client_id and self.withings_client_secret and self.withings_refresh_token)
        has_strava = hasattr(self, 'strava_handler') and self.strava_handler and self.strava_handler.is_authenticated()

        if has_garmin:
            configured_sources.append("Garmin (Full Historik)")
        if has_fitbit:
            configured_sources.append("Fitbit (Full Historik)")
        if has_withings:
            configured_sources.append("Withings (Full Historik)")
        if has_strava:
            configured_sources.append("Strava (Full Historik)")

        if not configured_sources:
            messagebox.showwarning(
                "Inga Källor Anslutna",
                "Du är inte ansluten till någon hälsokälla än.\n\nAnslut till Garmin Connect, Fitbit eller Withings via menyn längst upp först.",
                parent=self.root
            )
            return

        if not messagebox.askyesno(
            "Synkronisera Full Historik",
            f"Vill du synka och skriva över databasen med full historik från start för:\n• {', '.join(configured_sources)}?\n\nDetta hämtar all tillgänglig hälsodata från ditt konto från start.",
            parent=self.root
        ):
            return

        if hasattr(self, 'checkin_btn'):
            self.checkin_btn.config(state=tk.DISABLED, text="⏳ Full Synk...")
        if hasattr(self, 'charts_view') and hasattr(self.charts_view, 'dashboard_checkin_btn'):
            self.charts_view.dashboard_checkin_btn.config(state=tk.DISABLED, text="⏳ Full Synk...")

        self.update_status("🔄 Synkroniserar full historik för alla källor...", False)
        if hasattr(self, 'charts_view') and self.charts_view:
            self.charts_view.set_sync_status("⏳ Synkroniserar full historik...")

        completed = {"count": 0, "total": len(configured_sources)}

        def _on_source_finished(source_name):
            completed["count"] += 1
            if completed["count"] >= completed["total"]:
                def _update_ui():
                    if hasattr(self, 'checkin_btn'):
                        self.checkin_btn.config(state=tk.NORMAL, text="📥 Check-in")
                    if hasattr(self, 'charts_view') and hasattr(self.charts_view, 'dashboard_checkin_btn'):
                        self.charts_view.dashboard_checkin_btn.config(state=tk.NORMAL, text="📥 Check-in")
                    self.update_status("✅ Full historik har synkroniserats! Databasen har uppdaterats.", False)
                    if hasattr(self, 'charts_view') and self.charts_view:
                        self.charts_view.set_sync_status("✅ Full historik-synk klar!", is_done=True)
                        self.charts_view.set_range(3650)
                    messagebox.showinfo(
                        "Full Historik Synkroniserad",
                        "✅ Fullständig historik från start har hämtats!\n\nAll hälsodata för Garmin, Withings och Fitbit har skrivits över och uppdaterats i din lokala databas. Graferna visar nu hela din historik.",
                        parent=self.root
                    )
                self.root.after(0, _update_ui)

        # 1. Garmin Full Sync (3650 Days / Lifetime)
        if has_garmin:
            def _g_prog(curr, total, msg):
                def _update():
                    status_str = f"⏳ Synkar Garmin full historik ({curr}/{total} d)..."
                    self.update_status(status_str, False)
                    if hasattr(self, 'charts_view') and self.charts_view:
                        self.charts_view.set_sync_status(status_str)
                self.root.after(0, _update)
            try:
                self.garmin_handler.sync_garmin_history(days=3650, on_progress=_g_prog, on_complete=lambda: _on_source_finished("Garmin"))
            except Exception as e:
                logger.error(f"Garmin full sync error: {e}")
                _on_source_finished("Garmin")

        # 2. Fitbit Full Sync
        if has_fitbit:
            def _f_prog(curr, total, msg):
                def _update():
                    status_str = f"⏳ Synkar Fitbit full historik ({curr}/{total} d)..."
                    self.update_status(status_str, False)
                    if hasattr(self, 'charts_view') and self.charts_view:
                        self.charts_view.set_sync_status(status_str)
                self.root.after(0, _update)
            try:
                self.fitbit_handler.sync_fitbit_history(days=365, on_progress=_f_prog, on_complete=lambda: _on_source_finished("Fitbit"))
            except Exception as e:
                logger.error(f"Fitbit full sync error: {e}")
                _on_source_finished("Fitbit")

        # 3. Withings Full Sync (Lifetime)
        if has_withings:
            def _w_thread():
                def _w_update(text):
                    self.update_status(text, False)
                    if hasattr(self, 'charts_view') and self.charts_view:
                        self.charts_view.set_sync_status(text)
                self.root.after(0, lambda: _w_update("⏳ Synkar Withings full historik..."))
                try:
                    res = self.sync_withings(days=3650)
                    cnt = res.get("count", 0) if isinstance(res, dict) else 0
                    self.root.after(0, lambda: _w_update(f"✅ Withings full synk klar ({cnt} mätvärden sparades)"))
                except Exception as e:
                    logger.error(f"Withings full sync error: {e}")
                finally:
                    _on_source_finished("Withings")
            threading.Thread(target=_w_thread, daemon=True).start()

    def connect_to_withings(self):
        """Open dedicated popup dialog to configure Withings API credentials."""
        dialog = WithingsConnectDialog(
            self.root,
            client_id=self.withings_client_id,
            client_secret=self.withings_client_secret,
            refresh_token=self.withings_refresh_token,
            colors=self.colors
        )
        self.root.wait_window(dialog)

        if dialog.result:
            cid = dialog.result.get('client_id', '')
            csecret = dialog.result.get('client_secret', '')
            cb_url = dialog.result.get('callback_url', 'http://localhost:8000')
            raw_code = dialog.result.get('code_or_token', '')

            if not cid or not csecret:
                messagebox.showwarning("Inkomplett", "Ange både Client ID och Client Secret för Withings.", parent=self.root)
                return

            self.withings_client_id = cid
            self.withings_client_secret = csecret

            if dialog.result.get('auto_exchanged'):
                self.withings_access_token = dialog.result.get('access_token', '')
                self.withings_refresh_token = dialog.result.get('refresh_token', '')
                self.save_config()
                messagebox.showinfo("Withings Ansluten", "✅ Framgångsrikt ansluten till Withings! Hämtar mätvärden...", parent=self.root)
                self.perform_withings_checkin()
                return

            clean_code = _clean_oauth_code(raw_code)
            if clean_code:
                # If clean_code is already a refresh_token string
                if len(clean_code) > 30 and ('_' in clean_code or '-' in clean_code or '.' in clean_code or len(clean_code) > 40):
                    self.withings_refresh_token = clean_code
                    self.save_config()
                    self.perform_withings_checkin()
                else:
                    # Exchange authorization code for tokens using matching callback_url
                    try:
                        handler = WithingsDataHandler(client_id=cid, client_secret=csecret, db=self.db)
                        res = handler.exchange_code_for_token(clean_code, cid, csecret, redirect_uri=cb_url)
                        self.withings_access_token = res.get("access_token", "")
                        self.withings_refresh_token = res.get("refresh_token", "")
                        self.save_config()
                        messagebox.showinfo("Withings Ansluten", "✅ Framgångsrikt ansluten till Withings!", parent=self.root)
                        self.perform_withings_checkin()
                    except Exception as err:
                        messagebox.showerror("Withings Fel", f"Kunde inte byta koden mot token:\n\n{err}", parent=self.root)
            else:
                self.save_config()

    def perform_withings_checkin(self):
        """Perform a dedicated Withings Check-in for weight & body composition."""
        if not (self.withings_client_id and self.withings_client_secret and self.withings_refresh_token):
            messagebox.showwarning(
                "Withings Inte Konfigurerat",
                "Du måste ange Withings Client ID, Client Secret och Refresh Token först.",
                parent=self.root
            )
            self.connect_to_withings()
            return

        self.update_status("📥 Kör Withings Check-in...", False)
        if hasattr(self, 'charts_view') and self.charts_view:
            self.charts_view.set_sync_status("⏳ Synkroniserar Withings-data...")

        def _thread_target():
            res = self.sync_withings(days=365)
            def _update():
                cnt = res.get("count", 0) if isinstance(res, dict) else 0
                if cnt > 0:
                    self.update_status(f"✅ Withings Check-in genomförd ({cnt} mätvärden sparade)!", False)
                    if hasattr(self, 'charts_view') and self.charts_view:
                        self.charts_view.set_sync_status(f"✅ Withings synkning klar ({cnt} mätvärden)!", is_done=True)
                        self.charts_view.refresh_all_views()
                    messagebox.showinfo(
                        "Withings Check-in Klar",
                        f"✅ Hämta Withings-data lyckades!\n\n{cnt} vikt- och kroppssammansättningsmätningar sparades i databasen och graferna har uppdaterats.",
                        parent=self.root
                    )
                else:
                    err_msg = res.get("message") or res.get("error") if isinstance(res, dict) else "Inga mätvärden hittades."
                    self.update_status("⚠️ Withings: Inga mätvärden hittades.", False)
                    if hasattr(self, 'charts_view') and self.charts_view:
                        self.charts_view.set_sync_status("⚠️ Withings: Inga mätvärden hittades.", is_done=True)
                        self.charts_view.refresh_all_views()
                    messagebox.showwarning(
                        "Withings Check-in",
                        f"⚠️ Inga mätvärden hämtades från Withings API.\n\nDetaljer: {err_msg}\n\nTips: Se till att ditt Withings-konto har sparade mätningar under senaste året, eller importera din export-fil via 'Withings -> Importera Export-fil'.",
                        parent=self.root
                    )
            self.root.after(0, _update)

        threading.Thread(target=_thread_target, daemon=True).start()

    def sync_withings_full(self):
        """Sync full Withings history (30 days)."""
        self.perform_withings_checkin()

    def perform_garmin_checkin(self):
        """Perform a Garmin Check-in: fetch latest metrics & save into local SQLite DB, showing live progress & auto-updating charts."""
        if not self.authenticated or not self.garmin_handler:
            messagebox.showwarning("Inte Ansluten", "Du måste ansluta till Garmin Connect först.", parent=self.root)
            self.connect_to_garmin()
            return
            
        self.checkin_btn.config(state=tk.DISABLED, text="⏳ Synkar...")
        self.update_status("📥 Kör Garmin Check-in...", False)
        if hasattr(self, 'charts_view') and self.charts_view:
            self.charts_view.set_sync_status("⏳ Synkroniserar Garmin-data...")
        
        def _on_progress(curr, total, msg):
            def _update():
                status_str = f"⏳ Synkroniserar Garmin-data ({curr}/{total} d)..."
                self.update_status(status_str, False)
                if hasattr(self, 'charts_view') and self.charts_view:
                    self.charts_view.set_sync_status(f"⏳ Synkar Garmin ({curr}/{total} d)...")
            self.root.after(0, _update)

        def _on_sync_done():
            def _update_ui():
                self.checkin_btn.config(state=tk.NORMAL, text="📥 Check-in")
                self.update_status("✅ Check-in genomförd! Graferna har uppdaterats.", False)
                if hasattr(self, 'charts_view') and self.charts_view:
                    self.charts_view.set_sync_status("✅ Synkning klar!", is_done=True)
                    self.charts_view.refresh_charts()
                messagebox.showinfo("Garmin Check-in", "✅ Check-in genomförd!\n\nSenaste hälsodata och träningspass har hämtats från Garmin och graferna har uppdaterats automatiskt.", parent=self.root)
            self.root.after(0, _update_ui)
            
        try:
            sync_days = max(30, self.charts_view.days_range) if (hasattr(self, 'charts_view') and self.charts_view) else 30
            self.garmin_handler.sync_garmin_history(days=sync_days, on_progress=_on_progress, on_complete=_on_sync_done)
        except Exception as e:
            self.checkin_btn.config(state=tk.NORMAL, text="📥 Check-in")
            self.update_status(f"❌ Check-in fel: {e}", True)

    def connect_to_fitbit(self):
        """Open dedicated popup dialog to configure Fitbit API credentials."""
        current_cid = self.fitbit_handler.client_id if (hasattr(self, 'fitbit_handler') and self.fitbit_handler) else ""
        current_secret = self.fitbit_handler.client_secret if (hasattr(self, 'fitbit_handler') and self.fitbit_handler) else ""

        dialog = FitbitConnectDialog(self.root, client_id=current_cid, client_secret=current_secret, colors=self.colors)
        self.root.wait_window(dialog)

        if dialog.result:
            cid = dialog.result.get('client_id', '')
            csecret = dialog.result.get('client_secret', '')
            raw_code = dialog.result.get('code', '')

            if not cid or not csecret:
                messagebox.showwarning("Inkomplett", "Ange både Client ID och Client Secret för Fitbit.", parent=self.root)
                return

            clean_code = _clean_oauth_code(raw_code)
            if clean_code:
                try:
                    self.fitbit_handler.exchange_code_for_token(clean_code, cid, csecret)
                    messagebox.showinfo("Fitbit Ansluten", "✅ Framgångsrikt ansluten till Fitbit!", parent=self.root)
                    self.perform_fitbit_checkin()
                except Exception as err:
                    messagebox.showerror("Fitbit Fel", f"Kunde inte ansluta till Fitbit:\n\n{err}", parent=self.root)

    def perform_fitbit_checkin(self):
        """Fetch latest Fitbit data and insert into local SQLite DB."""
        if not hasattr(self, 'fitbit_handler') or not self.fitbit_handler:
            messagebox.showwarning("Inte Ansluten", "Fitbit-modulen är inte initierad.", parent=self.root)
            return

        self.update_status("📥 Kör Fitbit Check-in...", False)
        if hasattr(self, 'charts_view') and self.charts_view:
            self.charts_view.set_sync_status("⏳ Synkroniserar Fitbit-data...")

        def _on_prog(curr, total, msg):
            def _update():
                self.update_status(f"⏳ Synkar Fitbit ({curr}/{total} d)...", False)
                if hasattr(self, 'charts_view') and self.charts_view:
                    self.charts_view.set_sync_status(f"⏳ Synkar Fitbit ({curr}/{total} d)...")
            self.root.after(0, _update)

        def _on_done():
            def _update_ui():
                self.update_status("✅ Fitbit Check-in klar! Graferna har uppdaterats.", False)
                if hasattr(self, 'charts_view') and self.charts_view:
                    self.charts_view.set_sync_status("✅ Fitbit synk klar!", is_done=True)
                    self.charts_view.refresh_charts()
                messagebox.showinfo("Fitbit Check-in", "✅ Fitbit Check-in genomförd!\n\nSenaste Fitbit-data har hämtats och sparats i din lokala databas.", parent=self.root)
            self.root.after(0, _update_ui)

        self.fitbit_handler.sync_fitbit_history(days=7, on_progress=_on_prog, on_complete=_on_done)

    def import_fitbit_file(self):
        """Import Fitbit export archive file (CSV/JSON)."""
        from tkinter import filedialog
        filename = filedialog.askopenfilename(
            title="Välj Fitbit Export-fil",
            filetypes=[("Fitbit Files", "*.json;*.csv"), ("JSON Files", "*.json"), ("CSV Files", "*.csv"), ("All Files", "*.*")],
            parent=self.root
        )
        if not filename:
            return

        try:
            results = self.fitbit_handler.import_fitbit_export_file(filename)
            sleep_cnt = results.get("sleep", 0)
            act_cnt = results.get("activities", 0)
            
            if hasattr(self, 'charts_view') and self.charts_view:
                self.charts_view.refresh_charts()

            messagebox.showinfo(
                "Fitbit Import Klar",
                f"✅ Fitbit-data har importerats!\n\n"
                f"• Importerade sömndagar: {sleep_cnt}\n"
                f"• Importerade träningspass: {act_cnt}\n\n"
                f"Graferna har uppdaterats med den nya datan.",
                parent=self.root
            )
        except Exception as err:
            messagebox.showerror("Importfel", f"Kunde inte importera Fitbit-filen:\n\n{err}", parent=self.root)

    def connect_to_strava(self):
        """Open dedicated popup dialog to configure Strava API credentials."""
        current_cid = self.strava_handler.client_id if (hasattr(self, 'strava_handler') and self.strava_handler) else ""
        current_secret = self.strava_handler.client_secret if (hasattr(self, 'strava_handler') and self.strava_handler) else ""

        dialog = StravaConnectDialog(self.root, client_id=current_cid, client_secret=current_secret, colors=self.colors)
        self.root.wait_window(dialog)

        if dialog.result:
            cid = dialog.result.get('client_id', '')
            csecret = dialog.result.get('client_secret', '')
            raw_code = dialog.result.get('code', '')

            if not cid or not csecret:
                messagebox.showwarning("Inkomplett", "Ange både Client ID och Client Secret för Strava.", parent=self.root)
                return

            clean_code = _clean_oauth_code(raw_code)
            if clean_code:
                try:
                    self.strava_handler.exchange_code_for_token(clean_code, cid, csecret)
                    messagebox.showinfo("Strava Ansluten", "✅ Framgångsrikt ansluten till Strava!", parent=self.root)
                    self.perform_strava_checkin()
                except Exception as err:
                    messagebox.showerror("Strava Fel", f"Kunde inte ansluta till Strava:\n\n{err}", parent=self.root)

    def perform_strava_checkin(self):
        """Fetch latest Strava data and insert into local SQLite DB."""
        if not hasattr(self, 'strava_handler') or not self.strava_handler:
            messagebox.showwarning("Inte Ansluten", "Strava-modulen är inte initierad.", parent=self.root)
            return

        self.update_status("📥 Kör Strava Check-in...", False)
        if hasattr(self, 'charts_view') and self.charts_view:
            self.charts_view.set_sync_status("⏳ Synkroniserar Strava-data...")

        def _on_prog(curr, total, msg):
            def _update():
                self.update_status(f"⏳ Synkar Strava ({curr}/{total} pass)...", False)
                if hasattr(self, 'charts_view') and self.charts_view:
                    self.charts_view.set_sync_status(f"⏳ Synkar Strava ({curr}/{total} pass)...")
            self.root.after(0, _update)

        def _on_done():
            def _update_ui():
                self.update_status("✅ Strava Check-in klar! Graferna har uppdaterats.", False)
                if hasattr(self, 'charts_view') and self.charts_view:
                    self.charts_view.set_sync_status("✅ Strava synk klar!", is_done=True)
                    self.charts_view.refresh_all_views()
                messagebox.showinfo("Strava Check-in", "✅ Strava Check-in genomförd!\n\nSenaste Strava-pass har hämtats och sparats i din lokala databas.", parent=self.root)
            self.root.after(0, _update_ui)

        self.strava_handler.sync_strava_history(days=30, on_progress=_on_prog, on_complete=_on_done)

    def import_strava_file(self):
        """Import Strava export archive file (CSV/ZIP/JSON/GPX/FIT)."""
        from tkinter import filedialog
        filename = filedialog.askopenfilename(
            title="Välj Strava Export-fil",
            filetypes=[("Strava Export Files", "*.csv;*.zip;*.json;*.gpx;*.fit"), ("ZIP Archives", "*.zip"), ("CSV Files", "*.csv"), ("JSON Files", "*.json"), ("All Files", "*.*")],
            parent=self.root
        )
        if not filename:
            return

        try:
            results = self.strava_handler.import_strava_export_file(filename)
            act_cnt = results.get("activities", 0)
            
            if hasattr(self, 'charts_view') and self.charts_view:
                self.charts_view.refresh_all_views()

            messagebox.showinfo(
                "Strava Import Klar",
                f"✅ Strava-data har importerats!\n\n"
                f"• Importerade träningspass: {act_cnt}\n\n"
                f"Graferna och träningsloggen har uppdaterats med den nya datan.",
                parent=self.root
            )
        except Exception as err:
            messagebox.showerror("Importfel", f"Kunde inte importera Strava-filen:\n\n{err}", parent=self.root)

    def import_withings_file(self):
        """Import Withings export file (CSV/JSON)."""
        from tkinter import filedialog
        filename = filedialog.askopenfilename(
            title="Välj Withings Export-fil",
            filetypes=[("Withings Files", "*.json;*.csv"), ("JSON Files", "*.json"), ("CSV Files", "*.csv"), ("All Files", "*.*")],
            parent=self.root
        )
        if not filename:
            return

        try:
            handler = WithingsDataHandler(db=self.db)
            results = handler.import_withings_export_file(filename)
            count = results.get("count", 0)

            if hasattr(self, 'charts_view') and self.charts_view:
                self.charts_view.refresh_all_views()

            messagebox.showinfo(
                "Withings Import Klar",
                f"✅ Withings-data har importerats!\n\n"
                f"• Importerade vägningar & kroppsdata: {count}\n\n"
                f"Graferna har uppdaterats med den nya datan.",
                parent=self.root
            )
        except Exception as err:
            messagebox.showerror("Importfel", f"Kunde inte importera Withings-filen:\n\n{err}", parent=self.root)
            
    def create_widgets(self):
        """Create all UI widgets"""
        
        # Create Native Windows Drop-Down Menu Bar
        self.create_menu_bar()
        
        # Configure root grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        
        # Create main container frame using pack layout for clean 100% full-width resizing
        self.main_container = ttk.Frame(self.root, style='Main.TFrame')
        self.main_container.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.main_container.columnconfigure(0, weight=1)
        self.main_container.rowconfigure(0, weight=1)
        
        # --- LEFT PANE: Interactive Graphs & Dashboard ---
        self.left_pane = ttk.Frame(self.main_container, style='Main.TFrame')
        self.left_pane.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        self.charts_view = HealthChartsView(
            self.left_pane,
            db=self.db,
            colors=self.colors,
            on_toggle_chat=self.toggle_chat_pane,
            on_checkin=self.perform_unified_checkin
        )
        self.charts_view.pack(fill=tk.BOTH, expand=True)
        
        # --- RIGHT PANE: Chat & Control Interface ---
        main_frame = ttk.Frame(self.main_container, padding="15", style='Main.TFrame')
        self.right_pane = main_frame
        
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(3, weight=1)  # Chat display gets the extra space (now at row 3)
        
        # By default, start with ONLY the left dashboard pane visible (100% full width)
        self.chat_pane_visible = False
        
        # Row 0: Header Card
        header_card = ttk.Frame(main_frame, style='Card.TFrame', padding="20")
        header_card.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 15))
        header_card.columnconfigure(1, weight=1)
        
        # Try to load and display app icon
        self.logo_photo = None
        try:
            if getattr(sys, 'frozen', False):
                base_path = Path(sys._MEIPASS)
            else:
                base_path = Path(__file__).parent
            
            logo_path = base_path / "logo.png"
            if logo_path.exists():
                from PIL import Image, ImageTk
                img = Image.open(logo_path)
                img = img.resize((48, 48), Image.Resampling.LANCZOS)
                self.logo_photo = ImageTk.PhotoImage(img)
                logo_label = tk.Label(header_card, image=self.logo_photo, bg=self.colors['card_bg'])
                logo_label.grid(row=0, column=0, rowspan=2, padx=(0, 15))
        except Exception as e:
            logger.warning(f"Could not load logo image: {e}")
        
        # Title and subtitle container
        text_container = ttk.Frame(header_card, style='Card.TFrame')
        text_container.grid(row=0, column=1, rowspan=2, sticky=(tk.W, tk.N))
        
        title_label = ttk.Label(text_container, 
                               text="HealthChat",
                               style='Title.TLabel')
        title_label.grid(row=0, column=0, sticky=tk.W)
        
        subtitle_label = ttk.Label(text_container,
                                  text="AI-powered insights for your fitness data",
                                  foreground=self.colors['text_secondary'],
                                  background=self.colors['card_bg'],
                                  font=('Segoe UI', 10))
        subtitle_label.grid(row=1, column=0, sticky=tk.W, pady=(5, 0))
        
        # Right side buttons - larger icons
        button_container = ttk.Frame(header_card, style='Card.TFrame')
        button_container.grid(row=0, column=2, rowspan=2, padx=(10, 0))
        
        search_btn = ttk.Button(button_container,
                               text="🔍",
                               command=self.open_search,
                               style='Modern.TButton',
                               width=3)
        search_btn.grid(row=0, column=0, padx=3)
        self.create_tooltip(search_btn, "Search chat history")
        
        theme_btn = ttk.Button(button_container,
                              text="◐",
                              command=self.toggle_theme,
                              style='Modern.TButton',
                              width=3)
        theme_btn.grid(row=0, column=1, padx=3)
        self.create_tooltip(theme_btn, "Toggle dark mode")
        
        settings_btn = ttk.Button(button_container,
                                 text="Settings",
                                 command=self.open_settings,
                                 style='Modern.TButton')
        settings_btn.grid(row=0, column=2, padx=3)
        self.create_tooltip(settings_btn, "Settings")
        
        # Version label below Settings button
        version_label = ttk.Label(button_container,
                                 text=f"v{APP_VERSION}",
                                 foreground=self.colors['text_secondary'],
                                 background=self.colors['card_bg'],
                                 font=('Segoe UI', 8))
        version_label.grid(row=1, column=2, sticky=tk.E, padx=3)
        
        # Row 1: Control buttons card
        control_card = ttk.Frame(main_frame, style='Card.TFrame', padding="15")
        control_card.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 15))
        
        self.connect_btn = ttk.Button(control_card,
                                     text="▶ Connect to Garmin",
                                     command=self.connect_to_garmin,
                                     style='Accent.TButton')
        self.connect_btn.grid(row=0, column=0, padx=(0, 8))
        
        self.checkin_btn = ttk.Button(control_card,
                                     text="📥 Check-in",
                                     command=self.perform_unified_checkin,
                                     style='Accent.TButton',
                                     state=tk.NORMAL)
        self.checkin_btn.grid(row=0, column=1, padx=4)
        self.create_tooltip(self.checkin_btn, "Hämta och spara senaste data från alla anslutna källor (Garmin, Fitbit m.fl.)")
        
        self.refresh_btn = ttk.Button(control_card,
                                     text="Refresh",
                                     command=self.refresh_data,
                                     style='Modern.TButton',
                                     state=tk.DISABLED)
        self.refresh_btn.grid(row=0, column=2, padx=4)
        self.create_tooltip(self.refresh_btn, "Refresh Garmin data")
        
        self.reset_btn = ttk.Button(control_card,
                                   text="↺ Reset",
                                   command=self.reset_chat,
                                   style='Modern.TButton',
                                   state=tk.DISABLED)
        self.reset_btn.grid(row=0, column=2, padx=4)
        self.create_tooltip(self.reset_btn, "Clear chat history")
        
        self.save_prompts_btn = ttk.Button(control_card,
                                          text="📝 Prompts",
                                          command=self.open_saved_prompts,
                                          style='Modern.TButton')
        self.save_prompts_btn.grid(row=0, column=3, padx=4)
        self.create_tooltip(self.save_prompts_btn, "Manage saved prompts")
        
        self.save_chat_btn = ttk.Button(control_card,
                                       text="💾 Save Chat",
                                       command=self.save_chat_history,
                                       style='Modern.TButton',
                                       state=tk.DISABLED)
        self.save_chat_btn.grid(row=0, column=4, padx=4)
        self.create_tooltip(self.save_chat_btn, "Save this conversation with a custom name")
        
        self.view_chats_btn = ttk.Button(control_card,
                                        text="📂 History",
                                        command=self.open_chat_history_viewer,
                                        style='Modern.TButton')
        self.view_chats_btn.grid(row=0, column=5, padx=(4, 0))
        self.create_tooltip(self.view_chats_btn, "View saved chats")
        
        # Favorite button removed - feature was non-functional
        
        self.export_btn = ttk.Button(control_card,
                                     text="📄 Export Report",
                                     command=self.export_conversation_report,
                                     style='Modern.TButton',
                                     state=tk.DISABLED)
        self.export_btn.grid(row=0, column=6, padx=(4, 0))
        self.create_tooltip(self.export_btn, "Export conversation as PDF, DOCX, or TXT")
        
        # Status label
        self.status_label = ttk.Label(control_card,
                                     text="⚪  Not connected",
                                     style='Status.TLabel')
        self.status_label.grid(row=1, column=0, columnspan=7, sticky=tk.W, pady=(10, 0))
        
        # Smart Suggestions and Follow-up Questions removed from main grid
        # They were causing the chat display to shrink when shown
        # These features can be re-enabled later with a better UI approach (popup/sidebar)
        
        # Row 2: MFA card (initially hidden) - moved up from row 4
        # Row 2: MFA card (initially hidden) - moved up from row 4
        self.mfa_frame = ttk.LabelFrame(main_frame, 
                                       text="🔐 Multi-Factor Authentication", 
                                       style='Card.TLabelframe',
                                       padding="20")
        self.mfa_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=(0, 15))
        self.mfa_frame.grid_remove()  # Hide initially
        self.mfa_frame.columnconfigure(2, weight=1)
        
        ttk.Label(self.mfa_frame, 
                 text="Enter 6-digit code:", 
                 background=self.colors['card_bg'],
                 foreground=self.colors['text'],
                 font=('Segoe UI', 10)).grid(row=0, column=0, padx=(0, 15), pady=5)
        
        self.mfa_entry = ttk.Entry(self.mfa_frame, 
                                  width=15, 
                                  font=('Segoe UI', 12),
                                  style='Modern.TEntry')
        self.mfa_entry.grid(row=0, column=1, padx=(0, 15), pady=5)
        self.mfa_entry.bind('<Return>', lambda e: self.submit_mfa())
        
        self.mfa_btn = ttk.Button(self.mfa_frame,
                                 text="Submit Code",
                                 command=self.submit_mfa,
                                 style='Accent.TButton')
        self.mfa_btn.grid(row=0, column=2, sticky=tk.W, pady=5)
        
        # Row 3: Chat display card (gets extra space) - moved up from row 5
        chat_card = ttk.Frame(main_frame, style='Card.TFrame', padding="0")
        chat_card.grid(row=3, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 15))
        chat_card.columnconfigure(0, weight=1)
        chat_card.rowconfigure(0, weight=1)
        
        # Chat history (scrolled text) with modern styling
        self.chat_display = scrolledtext.ScrolledText(chat_card,
                                                      wrap=tk.WORD,
                                                      font=('Segoe UI', 11),  # Increased from 10 for better readability
                                                      bg=self.colors['card_bg'],
                                                      fg=self.colors['text'],  # Set text color for dark mode
                                                      relief=tk.FLAT,
                                                      borderwidth=0,
                                                      padx=20,
                                                      pady=15)
        self.chat_display.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.chat_display.config(state=tk.DISABLED)
        
        # Configure text tags with modern colors
        self.chat_display.tag_configure('user', 
                                       foreground=self.colors['accent'], 
                                       font=('Segoe UI', 10, 'bold'))
        self.chat_display.tag_configure('assistant', 
                                       foreground=self.colors['success'], 
                                       font=('Segoe UI', 10, 'bold'))
        self.chat_display.tag_configure('system', 
                                       foreground=self.colors['text_secondary'], 
                                       font=('Segoe UI', 9, 'italic'))
        self.chat_display.tag_configure('timestamp', 
                                       foreground=self.colors['text_secondary'], 
                                       font=('Segoe UI', 8))
        self.chat_display.tag_configure('bold', 
                                       font=('Segoe UI', 10, 'bold'))
        self.chat_display.tag_configure('header', 
                                       font=('Segoe UI', 11, 'bold'), 
                                       foreground=self.colors['text'])
        self.chat_display.tag_configure('table',
                                       font=('Courier New', 10),
                                       foreground=self.colors['text'], 
                                       spacing1=2, 
                                       spacing3=2)
        
        # Row 4: Input card (moved up from row 6)
        input_card = ttk.Frame(main_frame, style='Card.TFrame', padding="15")
        input_card.grid(row=4, column=0, sticky=(tk.W, tk.E))
        input_card.columnconfigure(0, weight=1)
        
        # Message input (multi-line Text widget) with modern styling
        self.message_entry = tk.Text(input_card, 
                                     height=3,
                                     font=('Segoe UI', 10),
                                     wrap=tk.WORD,
                                     relief=tk.FLAT,
                                     borderwidth=1,
                                     highlightthickness=1,
                                     highlightbackground=self.colors['border'],
                                     highlightcolor=self.colors['accent'],
                                     bg=self.colors['card_bg'],
                                     fg=self.colors['text'],
                                     insertbackground=self.colors['accent'])
        self.message_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 10))
        
        # Send button with modern accent style
        self.send_btn = ttk.Button(input_card,
                                   text="Send →",
                                   command=self.send_message,
                                   state=tk.DISABLED,
                                   style='Accent.TButton')
        self.send_btn.grid(row=0, column=1)
        
        # Helper text
        helper_text = ttk.Label(input_card,
                               text="Ctrl+Enter to send  •  Enter for new line",
                               foreground=self.colors['text_secondary'],
                               background=self.colors['card_bg'],
                               font=('Segoe UI', 8))
        helper_text.grid(row=1, column=0, columnspan=2, sticky=tk.W, pady=(8, 0))
        
        # Bind Ctrl+Enter to send (Enter alone creates new line)
        self.message_entry.bind('<Control-Return>', lambda e: self.send_message())
        self.message_entry.bind('<Control-Key-Return>', lambda e: self.send_message())
        self.message_entry.config(state=tk.DISABLED)
        
        # Row 5: Quick Questions card with customization
        quick_q_header = ttk.Frame(main_frame, style='Card.TFrame', padding="15 15 15 5")
        quick_q_header.grid(row=5, column=0, sticky=(tk.W, tk.E), pady=(15, 0))
        quick_q_header.columnconfigure(0, weight=1)
        
        ttk.Label(quick_q_header,
                 text="Quick Questions",
                 style='Heading.TLabel').grid(row=0, column=0, sticky=tk.W)
        
        ttk.Button(quick_q_header,
                  text="⚙️ Customize",
                  command=self.customize_quick_questions,
                  style='Modern.TButton',
                  width=12).grid(row=0, column=1, sticky=tk.E)
        
        # Quick questions buttons container
        self.examples_card = ttk.Frame(main_frame, style='Card.TFrame', padding="15 5 15 15")
        self.examples_card.grid(row=6, column=0, sticky=(tk.W, tk.E))
        self.examples_card.columnconfigure(0, weight=1)
        self.examples_card.columnconfigure(1, weight=1)
        
        # Load and display quick questions
        self.load_quick_questions()
        
    def toggle_chat_pane(self):
        """Toggle visibility of the right AI Chat side panel."""
        if self.chat_pane_visible:
            self.right_pane.pack_forget()
            self.chat_pane_visible = False
            if hasattr(self, 'charts_view') and self.charts_view:
                self.charts_view.update_chat_button(is_open=False)
                self.charts_view.refresh_all_views()
        else:
            self.right_pane.pack(side=tk.RIGHT, fill=tk.BOTH, expand=False)
            self.chat_pane_visible = True
            if hasattr(self, 'charts_view') and self.charts_view:
                self.charts_view.update_chat_button(is_open=True)
                self.charts_view.refresh_all_views()

    def add_message(self, sender, message, tag='user'):
        """Add a message to the chat display"""
        self.chat_display.config(state=tk.NORMAL)
        
        # Ensure message is properly decoded as UTF-8
        if isinstance(message, bytes):
            message = message.decode('utf-8', errors='replace')
        
        # Clean up any corrupted bullet points in the message
        message = message.replace('â€¢', '•')  # Fix corrupted bullets
        message = message.replace('â€"', '—')  # Fix em dash
        message = message.replace('â€™', "'")  # Fix apostrophe
        message = message.replace('â€œ', '"')  # Fix opening quote
        message = message.replace('â€', '"')   # Fix closing quote
        
        # Add timestamp
        timestamp = datetime.now().strftime("%H:%M")
        self.chat_display.insert(tk.END, f"[{timestamp}] ", 'timestamp')
        
        # Add sender
        self.chat_display.insert(tk.END, f"{sender}: ", tag)
        
        # Parse and add message with markdown formatting
        if tag == 'assistant':
            self._insert_markdown(message)
        else:
            self.chat_display.insert(tk.END, f"{message}\n\n")
        
        # Save to chat history (but not system messages)
        if tag != 'system':
            self.current_chat_history.append({
                'timestamp': datetime.now().isoformat(),
                'sender': sender,
                'message': message,
                'type': tag
            })
        
        self.chat_display.config(state=tk.DISABLED)
        self.chat_display.see(tk.END)
        
    def _insert_markdown(self, text):
        """Insert text with basic markdown formatting (headers, bold, bullets, tables)"""
        import re
        
        lines = text.split('\n')

        i = 0
        while i < len(lines):
            line = lines[i]

            # Detect table (lines with | characters)
            if '|' in line and line.strip().startswith('|'):
                # Collect all consecutive table lines
                table_lines = []
                while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                # Render the collected table with proper column alignment
                self._render_table(table_lines)
                self.chat_display.insert(tk.END, '\n')
                continue
            
            # Handle headers (#### or ### or ## or #)
            if line.startswith('#### '):
                header_text = line[5:]
                self._insert_inline_formatting(header_text)
                self.chat_display.insert(tk.END, '\n', 'header')
            elif line.startswith('### '):
                header_text = line[4:]
                self._insert_inline_formatting(header_text)
                self.chat_display.insert(tk.END, '\n', 'header')
            elif line.startswith('## '):
                header_text = line[3:]
                self._insert_inline_formatting(header_text)
                self.chat_display.insert(tk.END, '\n', 'header')
            elif line.startswith('# '):
                header_text = line[2:]
                self._insert_inline_formatting(header_text)
                self.chat_display.insert(tk.END, '\n', 'header')
            # Handle bullets (- item or * item)
            elif line.strip().startswith(('- ', '* ')):
                bullet_text = '  • ' + line.strip()[2:]
                self._insert_inline_formatting(bullet_text)
                self.chat_display.insert(tk.END, '\n')
            # Handle numbered lists (1. item)
            elif re.match(r'^\d+\.\s', line.strip()):
                self._insert_inline_formatting(line)
                self.chat_display.insert(tk.END, '\n')
            # Regular line with possible inline formatting
            else:
                self._insert_inline_formatting(line)
                self.chat_display.insert(tk.END, '\n')
            
            i += 1
        
        self.chat_display.insert(tk.END, "\n")
    
    def _render_table(self, lines):
        """Render markdown table lines with proper column alignment"""
        import re

        rows = []
        for line in lines:
            stripped = line.strip()
            # Skip separator lines (only |, -, :, +, space)
            if re.match(r'^[\|\-\+\:\s]+$', stripped):
                continue
            # Parse cells and strip bold markers for display
            cells = [re.sub(r'\*\*(.+?)\*\*', r'\1', c.strip())
                     for c in stripped.strip('|').split('|')]
            rows.append(cells)

        if not rows:
            return

        # Normalize column count
        num_cols = max(len(row) for row in rows)
        rows = [row + [''] * (num_cols - len(row)) for row in rows]

        # Calculate column widths
        col_widths = [max(len(row[j]) for row in rows) for j in range(num_cols)]

        # Render header row
        header_str = ' | '.join(rows[0][j].ljust(col_widths[j]) for j in range(num_cols))
        self.chat_display.insert(tk.END, header_str + '\n', 'table')

        # Separator under header
        sep_str = '-+-'.join('-' * col_widths[j] for j in range(num_cols))
        self.chat_display.insert(tk.END, sep_str + '\n', 'table')

        # Data rows
        for row in rows[1:]:
            row_str = ' | '.join(row[j].ljust(col_widths[j]) for j in range(num_cols))
            self.chat_display.insert(tk.END, row_str + '\n', 'table')

    def _insert_inline_formatting(self, text):
        """Insert text with inline bold formatting (**text**)"""
        import re
        
        # Split by bold markers **text**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                # Bold text
                bold_text = part[2:-2]
                self.chat_display.insert(tk.END, bold_text, 'bold')
            else:
                # Regular text
                self.chat_display.insert(tk.END, part)
        
    def update_status(self, message, is_error=False):
        """Update the status label"""
        self.status_label.config(text=message)
        if is_error:
            self.status_label.config(foreground='#e74c3c')
        else:
            self.status_label.config(foreground='#27ae60')
            
    def open_settings(self):
        """Open settings dialog"""
        current_config = {
            'ai_provider': self.ai_provider,
            'xai_api_key': self.xai_api_key or '',
            'xai_model': self.xai_model,
            'openai_api_key': self.openai_api_key or '',
            'openai_model': self.openai_model,
            'azure_api_key': self.azure_api_key or '',
            'azure_endpoint': self.azure_endpoint or '',
            'azure_deployment': self.azure_deployment or '',
            'gemini_api_key': self.gemini_api_key or '',
            'gemini_model': self.gemini_model,
            'anthropic_api_key': self.anthropic_api_key or '',
            'anthropic_model': self.anthropic_model,
            'ollama_model': self.ollama_model,
            'ollama_base_url': self.ollama_base_url,
            'garmin_email': self.garmin_email or '',
            'garmin_password': self.garmin_password or '',
            'withings_client_id': self.withings_client_id or '',
            'withings_client_secret': self.withings_client_secret or '',
            'withings_refresh_token': self.withings_refresh_token or ''
        }
        
        dialog = SettingsDialog(self.root, current_config, colors=self.colors)
        self.root.wait_window(dialog)
        
        if dialog.result:
            # Update AI provider
            self.ai_provider = dialog.result.get('ai_provider', 'xai')
            
            # Update all provider API keys
            self.xai_api_key = dialog.result.get('xai_api_key', '')
            self.xai_model = dialog.result.get('xai_model', 'grok-3')
            self.openai_api_key = dialog.result.get('openai_api_key', '')
            self.openai_model = dialog.result.get('openai_model', 'gpt-4o')
            self.azure_api_key = dialog.result.get('azure_api_key', '')
            self.azure_endpoint = dialog.result.get('azure_endpoint', '')
            self.azure_deployment = dialog.result.get('azure_deployment', '')
            self.gemini_api_key = dialog.result.get('gemini_api_key', '')
            self.gemini_model = dialog.result.get('gemini_model', 'gemini-1.5-flash')
            self.anthropic_api_key = dialog.result.get('anthropic_api_key', '')
            self.anthropic_model = dialog.result.get('anthropic_model', 'claude-sonnet-4-6')
            self.ollama_model = dialog.result.get('ollama_model', 'llama3.2')
            self.ollama_base_url = dialog.result.get('ollama_base_url', 'http://localhost:11434/v1')
            
            # Update Garmin credentials
            self.garmin_email = dialog.result.get('garmin_email', '')
            self.garmin_password = dialog.result.get('garmin_password', '')

            # Update Withings credentials
            self.withings_client_id = dialog.result.get('withings_client_id', '')
            self.withings_client_secret = dialog.result.get('withings_client_secret', '')
            self.withings_refresh_token = dialog.result.get('withings_refresh_token', '')
            
            self.save_config()
            
            # If already authenticated, reinitialize AI client
            if self.ai_client:
                try:
                    self.initialize_ai_client()
                    provider_names = {
                        'xai': 'xAI (Grok)',
                        'openai': 'OpenAI (ChatGPT)',
                        'azure': 'Azure OpenAI',
                        'gemini': 'Google Gemini',
                        'anthropic': 'Anthropic (Claude)',
                        'ollama': 'Ollama (Local)'
                    }
                    provider_name = provider_names.get(self.ai_provider, self.ai_provider)
                    self.add_message("System", f"Settings updated! Now using: {provider_name}", 'system')
                except Exception as e:
                    self.add_message("System", f"Error updating AI client: {e}", 'system')
                    
    def initialize_ai_client(self):
        """Initialize AI client based on selected provider"""
        try:
            provider = self.ai_provider
            
            if provider == 'xai' and self.xai_api_key:
                from ai_client import AIClient
                self.ai_client = AIClient(provider='xai', api_key=self.xai_api_key, model=self.xai_model)
                logger.info(f"AI client initialized: xAI ({self.xai_model})")
                return True
                
            elif provider == 'openai' and self.openai_api_key:
                from ai_client import AIClient
                self.ai_client = AIClient(provider='openai', api_key=self.openai_api_key, model=self.openai_model)
                logger.info(f"AI client initialized: OpenAI ({self.openai_model})")
                return True
                
            elif provider == 'azure' and self.azure_api_key and self.azure_endpoint:
                from ai_client import AIClient
                self.ai_client = AIClient(
                    provider='azure',
                    api_key=self.azure_api_key,
                    azure_endpoint=self.azure_endpoint,
                    azure_deployment=self.azure_deployment
                )
                logger.info(f"AI client initialized: Azure OpenAI")
                return True
                
            elif provider == 'gemini' and self.gemini_api_key:
                from ai_client import AIClient
                self.ai_client = AIClient(provider='gemini', api_key=self.gemini_api_key, model=self.gemini_model)
                logger.info(f"AI client initialized: Google Gemini ({self.gemini_model})")
                return True
                
            elif provider == 'anthropic' and self.anthropic_api_key:
                from ai_client import AIClient
                self.ai_client = AIClient(provider='anthropic', api_key=self.anthropic_api_key, model=self.anthropic_model)
                logger.info(f"AI client initialized: Anthropic ({self.anthropic_model})")
                return True
                
            elif provider == 'ollama':
                from ai_client import AIClient
                self.ai_client = AIClient(
                    provider='ollama',
                    api_key='ollama',
                    model=self.ollama_model,
                    ollama_base_url=self.ollama_base_url
                )
                logger.info(f"AI client initialized: Ollama ({self.ollama_model}) at {self.ollama_base_url}")
                return True
            else:
                logger.warning(f"No valid API key for provider: {provider}")
                return False
                
        except Exception as e:
            logger.error(f"Error initializing AI client: {e}")
            return False
    
    def get_current_ai_key(self):
        """Get the API key for the currently selected provider"""
        provider = self.ai_provider
        if provider == 'xai':
            return self.xai_api_key
        elif provider == 'openai':
            return self.openai_api_key
        elif provider == 'azure':
            return self.azure_api_key
        elif provider == 'gemini':
            return self.gemini_api_key
        elif provider == 'anthropic':
            return self.anthropic_api_key
        elif provider == 'ollama':
            return 'ollama'  # Always valid - no real key needed
        return None
    
    def connect_to_garmin(self):
        """Initialize and authenticate with Garmin Connect"""
        # Check if all credentials are configured
        current_ai_key = self.get_current_ai_key()
        
        if not current_ai_key or not self.garmin_email or not self.garmin_password:
            provider_names = {
                'xai': 'xAI',
                'openai': 'OpenAI',
                'azure': 'Azure OpenAI',
                'gemini': 'Google Gemini',
                'anthropic': 'Anthropic',
                'ollama': 'Ollama (Local - no key needed)'
            }
            provider_name = provider_names.get(self.ai_provider, self.ai_provider)
            
            if self.ai_provider == 'ollama':
                # Ollama needs no API key, so only Garmin creds are missing
                messagebox.showerror(
                    "Configuration Required",
                    f"Please configure your Garmin credentials in Settings before connecting.\n\n"
                    f"You need:\n"
                    f"- Garmin Connect email\n"
                    f"- Garmin Connect password\n\n"
                    f"(Ollama requires no API key - just make sure Ollama is running)",
                    parent=self.root
                )
            else:
                messagebox.showerror(
                    "Configuration Required",
                    f"Please configure all your credentials in Settings before connecting to Garmin.\n\n"
                    f"You need:\n"
                    f"- {provider_name} API key\n"
                    f"- Garmin Connect email\n"
                    f"- Garmin Connect password",
                    parent=self.root
                )
            self.open_settings()
            return
            
        self.connect_btn.config(state=tk.DISABLED, text="Connecting...")
        self.update_status("Connecting to Garmin...", False)
        
        # Run in thread to prevent UI freezing
        thread = threading.Thread(target=self._authenticate_garmin)
        thread.daemon = True
        thread.start()
        
    def sync_withings(self, days: int = 365) -> Dict[str, Any]:
        """Sync Withings weight and body composition data (runs safely in background)"""
        if self.withings_client_id and self.withings_client_secret and self.withings_refresh_token:
            try:
                handler = WithingsDataHandler(
                    client_id=self.withings_client_id,
                    client_secret=self.withings_client_secret,
                    refresh_token=self.withings_refresh_token,
                    access_token=self.withings_access_token,
                    db=self.db
                )
                res = handler.sync_withings_data(days=days)
                if res.get("access_token"):
                    self.withings_access_token = res["access_token"]
                if res.get("refresh_token"):
                    self.withings_refresh_token = res["refresh_token"]
                    self.save_config()
                if res.get("success"):
                    logger.info(f"Successfully synced {res.get('count')} Withings measurements")
                return res
            except Exception as e:
                logger.warning(f"Error syncing Withings data: {e}")
                return {"success": False, "error": str(e), "count": 0}
        return {"success": False, "error": "Saknar Withings OAuth2 credentials", "count": 0}

    def _authenticate_garmin(self):
        """Authenticate with Garmin (runs in thread)"""
        try:
            # Initialize AI client with current provider
            if not self.initialize_ai_client():
                self.root.after(0, lambda: self._on_auth_failure("Failed to initialize AI client. Please check your settings."))
                return
            
            # Initialize Garmin handler with stored credentials
            self.garmin_handler = GarminDataHandler(self.garmin_email, self.garmin_password)
            result = self.garmin_handler.authenticate()
            
            if result.get('success'):
                self.authenticated = True
                self.mfa_required = False
                self.sync_withings()
                self.root.after(0, lambda: self._on_auth_success())
            elif result.get('mfa_required'):
                self.mfa_required = True
                self.authenticated = False
                self.root.after(0, lambda: self._show_mfa_input())
            else:
                error_msg = result.get('error', 'Unknown Garmin authentication error')
                self.root.after(0, lambda: self._on_auth_failure(error_msg))
                
        except Exception as e:
            self.root.after(0, lambda: self._on_auth_failure(str(e)))
            
    def _show_mfa_input(self):
        """Show MFA input frame"""
        self.mfa_frame.grid()
        self.update_status("🔐 MFA Required: Enter your 6-digit code", False)
        self.mfa_entry.focus()
        
    def submit_mfa(self):
        """Submit MFA code"""
        mfa_code = self.mfa_entry.get().strip()
        
        if not mfa_code or len(mfa_code) != 6:
            self.update_status("❌ Please enter a valid 6-digit MFA code", True)
            return
            
        self.mfa_btn.config(state=tk.DISABLED)
        self.update_status("Submitting MFA code...", False)
        
        # Run in thread
        thread = threading.Thread(target=self._submit_mfa_code, args=(mfa_code,))
        thread.daemon = True
        thread.start()
        
    def _submit_mfa_code(self, mfa_code):
        """Submit MFA code (runs in thread)"""
        try:
            result = self.garmin_handler.submit_mfa(mfa_code)
            
            if result.get('success'):
                self.authenticated = True
                self.mfa_required = False
                self.root.after(0, lambda: self._on_auth_success())
                self.root.after(0, lambda: self.mfa_frame.grid_remove())
            else:
                error_msg = result.get('error', 'Unknown MFA error')
                self.root.after(0, lambda: self._on_auth_failure(error_msg))
                self.root.after(0, lambda: self.mfa_btn.config(state=tk.NORMAL))
                
        except Exception as e:
            self.root.after(0, lambda: self._on_auth_failure(str(e)))
            self.root.after(0, lambda: self.mfa_btn.config(state=tk.NORMAL))
            
    def _on_auth_success(self):
        """Handle successful authentication"""
        self.update_status("✅ Connected to Garmin Connect!", False)
        self.message_entry.config(state=tk.NORMAL)
        self.send_btn.config(state=tk.NORMAL)
        self.refresh_btn.config(state=tk.NORMAL)
        self.reset_btn.config(state=tk.NORMAL)
        self.save_chat_btn.config(state=tk.NORMAL)
        if hasattr(self, 'checkin_btn') and self.checkin_btn:
            self.checkin_btn.config(state=tk.NORMAL)
        # favorite_btn removed
        self.export_btn.config(state=tk.NORMAL)
        self.connect_btn.config(state=tk.NORMAL, text="✅ Connected (Reconnect)")
        self.message_entry.focus()
        
        self.add_message("System",
                        "Connected to Garmin Connect! You can now ask questions about your fitness data.",
                        'system')
        
        # Refresh charts directly from local SQLite database (no automatic background sync on startup)
        if hasattr(self, 'charts_view') and self.charts_view:
            self.charts_view.refresh_charts()
        
    def _on_auth_failure(self, error_msg: str):
        """Handle failed authentication"""
        self.update_status(f"❌ {error_msg}", True)
        self.connect_btn.config(state=tk.NORMAL, text="▶ Connect to Garmin")
        messagebox.showerror("Garmin Connection Error", f"Could not connect to Garmin Connect:\n\n{error_msg}", parent=self.root)

    def send_message(self):
        """Send a message to the chatbot"""
        if not self.authenticated:
            self.update_status("❌ Please connect to Garmin first", True)
            return
            
        # Get message from Text widget
        message = self.message_entry.get("1.0", tk.END).strip()
        if not message:
            return
            
        # Add user message to display
        self.add_message("You", message, 'user')
        
        # Clear input
        self.message_entry.delete("1.0", tk.END)
        
        # Disable input while processing
        self.message_entry.config(state=tk.DISABLED)
        self.send_btn.config(state=tk.DISABLED)
        
        # Process in thread
        thread = threading.Thread(target=self._process_message, args=(message,))
        thread.daemon = True
        thread.start()
        
        # Return 'break' to prevent default behavior when called from key binding
        return 'break'
        
    def _process_message(self, message):
        """Process the message and get AI response (runs in thread)"""
        try:
            # Determine what data to fetch
            query_lower = message.lower()
            
            # Detect if user wants activities by date range or by count
            activity_limit = 5  # Default
            use_date_range = False
            garmin_context = ""
            
            # Check for date range requests (last X days/weeks/months)
            import re
            from datetime import datetime, timedelta
            
            # Match "last/past X days/weeks/months" OR "last/this month/week/year"
            time_period_match = re.search(r'(?:last|past)\s+(\d+)\s+(day|week|month)s?', query_lower)
            simple_period_match = re.search(r'(?:last|past|this)\s+(month|week|year)', query_lower)
            
            if time_period_match:
                number = int(time_period_match.group(1))
                unit = time_period_match.group(2)
                
                # Calculate date range
                end_date = datetime.now()
                if unit == "day":
                    start_date = end_date - timedelta(days=number)
                elif unit == "week":
                    start_date = end_date - timedelta(weeks=number)
                elif unit == "month":
                    start_date = end_date - timedelta(days=number * 30)  # Approximate
                
                logger.info(f"Detected date range query: last {number} {unit}(s)")
                logger.info(f"Date range: {start_date.strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')}")
                use_date_range = True
                
            elif simple_period_match:
                # Handle "last month", "this week", etc.
                period = simple_period_match.group(1)
                prefix = simple_period_match.group(0).split()[0]  # "last", "this", or "past"
                
                end_date = datetime.now()
                
                if period == "month":
                    if prefix == "this":
                        # Current month: from 1st to today
                        start_date = end_date.replace(day=1)
                    else:
                        # Last month: 30 days ago
                        start_date = end_date - timedelta(days=30)
                elif period == "week":
                    if prefix == "this":
                        # Current week: last 7 days
                        start_date = end_date - timedelta(days=end_date.weekday())
                    else:
                        # Last week: 7 days ago
                        start_date = end_date - timedelta(days=7)
                elif period == "year":
                    if prefix == "this":
                        # Current year: from Jan 1 to today
                        start_date = end_date.replace(month=1, day=1)
                    else:
                        # Last year: 365 days ago
                        start_date = end_date - timedelta(days=365)
                
                logger.info(f"Detected simple period query: {prefix} {period}")
                logger.info(f"Date range: {start_date.strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')}")
                use_date_range = True
                
            if use_date_range:
                # Fetch activities by date range
                try:
                    activities = self.garmin_handler.get_activities_by_date(
                        start_date.strftime("%Y-%m-%d"),
                        end_date.strftime("%Y-%m-%d")
                    )
                    
                    if activities:
                        # Format activities for context
                        context_parts = [f"=== Activities from {start_date.strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')} ({len(activities)} activities) ==="]
                        for i, activity in enumerate(activities, 1):
                            act_name = activity.get("activityName", "Unknown")
                            act_type = activity.get("activityType", {}).get("typeKey", "Unknown")
                            distance = activity.get("distance", 0) / 1000 if activity.get("distance") else 0
                            duration = activity.get("duration", 0) / 60 if activity.get("duration") else 0
                            calories = activity.get("calories", "N/A")
                            start_time = activity.get("startTimeLocal", "N/A")
                            
                            context_parts.append(f"{i}. {act_name} ({act_type})")
                            context_parts.append(f"   Date: {start_time}")
                            context_parts.append(f"   Distance: {distance:.2f} km")
                            context_parts.append(f"   Duration: {duration:.1f} minutes")
                            context_parts.append(f"   Calories: {calories}")
                            context_parts.append("")
                        
                        garmin_context = "\n".join(context_parts)
                        logger.info(f"Fetched {len(activities)} activities for date range")
                    else:
                        garmin_context = f"=== Activities from {start_date.strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')} ===\nNo workouts or activities recorded during this period."
                        logger.info("No activities found for requested date range")
                except Exception as e:
                    logger.error(f"Error fetching activities by date: {e}")
                    # Fall back to regular method
                    use_date_range = False
            
            # If not using date range, detect count-based requests
            if not use_date_range:
                # Check for requests for more activities
                if any(phrase in query_lower for phrase in [
                    "show me more", "more activities", "all activities", "all my activities",
                    "show all", "recent activities"
                ]):
                    activity_limit = 30  # Fetch more for these queries
                    logger.info(f"Detected request for more activities, fetching {activity_limit}")
                
                # Check for specific number requests
                number_match = re.search(r'(?:last|past|recent)\s+(\d+)', query_lower)
                if number_match:
                    requested_count = int(number_match.group(1))
                    activity_limit = min(requested_count, 50)  # Cap at 50
                    logger.info(f"User requested {requested_count} activities, fetching {activity_limit}")
                
                # Fetch appropriate data using regular method (supporting English & Swedish queries)
                if any(word in query_lower for word in ["pt", "coach", "personlig tränare", "analysera", "utvärdera", "hälsa", "form", "helhet", "all data", "overview", "summary"]):
                    garmin_context = self.garmin_handler.format_data_for_context("comprehensive", activity_limit=15)
                elif any(word in query_lower for word in ["sömn", "sömnskikringar", "sova", "sleep", "rest", "bed"]):
                    garmin_context = self.garmin_handler.format_data_for_context("sleep")
                elif any(word in query_lower for word in ["body battery", "energi", "återhämtning", "återhämtningskapacitet", "energy"]):
                    garmin_context = self.garmin_handler.format_data_for_context("comprehensive", activity_limit=10)
                elif any(word in query_lower for word in ["träning", "träningspass", "pass", "löpning", "cykling", "activity", "activities", "workout", "run", "walk", "bike", "exercise"]):
                    garmin_context = self.garmin_handler.format_data_for_context("comprehensive", activity_limit=15)
                elif any(word in query_lower for word in ["steg", "stegräknare", "sträcka", "kalorier", "step", "walk", "distance", "calorie"]):
                    garmin_context = self.garmin_handler.format_data_for_context("summary")
                elif any(word in query_lower for word in ["stress", "stressed", "tension"]):
                    garmin_context = self.garmin_handler.format_data_for_context("stress")
                elif any(word in query_lower for word in ["respiration", "andning", "breathing", "breath"]):
                    garmin_context = self.garmin_handler.format_data_for_context("respiration")
                elif any(word in query_lower for word in ["hydration", "vatten", "vätska", "water", "drink", "fluid"]):
                    garmin_context = self.garmin_handler.format_data_for_context("hydration")
                elif any(word in query_lower for word in ["nutrition", "mat", "kost", "food", "eat", "meal", "diet", "protein", "carbs", "fat", "macros", "calories consumed", "food log", "logged"]):
                    garmin_context = self.garmin_handler.format_data_for_context("nutrition")
                elif any(word in query_lower for word in ["trappor", "våningar", "floor", "climb", "stairs", "elevation"]):
                    garmin_context = self.garmin_handler.format_data_for_context("floors")
                elif any(word in query_lower for word in ["intense", "intensity", "vigorous", "moderate"]):
                    garmin_context = self.garmin_handler.format_data_for_context("intensity")
                elif any(word in query_lower for word in ["spo2", "syre", "oxygen", "pulse ox"]):
                    garmin_context = self.garmin_handler.format_data_for_context("spo2")
                elif any(word in query_lower for word in ["hrv", "pulsvariabilitet", "heart rate variability", "variability"]):
                    garmin_context = self.garmin_handler.format_data_for_context("hrv")
                elif any(word in query_lower for word in ["vo2", "kondition", "fitness age", "training status", "training load"]):
                    garmin_context = self.garmin_handler.format_data_for_context("training")
                else:
                    garmin_context = self.garmin_handler.format_data_for_context("comprehensive", activity_limit=activity_limit)
            
            # Add conversation context for memory
            context_summary = ""
            if self.conversation_context:
                recent_convs = self.conversation_context[-5:]
                context_summary = "\n\nPrevious conversation context:\n"
                for conv in recent_convs:
                    sender = conv.get('sender', 'User')
                    msg = conv.get('message', '')[:100]  # First 100 chars
                    context_summary += f"{sender}: {msg}...\n"
            
            enhanced_context = garmin_context + context_summary
            
            # Get AI response
            response = self.ai_client.chat(message, enhanced_context)
            
            # Add response to display
            self.root.after(0, lambda: self.add_message("HealthChat", response, 'assistant'))
            
            # Follow-up suggestions disabled for better UX - they took up too much vertical space
            # self.root.after(0, lambda: self.show_followup_buttons(response))
            
            # Update conversation context
            self.conversation_context.append({
                'sender': 'You',
                'message': message,
                'timestamp': datetime.now().isoformat()
            })
            self.conversation_context.append({
                'sender': 'HealthChat',
                'message': response,
                'timestamp': datetime.now().isoformat()
            })
            
            # Keep context manageable
            if len(self.conversation_context) > self.max_context_messages:
                self.conversation_context = self.conversation_context[-self.max_context_messages:]
            
        except Exception as e:
            error_msg = f"Sorry, I encountered an error: {str(e)}"
            self.root.after(0, lambda: self.add_message("System", error_msg, 'system'))
            
        finally:
            # Re-enable input
            self.root.after(0, lambda: self.message_entry.config(state=tk.NORMAL))
            self.root.after(0, lambda: self.send_btn.config(state=tk.NORMAL))
            self.root.after(0, lambda: self.message_entry.focus())
            
    def use_example(self, question):
        """Use an example question"""
        if not self.authenticated:
            self.update_status("❌ Please connect to Garmin first", True)
            return
            
        self.message_entry.delete("1.0", tk.END)
        self.message_entry.insert("1.0", question)
        self.send_message()
    
    def load_quick_questions(self):
        """Load and display quick questions (default or custom)"""
        # Clear existing buttons
        for widget in self.examples_card.winfo_children():
            widget.destroy()
        
        # Load custom questions or use defaults
        quick_questions_file = self.config_dir / "quick_questions.json"
        
        try:
            if quick_questions_file.exists():
                with open(quick_questions_file, 'r') as f:
                    questions = json.load(f)
            else:
                # Default questions
                questions = [
                    "🏋️‍♂️ Analysera min träning som min Personliga Tränare (PT)",
                    "🏃‍♂️ Hur har mina löppass och tempo utvecklats?",
                    "💤 Hur ser min sömn och återhämtning ut?",
                    "🔥 Hur många kalorier har jag bränt denna vecka?",
                ]
        except Exception as e:
            logger.error(f"Error loading quick questions: {e}")
            questions = [
                "🏋️‍♂️ Analysera min träning som min Personliga Tränare (PT)",
                "🏃‍♂️ Hur har mina löppass och tempo utvecklats?",
                "💤 Hur ser min sömn och återhämtning ut?",
                "🔥 Hur många kalorier har jag bränt denna vecka?",
            ]
        
        # Create buttons with rounded hover effect
        for i, question in enumerate(questions[:8]):  # Max 8 questions (4x2 grid)
            btn = ttk.Button(self.examples_card,
                           text=question,
                           style='QuickQuestion.TButton',
                           command=lambda q=question: self.use_example(q))
            btn.grid(row=i//2, column=i%2, padx=6, pady=6, sticky=(tk.W, tk.E))
    
    def customize_quick_questions(self):
        """Open dialog to customize quick questions"""
        CustomizeQuestionsDialog(self.root, self)
        
    def refresh_data(self):
        """Refresh Garmin data"""
        self.refresh_btn.config(state=tk.DISABLED)
        self.update_status("Refreshing data...", False)
        
        thread = threading.Thread(target=self._refresh_data)
        thread.daemon = True
        thread.start()
        
    def _refresh_data(self):
        """Refresh data (runs in thread)"""
        try:
            result = self.garmin_handler.authenticate()
            if result.get('success'):
                self.root.after(0, lambda: self.update_status("✅ Data refreshed!", False))
                self.root.after(0, lambda: self.add_message("System", "Data refreshed successfully!", 'system'))
            elif result.get('mfa_required'):
                # MFA is required for refresh
                self.mfa_required = True
                self.authenticated = False
                self.root.after(0, lambda: self._show_mfa_input())
                self.root.after(0, lambda: self.update_status("🔐 MFA Required: Enter your 6-digit code", False))
            else:
                error_msg = result.get('error', 'Unknown error')
                self.root.after(0, lambda: self.update_status(f"❌ {error_msg}", True))
        except Exception as e:
            self.root.after(0, lambda: self.update_status(f"❌ Error: {str(e)}", True))
        finally:
            self.root.after(0, lambda: self.refresh_btn.config(state=tk.NORMAL))
            
    def reset_chat(self):
        """Reset the conversation"""
        if self.ai_client:
            self.ai_client.reset_conversation()
            
        self.chat_display.config(state=tk.NORMAL)
        self.chat_display.delete(1.0, tk.END)
        self.chat_display.config(state=tk.DISABLED)
        self.current_chat_history = []
        
        self.add_message("System", "Conversation reset!", 'system')
        self.update_status("✅ Chat reset", False)
    
    def open_saved_prompts(self):
        """Open dialog to manage saved prompts"""
        SavedPromptsDialog(self.root, self)
    
    def load_saved_prompts(self):
        """Load saved prompts from file, using PT Coach examples as defaults"""
        default_prompts = [
            {
                'name': '🏋️‍♂️ PT Coach: Tränings- & Återhämtningsanalys',
                'prompt': 'Du är min professionella Personliga Tränare (COACH AI) och hälsocoach.\n\nAnalysera min tränings- och hälsodata (träningspass, tempo, pulszoner, sömn, stress, HRV och Body Battery):\n1. UTVÄRDERING: Ge en ärlig, strukturerad och konstruktiv analys av min senaste träning och belastning.\n2. ÅTERHÄMTNING & STRESS: Granska min sömn och återhämtningsstatus – riskerar jag överträning? (Om RHR stigit eller HRV/Body Battery fallit, rekommendera aktiv vila).\n3. PT-RÅD & PROGRESSIV BELASTNING: Ge konkreta råd för mina kommande träningspass med tillämpning av progressive overload när återhämtningen är god.\n4. TONSÄTTNING: Svara på svenska i en uppmuntrande, artig och extremt kunnig stil!',
                'created': datetime.now().isoformat()
            },
            {
                'name': '💤 Sömn & Återhämtnings-Check',
                'prompt': 'Analysera min senaste sömn- och stressdata tillsammans med mitt Body Battery på svenska. Behöver jag vila mer eller är jag redo för ett hårt träningspass idag?',
                'created': datetime.now().isoformat()
            },
            {
                'name': '🏃‍♂️ Löpanalys & Tempotrender',
                'prompt': 'Hur har mina löppass utvecklats gällande tempo, snittpuls och distans? Ge mig 3 konkreta insikter för att förbättra min uthållighet på svenska.',
                'created': datetime.now().isoformat()
            }
        ]
        try:
            if self.saved_prompts_file.exists():
                with open(self.saved_prompts_file, 'r') as f:
                    data = json.load(f)
                    if data:
                        return data
            # Save defaults if file does not exist or is empty
            with open(self.saved_prompts_file, 'w') as f:
                json.dump(default_prompts, f, indent=2)
            return default_prompts
        except Exception as e:
            logger.error(f"Error loading saved prompts: {e}")
            return default_prompts
    
    def save_prompt(self, name, prompt):
        """Save a prompt for reuse"""
        try:
            prompts = self.load_saved_prompts()
            prompts.append({'name': name, 'prompt': prompt, 'created': datetime.now().isoformat()})
            with open(self.saved_prompts_file, 'w', encoding='utf-8') as f:
                json.dump(prompts, f, indent=2, ensure_ascii=False)
            logger.info(f"Saved prompt: {name}")
        except Exception as e:
            logger.error(f"Error saving prompt: {e}")
            
    def update_saved_prompt(self, index, name, prompt):
        """Update an existing saved prompt by index"""
        try:
            prompts = self.load_saved_prompts()
            if 0 <= index < len(prompts):
                prompts[index]['name'] = name
                prompts[index]['prompt'] = prompt
                prompts[index]['updated'] = datetime.now().isoformat()
                with open(self.saved_prompts_file, 'w', encoding='utf-8') as f:
                    json.dump(prompts, f, indent=2, ensure_ascii=False)
                logger.info(f"Updated prompt at index {index}: {name}")
        except Exception as e:
            logger.error(f"Error updating prompt: {e}")
    
    def delete_saved_prompt(self, index):
        """Delete a saved prompt"""
        try:
            prompts = self.load_saved_prompts()
            if 0 <= index < len(prompts):
                deleted = prompts.pop(index)
                with open(self.saved_prompts_file, 'w') as f:
                    json.dump(prompts, f, indent=2)
                logger.info(f"Deleted prompt: {deleted['name']}")
        except Exception as e:
            logger.error(f"Error deleting prompt: {e}")
    
    def save_chat_history(self):
        """Save current chat session to file with custom name"""
        if not self.current_chat_history:
            messagebox.showinfo("No Chat History", "There's no chat history to save yet!", parent=self.root)
            return
        
        # Prompt for custom name
        dialog = tk.Toplevel(self.root)
        dialog.title("Save Chat Session")
        dialog.geometry("500x200")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Set window icon
        try:
            if getattr(sys, 'frozen', False):
                base_path = Path(sys._MEIPASS)
            else:
                base_path = Path(__file__).parent
            
            icon_path = base_path / "logo.ico"
            if icon_path.exists():
                dialog.iconbitmap(str(icon_path))
        except Exception as e:
            logger.debug(f"Could not load Save dialog icon: {e}")
        
        # Center on parent
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - (dialog.winfo_width() // 2)
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - (dialog.winfo_height() // 2)
        dialog.geometry(f"+{x}+{y}")
        
        frame = ttk.Frame(dialog, padding="20")
        frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        dialog.columnconfigure(0, weight=1)
        dialog.rowconfigure(0, weight=1)
        frame.columnconfigure(1, weight=1)
        
        ttk.Label(frame, text="💾 Save Chat Session", 
                 font=('Segoe UI', 12, 'bold')).grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=(0, 15))
        
        ttk.Label(frame, text="Session Name:", 
                 font=('Segoe UI', 10)).grid(row=1, column=0, sticky=tk.W, pady=5)
        
        name_var = tk.StringVar()
        name_entry = ttk.Entry(frame, textvariable=name_var, font=('Segoe UI', 10), width=40)
        name_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=5, padx=(10, 0))
        name_entry.focus()
        
        ttk.Label(frame, text="(Optional - leave blank for date/time only)", 
                 font=('Segoe UI', 8), foreground='gray').grid(row=2, column=1, sticky=tk.W, padx=(10, 0))
        
        result = {'saved': False}
        
        def save_with_name():
            custom_name = name_var.get().strip()
            try:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                
                # Create filename with custom name if provided
                if custom_name:
                    # Sanitize the custom name for filename
                    safe_name = "".join(c for c in custom_name if c.isalnum() or c in (' ', '-', '_')).strip()
                    safe_name = safe_name.replace(' ', '_')
                    filename = self.chat_history_dir / f"chat_{timestamp}_{safe_name}.json"
                else:
                    filename = self.chat_history_dir / f"chat_{timestamp}.json"
                
                with open(filename, 'w') as f:
                    json.dump({
                        'saved_at': datetime.now().isoformat(),
                        'custom_name': custom_name if custom_name else None,
                        'messages': self.current_chat_history
                    }, f, indent=2)
                
                result['saved'] = True
                dialog.destroy()
                
                display_name = custom_name if custom_name else "chat session"
                messagebox.showinfo("Chat Saved", 
                                  f"'{display_name}' saved successfully!\n\nLocation: {filename}", 
                                  parent=self.root)
                logger.info(f"Saved chat history to: {filename}")
                
            except Exception as e:
                logger.error(f"Error saving chat history: {e}")
                messagebox.showerror("Save Error", f"Failed to save chat history: {e}", parent=dialog)
        
        def cancel():
            dialog.destroy()
        
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=3, column=0, columnspan=2, pady=(20, 0))
        
        ttk.Button(button_frame, text="Save Chat", command=save_with_name, width=10).grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="Cancel", command=cancel, width=10).grid(row=0, column=1, padx=5)
        
        # Bind Enter key to save
        name_entry.bind('<Return>', lambda e: save_with_name())
        dialog.bind('<Escape>', lambda e: cancel())
    
    
    def open_chat_history_viewer(self):
        """Open dialog to view saved chat histories"""
        ChatHistoryViewer(self.root, self)
    
    def load_conversation_history(self):
        """Load recent conversation context for AI memory"""
        try:
            # Load last 5 chat files for context
            chat_files = sorted(self.chat_history_dir.glob("chat_*.json"),
                              key=lambda f: f.stat().st_mtime,
                              reverse=True)[:5]
            
            for file in chat_files:
                with open(file, 'r') as f:
                    data = json.load(f)
                    messages = data.get('messages', [])
                    # Add to context (last message from each chat)
                    if messages:
                        self.conversation_context.extend(messages[-3:])  # Last 3 from each
            
            # Keep only most recent messages
            self.conversation_context = self.conversation_context[-self.max_context_messages:]
        except Exception as e:
            logger.error(f"Error loading conversation history: {e}")
    
    def toggle_theme(self):
        """Toggle between light and dark mode"""
        # Check current theme
        if not hasattr(self, 'dark_mode'):
            self.dark_mode = False
        
        self.dark_mode = not self.dark_mode
        
        if self.dark_mode:
            # Dark mode colors
            self.colors = {
                'bg': '#202020',
                'card_bg': '#2D2D30',
                'accent': '#60A5FA',
                'accent_hover': '#3B82F6',
                'accent_light': '#1E3A5F',
                'text': '#E5E5E5',
                'text_secondary': '#A0A0A0',
                'border': '#3E3E42',
                'success': '#10B981',
                'warning': '#F59E0B',
                'shadow': '#00000040',
            }
        else:
            # Light mode colors (original)
            self.colors = {
                'bg': '#F3F3F3',
                'card_bg': '#FFFFFF',
                'accent': '#0078D4',
                'accent_hover': '#106EBE',
                'accent_light': '#E6F2FA',
                'text': '#1F1F1F',
                'text_secondary': '#605E5C',
                'border': '#EDEBE9',
                'success': '#107C10',
                'warning': '#D83B01',
                'shadow': '#00000010',
            }
        
        # Apply new theme immediately
        self.apply_theme()
        
        # Save theme preference
        self.save_config()
        logger.info(f"Theme toggled to {'dark' if self.dark_mode else 'light'} mode and saved")
    
    def apply_theme(self):
        """Apply current theme colors to all UI elements"""
        # Update root background
        self.root.configure(bg=self.colors['bg'])
        
        # Re-configure ttk styles with new colors
        style = ttk.Style()
        
        # Base frame style
        style.configure('TFrame',
                       background=self.colors['bg'])
        
        # Card frame style
        style.configure('Card.TFrame',
                       background=self.colors['card_bg'],
                       relief='flat',
                       borderwidth=1)
        
        # Label styles
        style.configure('Title.TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 18, 'bold'))
        
        style.configure('Heading.TLabel',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 11, 'bold'))
        
        style.configure('TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 10))
        
        style.configure('Status.TLabel',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text_secondary'],
                       font=('Segoe UI', 9))
        
        # Button styles
        style.configure('Modern.TButton',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       bordercolor=self.colors['border'],
                       borderwidth=1,
                       relief='flat',
                       padding=(12, 6),
                       font=('Segoe UI', 11))
        style.map('Modern.TButton',
                 background=[('active', self.colors['accent_light'] if self.dark_mode else self.colors['border']), 
                           ('pressed', self.colors['accent_light'] if self.dark_mode else self.colors['border'])],
                 foreground=[('active', self.colors['accent'] if self.dark_mode else self.colors['text']),
                           ('pressed', self.colors['accent'] if self.dark_mode else self.colors['text'])])
        
        style.configure('Accent.TButton',
                       background=self.colors['accent'],
                       foreground='white',
                       borderwidth=0,
                       relief='flat',
                       padding=(16, 8),
                       font=('Segoe UI', 10, 'bold'))
        style.map('Accent.TButton',
                 background=[('active', self.colors['accent_hover']), 
                           ('pressed', self.colors['accent_hover'])],
                 foreground=[('active', 'white'), ('pressed', 'white')])
        
        # Quick Question button style
        style.configure('QuickQuestion.TButton',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       bordercolor=self.colors['border'],
                       borderwidth=1,
                       relief='flat',
                       padding=(12, 8),
                       font=('Segoe UI', 10))
        style.map('QuickQuestion.TButton',
                 background=[('active', self.colors['accent_light']), 
                           ('pressed', self.colors['accent'])],
                 foreground=[('active', self.colors['accent']),
                           ('pressed', 'white')],
                 relief=[('active', 'raised'),
                        ('pressed', 'sunken')])
        
        # LabelFrame style
        style.configure('Card.TLabelframe',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       bordercolor=self.colors['border'],
                       borderwidth=1,
                       relief='flat')
        style.configure('Card.TLabelframe.Label',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 10, 'bold'))
        
        # Update chat display
        self.chat_display.config(bg=self.colors['card_bg'], fg=self.colors['text'])
        
        # Update chat display tags
        self.chat_display.tag_configure('user', foreground=self.colors['accent'])
        self.chat_display.tag_configure('assistant', foreground=self.colors['success'])
        self.chat_display.tag_configure('system', foreground=self.colors['text_secondary'])
        self.chat_display.tag_configure('timestamp', foreground=self.colors['text_secondary'])
        self.chat_display.tag_configure('header', foreground=self.colors['text'])
        self.chat_display.tag_configure('table', foreground=self.colors['text'])
        
        # Update message entry
        self.message_entry.config(
            bg=self.colors['card_bg'],
            fg=self.colors['text'],
            insertbackground=self.colors['accent'],
            highlightbackground=self.colors['border'],
            highlightcolor=self.colors['accent']
        )
        
        # Update MFA entry if it exists
        try:
            self.mfa_entry.config(
                fieldbackground=self.colors['card_bg'],
                foreground=self.colors['text']
            )
        except:
            pass
        
        # Update suggestions label
        try:
            self.suggestions_label.config(
                background=self.colors['card_bg'],
                foreground=self.colors['text_secondary']
            )
        except:
            pass
        
        # Update all frames and labels recursively
        self._update_widget_colors(self.root)
    
    def _update_widget_colors(self, widget):
        """Recursively update widget colors"""
        try:
            widget_type = widget.winfo_class()
            
            # Update Frame backgrounds
            if widget_type in ('TFrame', 'Frame'):
                try:
                    widget.configure(background=self.colors['bg'])
                except:
                    pass
            
            # Update Label colors
            elif widget_type in ('TLabel', 'Label'):
                try:
                    widget.configure(
                        background=self.colors['bg'],
                        foreground=self.colors['text']
                    )
                except:
                    pass
            
            # Update specific styled labels
            if hasattr(widget, 'cget'):
                try:
                    # Check if it's a card background element
                    if 'Card' in str(widget.winfo_parent()):
                        widget.configure(background=self.colors['card_bg'])
                except:
                    pass
            
            # Recursively update children
            for child in widget.winfo_children():
                self._update_widget_colors(child)
        
        except:
            pass
    
    def open_search(self):
        """Open search dialog for chat history"""
        SearchDialog(self.root, self)
    
    # toggle_favorite_chat method removed - feature was non-functional
    
    def show_smart_suggestions(self):
        """Generate and display smart suggestions based on user data"""
        if not self.authenticated:
            return
        
        # Show suggestions frame
        self.suggestions_frame.grid()
        
        # Generate suggestions (this would analyze actual data)
        suggestions = []
        
        # Check when they last asked about certain topics
        recent_topics = [msg.get('message', '').lower() for msg in self.current_chat_history[-10:]]
        
        if not any('sleep' in topic for topic in recent_topics):
            suggestions.append("You haven't checked your sleep data recently")
        
        if not any('steps' in topic or 'walking' in topic for topic in recent_topics):
            suggestions.append("How about reviewing your step count?")
        
        if not any('heart' in topic for topic in recent_topics):
            suggestions.append("Check your heart rate trends")
        
        if suggestions:
            suggestion_text = "Ὂ1 " + " • ".join(suggestions[:2])
            self.suggestions_label.config(text=suggestion_text)
        else:
            self.suggestions_frame.grid_remove()
    
    def show_followup_buttons(self, response_text):
        """Show context-aware follow-up buttons after AI response"""
        # Clear existing buttons
        for widget in self.followup_frame.winfo_children():
            widget.destroy()
        
        # Generate follow-up questions based on response
        followups = []
        
        response_lower = response_text.lower()
        
        if 'steps' in response_lower or 'walking' in response_lower:
            followups = [
                "Compare to last week",
                "Show me a weekly trend",
                "What's my daily average?"
            ]
        elif 'sleep' in response_lower:
            followups = [
                "How does this compare to my goal?",
                "Show sleep quality trends",
                "What affects my sleep?"
            ]
        elif 'workout' in response_lower or 'activity' in response_lower:
            followups = [
                "Show workout details",
                "Compare to previous workouts",
                "What's my weekly total?"
            ]
        elif 'heart' in response_lower:
            followups = [
                "Show resting heart rate trend",
                "Compare to healthy range",
                "What's my max heart rate?"
            ]
        else:
            followups = [
                "Tell me more",
                "Show details",
                "Any recommendations?"
            ]
        
        if followups:
            self.followup_frame.grid()
            
            ttk.Label(self.followup_frame,
                     text="Quick follow-ups:",
                     background=self.colors['card_bg'],
                     foreground=self.colors['text_secondary'],
                     font=('Segoe UI', 9)).grid(row=0, column=0, sticky=tk.W, padx=(0, 10))
            
            for i, question in enumerate(followups[:3]):
                btn = ttk.Button(self.followup_frame,
                               text=question,
                               style='Modern.TButton',
                               command=lambda q=question: self.use_example(q))
                btn.grid(row=0, column=i+1, padx=5)
        else:
            self.followup_frame.grid_remove()
    
    def export_conversation_report(self):
        """Export current conversation as a formatted document"""
        ExportReportDialog(self.root, self)
    
    def create_tooltip(self, widget, text):
        """Create a tooltip for a widget"""
        tooltip = ToolTip(widget, text)


class CustomizeQuestionsDialog(tk.Toplevel):
    """Dialog for customizing quick questions"""
    
    def __init__(self, parent, app):
        super().__init__(parent)
        
        # Make transient first
        self.transient(parent)
        
        self.title("Customize Quick Questions")
        self.app = app
        
        # Calculate centered position
        width = 700
        height = 550
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (width // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (height // 2)
        self.geometry(f"{width}x{height}+{x}+{y}")
        self.withdraw()
        
        # Apply colors
        self.colors = app.colors
        self.configure(bg=self.colors['bg'])
        
        # Set icon
        try:
            if getattr(sys, 'frozen', False):
                base_path = Path(sys._MEIPASS)
            else:
                base_path = Path(__file__).parent
            icon_path = base_path / "logo.ico"
            if icon_path.exists():
                self.iconbitmap(str(icon_path))
        except Exception as e:
            logger.debug(f"Could not load Customize Questions dialog icon: {e}")
        
        # Main frame
        main_frame = ttk.Frame(self, padding="20", style='TFrame')
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(1, weight=1)
        
        # Title
        ttk.Label(main_frame,
                 text="⚙️ Customize Quick Questions",
                 font=('Segoe UI', 14, 'bold'),
                 background=self.colors['bg'],
                 foreground=self.colors['text']).grid(row=0, column=0, sticky=tk.W, pady=(0, 15))
        
        # Instructions
        ttk.Label(main_frame,
                 text="Add up to 8 quick questions (one per line). Leave blank to use defaults.",
                 font=('Segoe UI', 9),
                 background=self.colors['bg'],
                 foreground=self.colors['text_secondary']).grid(row=1, column=0, sticky=tk.W, pady=(0, 10))
        
        # Text area for questions
        text_frame = ttk.Frame(main_frame, style='Card.TFrame')
        text_frame.grid(row=2, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 15))
        text_frame.columnconfigure(0, weight=1)
        text_frame.rowconfigure(0, weight=1)
        
        self.questions_text = tk.Text(text_frame,
                                     height=15,
                                     font=('Segoe UI', 10),
                                     wrap=tk.WORD,
                                     relief=tk.FLAT,
                                     borderwidth=0,
                                     padx=10,
                                     pady=10,
                                     bg=self.colors['card_bg'],
                                     fg=self.colors['text'],
                                     insertbackground=self.colors['accent'])
        self.questions_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=self.questions_text.yview)
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        self.questions_text.config(yscrollcommand=scrollbar.set)
        
        # Load current questions
        self.load_questions()
        
        # Buttons
        button_frame = ttk.Frame(main_frame, style='TFrame')
        button_frame.grid(row=3, column=0, pady=(0, 0))
        
        ttk.Button(button_frame,
                  text="💾 Save",
                  command=self.save_questions,
                  style='Accent.TButton').grid(row=0, column=0, padx=5)
        
        ttk.Button(button_frame,
                  text="↺ Reset to Defaults",
                  command=self.reset_to_defaults,
                  style='Modern.TButton').grid(row=0, column=1, padx=5)
        
        ttk.Button(button_frame,
                  text="Cancel",
                  command=self.destroy,
                  style='Modern.TButton').grid(row=0, column=2, padx=5)
        
        # Show window
        self.deiconify()
        self.grab_set()
    
    def load_questions(self):
        """Load current questions into text area"""
        quick_questions_file = self.app.config_dir / "quick_questions.json"
        
        try:
            if quick_questions_file.exists():
                with open(quick_questions_file, 'r') as f:
                    questions = json.load(f)
            else:
                questions = [
                    "How many steps did I take today?",
                    "What was my last workout?",
                    "How did I sleep last night?",
                    "Show me my recent activities",
                ]
        except Exception as e:
            logger.error(f"Error loading questions: {e}")
            questions = []
        
        self.questions_text.delete("1.0", tk.END)
        self.questions_text.insert("1.0", "\n".join(questions))
    
    def save_questions(self):
        """Save custom questions"""
        # Get text and split by lines
        text = self.questions_text.get("1.0", tk.END).strip()
        questions = [q.strip() for q in text.split('\n') if q.strip()]
        
        # Limit to 8 questions
        if len(questions) > 8:
            messagebox.showwarning(
                "Too Many Questions",
                f"Maximum 8 questions allowed. Only the first 8 will be saved.",
                parent=self
            )
            questions = questions[:8]
        
        # Save to file
        quick_questions_file = self.app.config_dir / "quick_questions.json"
        
        try:
            with open(quick_questions_file, 'w') as f:
                json.dump(questions, f, indent=2)
            
            # Reload questions in main window
            self.app.load_quick_questions()
            
            messagebox.showinfo(
                "Saved",
                f"Saved {len(questions)} quick question(s)!",
                parent=self
            )
            self.destroy()
            
        except Exception as e:
            logger.error(f"Error saving questions: {e}")
            messagebox.showerror(
                "Save Error",
                f"Failed to save questions: {e}",
                parent=self
            )
    
    def reset_to_defaults(self):
        """Reset to default questions"""
        defaults = [
            "How many steps did I take today?",
            "What was my last workout?",
            "How did I sleep last night?",
            "Show me my recent activities",
        ]
        
        self.questions_text.delete("1.0", tk.END)
        self.questions_text.insert("1.0", "\n".join(defaults))


class ToolTip:
    """Simple tooltip class"""
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tooltip_window = None
        self.widget.bind('<Enter>', self.show_tooltip)
        self.widget.bind('<Leave>', self.hide_tooltip)
    
    def show_tooltip(self, event=None):
        """Display tooltip"""
        if self.tooltip_window or not self.text:
            return
        
        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 5
        
        self.tooltip_window = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        
        label = tk.Label(tw, text=self.text, 
                        background="#2D2D30",
                        foreground="#E5E5E5",
                        relief=tk.SOLID,
                        borderwidth=1,
                        font=('Segoe UI', 9),
                        padx=8,
                        pady=4)
        label.pack()
    
    def hide_tooltip(self, event=None):
        """Hide tooltip"""
        if self.tooltip_window:
            self.tooltip_window.destroy()
            self.tooltip_window = None


class SavedPromptsDialog(tk.Toplevel):
    """Dialog for managing saved prompts with full text preview"""
    
    def __init__(self, parent, app):
        super().__init__(parent)
        self.title("Saved Prompts")
        self.app = app
        
        # Calculate centered position BEFORE setting geometry
        width = 760
        height = 560
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (width // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (height // 2)
        self.geometry(f"{width}x{height}+{x}+{y}")
        self.withdraw()
        
        # Apply Fluent Design colors
        self.colors = app.colors
        self.configure(bg=self.colors['bg'])
        # Set window icon
        try:
            if getattr(sys, 'frozen', False):
                base_path = Path(sys._MEIPASS)
            else:
                base_path = Path(__file__).parent
            
            icon_path = base_path / "logo.ico"
            if icon_path.exists():
                self.iconbitmap(str(icon_path))
        except Exception as e:
            logger.debug(f"Could not load Saved Prompts dialog icon: {e}")
        
        # Make modal
        self.transient(parent)
        
        # Configure ttk styles for this dialog
        style = ttk.Style()
        style.configure('SavedPrompts.TFrame', background=self.colors['bg'])
        style.configure('SavedPrompts.Card.TFrame', background=self.colors['card_bg'])
        style.configure('SavedPrompts.TLabel', 
                       background=self.colors['bg'], 
                       foreground=self.colors['text'])
        style.configure('SavedPrompts.Title.TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 14, 'bold'))
        style.configure('SavedPrompts.SubTitle.TLabel',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text_secondary'],
                       font=('Segoe UI', 9, 'bold'))
        
        main_frame = ttk.Frame(self, padding="20", style='SavedPrompts.TFrame')
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(1, weight=1)
        
        # Title
        title = ttk.Label(main_frame, text="Saved Prompts", style='SavedPrompts.Title.TLabel')
        title.grid(row=0, column=0, sticky=tk.W, pady=(0, 10))
        
        # Split frame for Listbox (left/top) and Full Text Preview (right/bottom)
        paned = ttk.PanedWindow(main_frame, orient=tk.VERTICAL)
        paned.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 15))
        
        # Listbox frame
        list_frame = ttk.Frame(paned, style='SavedPrompts.Card.TFrame', padding="10")
        list_frame.columnconfigure(0, weight=1)
        list_frame.rowconfigure(1, weight=1)
        
        ttk.Label(list_frame, text="SELECT PROMPT:", style='SavedPrompts.SubTitle.TLabel').grid(row=0, column=0, sticky=tk.W, pady=(0, 5))
        
        self.prompts_listbox = tk.Listbox(list_frame, 
                                          font=('Segoe UI', 10), 
                                          height=6,
                                          bg=self.colors['card_bg'],
                                          fg=self.colors['text'],
                                          selectbackground=self.colors['accent'],
                                          selectforeground='white',
                                          relief=tk.FLAT,
                                          borderwidth=1)
        self.prompts_listbox.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        list_scroll = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.prompts_listbox.yview)
        list_scroll.grid(row=1, column=1, sticky=(tk.N, tk.S))
        self.prompts_listbox.config(yscrollcommand=list_scroll.set)
        
        paned.add(list_frame, weight=1)
        
        # Preview text frame
        preview_frame = ttk.Frame(paned, style='SavedPrompts.Card.TFrame', padding="10")
        preview_frame.columnconfigure(0, weight=1)
        preview_frame.rowconfigure(1, weight=1)
        
        preview_header = ttk.Frame(preview_frame, style='SavedPrompts.Card.TFrame')
        preview_header.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 5))
        preview_header.columnconfigure(0, weight=1)
        
        ttk.Label(preview_header, text="FULL PROMPT TEXT (EDITABLE):", style='SavedPrompts.SubTitle.TLabel').grid(row=0, column=0, sticky=tk.W)
        ttk.Button(preview_header, text="💾 Save Prompt Text Changes", command=self.save_inline_changes).grid(row=0, column=1, sticky=tk.E)
        
        from tkinter import scrolledtext
        self.preview_text = scrolledtext.ScrolledText(preview_frame,
                                                     font=('Segoe UI', 10),
                                                     wrap=tk.WORD,
                                                     height=10,
                                                     bg=self.colors['bg'],
                                                     fg=self.colors['text'],
                                                     insertbackground=self.colors['text'],
                                                     relief=tk.FLAT,
                                                     borderwidth=1)
        self.preview_text.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        paned.add(preview_frame, weight=2)
        
        # Selection bindings
        self.prompts_listbox.bind('<<ListboxSelect>>', self.on_select_prompt)
        self.prompts_listbox.bind('<Double-Button-1>', lambda e: self.use_prompt())
        self.prompts_listbox.bind('<KeyRelease-Up>', self.on_select_prompt)
        self.prompts_listbox.bind('<KeyRelease-Down>', self.on_select_prompt)
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=2, column=0, pady=(0, 5))
        
        ttk.Button(button_frame, text="Use Selected", command=self.use_prompt).grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="Edit Prompt", command=self.edit_prompt).grid(row=0, column=1, padx=5)
        ttk.Button(button_frame, text="New Prompt", command=self.new_prompt).grid(row=0, column=2, padx=5)
        ttk.Button(button_frame, text="Delete", command=self.delete_prompt).grid(row=0, column=3, padx=5)
        ttk.Button(button_frame, text="Close", command=self.destroy).grid(row=0, column=4, padx=5)
        
        self.load_prompts()
        self.deiconify()
        self.grab_set()
    
    def clean_name_for_display(self, name: str) -> str:
        """Strip non-printable emoji surrogates to avoid rectangular boxes in Tkinter Listbox"""
        import re
        cleaned = re.sub(r'[\U00010000-\U0010ffff]', '', name).strip()
        return cleaned or name

    def load_prompts(self, selected_index=0):
        """Load and display saved prompts"""
        self.prompts_listbox.delete(0, tk.END)
        self.prompts = self.app.load_saved_prompts()
        
        for prompt in self.prompts:
            clean_title = self.clean_name_for_display(prompt['name'])
            self.prompts_listbox.insert(tk.END, f"  {clean_title}")
            
        if self.prompts:
            idx = min(selected_index, len(self.prompts) - 1)
            self.prompts_listbox.selection_set(idx)
            self.on_select_prompt()

    def on_select_prompt(self, event=None):
        """Update full prompt preview text when selecting item in listbox"""
        selection = self.prompts_listbox.curselection()
        if not selection:
            return
        index = selection[0]
        if 0 <= index < len(self.prompts):
            prompt_obj = self.prompts[index]
            full_text = prompt_obj.get('prompt', '')
            self.preview_text.delete('1.0', tk.END)
            self.preview_text.insert('1.0', full_text)
            
    def save_inline_changes(self):
        """Save direct edits made in the preview text area"""
        selection = self.prompts_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a prompt to save changes for.", parent=self)
            return
        index = selection[0]
        prompt_obj = self.prompts[index]
        new_text = self.preview_text.get("1.0", tk.END).strip()
        if not new_text:
            messagebox.showerror("Error", "Prompt text cannot be empty.", parent=self)
            return
        self.app.update_saved_prompt(index, prompt_obj['name'], new_text)
        messagebox.showinfo("Saved", f"Successfully saved prompt text changes for '{prompt_obj['name']}'!", parent=self)
        self.load_prompts(selected_index=index)

    def edit_prompt(self):
        """Edit selected prompt name and text in a modal dialog"""
        selection = self.prompts_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a prompt to edit.", parent=self)
            return
        index = selection[0]
        prompt_obj = self.prompts[index]
        
        dialog = tk.Toplevel(self)
        dialog.title("Edit Prompt")
        dialog.geometry("560x380")
        dialog.transient(self)
        dialog.grab_set()
        
        frame = ttk.Frame(dialog, padding="20")
        frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        dialog.columnconfigure(0, weight=1)
        dialog.rowconfigure(0, weight=1)
        frame.columnconfigure(1, weight=1)
        frame.rowconfigure(1, weight=1)
        
        ttk.Label(frame, text="Prompt Name:", font=('Segoe UI', 10, 'bold')).grid(row=0, column=0, sticky=tk.W, pady=5)
        name_entry = ttk.Entry(frame, font=('Segoe UI', 10))
        name_entry.insert(0, prompt_obj['name'])
        name_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=5, padx=(10, 0))
        
        ttk.Label(frame, text="Prompt Text:", font=('Segoe UI', 10, 'bold')).grid(row=1, column=0, sticky=(tk.W, tk.N), pady=5)
        
        from tkinter import scrolledtext
        prompt_text = scrolledtext.ScrolledText(frame, font=('Segoe UI', 10), wrap=tk.WORD, height=10)
        prompt_text.insert("1.0", prompt_obj['prompt'])
        prompt_text.grid(row=1, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5, padx=(10, 0))
        
        def save():
            name = name_entry.get().strip()
            prompt = prompt_text.get("1.0", tk.END).strip()
            
            if not name or not prompt:
                messagebox.showerror("Validation Error", "Please enter both name and prompt text", parent=dialog)
                return
            
            self.app.update_saved_prompt(index, name, prompt)
            self.load_prompts(selected_index=index)
            dialog.destroy()
        
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=2, column=0, columnspan=2, pady=(15, 0))
        
        ttk.Button(button_frame, text="Save Changes", command=save).grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="Cancel", command=dialog.destroy).grid(row=0, column=1, padx=5)
    
    def new_prompt(self):
        """Create new saved prompt"""
        dialog = tk.Toplevel(self)
        dialog.title("New Prompt")
        dialog.geometry("560x380")
        dialog.transient(self)
        dialog.grab_set()
        
        frame = ttk.Frame(dialog, padding="20")
        frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        dialog.columnconfigure(0, weight=1)
        dialog.rowconfigure(0, weight=1)
        frame.columnconfigure(1, weight=1)
        frame.rowconfigure(1, weight=1)
        
        ttk.Label(frame, text="Prompt Name:", font=('Segoe UI', 10, 'bold')).grid(row=0, column=0, sticky=tk.W, pady=5)
        name_entry = ttk.Entry(frame, font=('Segoe UI', 10))
        name_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=5, padx=(10, 0))
        
        ttk.Label(frame, text="Prompt Text:", font=('Segoe UI', 10, 'bold')).grid(row=1, column=0, sticky=(tk.W, tk.N), pady=5)
        
        from tkinter import scrolledtext
        prompt_text = scrolledtext.ScrolledText(frame, font=('Segoe UI', 10), wrap=tk.WORD, height=10)
        prompt_text.grid(row=1, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5, padx=(10, 0))
        
        def save():
            name = name_entry.get().strip()
            prompt = prompt_text.get("1.0", tk.END).strip()
            
            if not name or not prompt:
                messagebox.showerror("Validation Error", "Please enter both name and prompt text", parent=dialog)
                return
            
            self.app.save_prompt(name, prompt)
            self.load_prompts(selected_index=len(self.prompts))
            dialog.destroy()
        
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=2, column=0, columnspan=2, pady=(15, 0))
        
        ttk.Button(button_frame, text="Save Prompt", command=save).grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="Cancel", command=dialog.destroy).grid(row=0, column=1, padx=5)
    
    def use_prompt(self):
        """Use selected prompt"""
        selection = self.prompts_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a prompt to use", parent=self)
            return
        
        index = selection[0]
        prompt = self.prompts[index]
        
        # Insert into message entry and close dialog
        self.app.message_entry.delete("1.0", tk.END)
        self.app.message_entry.insert("1.0", prompt['prompt'])
        self.destroy()
    
    def delete_prompt(self):
        """Delete selected prompt"""
        selection = self.prompts_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a prompt to delete", parent=self)
            return
        
        index = selection[0]
        prompt = self.prompts[index]
        
        if messagebox.askyesno("Confirm Delete", f"Delete prompt '{prompt['name']}'?", parent=self):
            self.app.delete_saved_prompt(index)
            self.load_prompts()






class SearchDialog(tk.Toplevel):
    """Dialog for searching chat history"""
    
    def __init__(self, parent, app):
        super().__init__(parent)
        self.title("Search Chat History")
        self.app = app
        
        # Calculate centered position BEFORE setting geometry
        width = 700
        height = 500
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (width // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (height // 2)
        self.geometry(f"{width}x{height}+{x}+{y}")
        
        # Apply Fluent Design colors
        self.colors = app.colors
        self.configure(bg=self.colors['bg'])
        
        # Set window icon
        try:
            if getattr(sys, 'frozen', False):
                base_path = Path(sys._MEIPASS)
            else:
                base_path = Path(__file__).parent
            
            icon_path = base_path / "logo.ico"
            if icon_path.exists():
                self.iconbitmap(str(icon_path))
        except Exception as e:
            logger.debug(f"Could not load Search dialog icon: {e}")
        
        self.transient(parent)
        self.grab_set()
        
        main_frame = ttk.Frame(self, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(2, weight=1)
        
        # Title
        ttk.Label(main_frame, text="History", 
                 font=('Segoe UI', 14, 'bold')).grid(row=0, column=0, sticky=tk.W, pady=(0, 15))
        
        # Search box
        search_frame = ttk.Frame(main_frame, style='Search.Card.TFrame', padding="15")
        search_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 15))
        search_frame.columnconfigure(0, weight=1)
        
        self.search_entry = ttk.Entry(search_frame, font=('Segoe UI', 10), style='Modern.TEntry')
        self.search_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 10))
        self.search_entry.bind('<Return>', lambda e: self.perform_search())
        
        ttk.Button(search_frame, text="Search", 
                  command=self.perform_search).grid(row=0, column=1)
        
        # Results
        results_frame = ttk.Frame(main_frame)
        results_frame.grid(row=2, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        results_frame.columnconfigure(0, weight=1)
        results_frame.rowconfigure(0, weight=1)
        
        self.results_text = scrolledtext.ScrolledText(results_frame,
                                                      wrap=tk.WORD,
                                                      font=('Segoe UI', 9),
                                                      state=tk.DISABLED,
                                                      bg=self.colors['card_bg'],
                                                      fg=self.colors['text'],
                                                      relief=tk.FLAT,
                                                      borderwidth=0,
                                                      padx=15,
                                                      pady=15)
        self.results_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Close button
        ttk.Button(main_frame, text="Close", 
                  command=self.destroy).grid(row=3, column=0, pady=(15, 0))
    
    def perform_search(self):
        """Search through all saved chats"""
        query = self.search_entry.get().strip().lower()
        if not query:
            return
        
        self.results_text.config(state=tk.NORMAL)
        self.results_text.delete(1.0, tk.END)
        
        try:
            chat_files = sorted(self.app.chat_history_dir.glob("chat_*.json"),
                              key=lambda f: f.stat().st_mtime,
                              reverse=True)
            
            results_found = 0
            for file in chat_files:
                with open(file, 'r') as f:
                    data = json.load(f)
                    messages = data.get('messages', [])
                    
                    for msg in messages:
                        text = msg.get('message', '').lower()
                        if query in text:
                            results_found += 1
                            sender = msg.get('sender', 'Unknown')
                            timestamp = msg.get('timestamp', '')
                            
                            self.results_text.insert(tk.END, f"\n{sender} ({timestamp[:10]}):\n", 'bold')
                            self.results_text.insert(tk.END, f"{msg.get('message', '')}...\n\n")
                            
                            if results_found >= 20:  # Limit results
                                break
                
                if results_found >= 20:
                    break
            
            if results_found == 0:
                self.results_text.insert(tk.END, f"No results found for '{query}'")
            else:
                self.results_text.insert(tk.END, f"\n--- {results_found} results found ---")
        
        except Exception as e:
            self.results_text.insert(tk.END, f"Error searching: {e}")
        
        self.results_text.config(state=tk.DISABLED)


class ExportReportDialog(tk.Toplevel):
    """Dialog for exporting conversation reports"""
    
    def __init__(self, parent, app):
        super().__init__(parent)
        
        # Make transient first
        self.transient(parent)
        
        self.title("Export Report")
        self.app = app
        
        # Calculate centered position
        width = 500
        height = 400
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (width // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (height // 2)
        self.geometry(f"{width}x{height}+{x}+{y}")
        self.withdraw()
        
        main_frame = ttk.Frame(self, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        
        # Title
        ttk.Label(main_frame, text="📄 Export Report", 
                 font=('Segoe UI', 14, 'bold')).grid(row=0, column=0, sticky=tk.W, pady=(0, 20))
        
        # Export options
        ttk.Label(main_frame, text="Choose export format:", 
                 font=('Segoe UI', 10)).grid(row=1, column=0, sticky=tk.W, pady=(0, 10))
        
        self.export_format = tk.StringVar(value="pdf")
        
        ttk.Radiobutton(main_frame, text="PDF Document", 
                       variable=self.export_format, value="pdf").grid(row=2, column=0, sticky=tk.W, pady=3)
        ttk.Radiobutton(main_frame, text="Word Document (.docx)", 
                       variable=self.export_format, value="docx").grid(row=3, column=0, sticky=tk.W, pady=3)
        ttk.Radiobutton(main_frame, text="Text File (.txt)", 
                       variable=self.export_format, value="txt").grid(row=4, column=0, sticky=tk.W, pady=3)
        
        # Include options
        ttk.Label(main_frame, text="\nInclude:", 
                 font=('Segoe UI', 10, 'bold')).grid(row=5, column=0, sticky=tk.W, pady=(15, 10))
        
        self.include_timestamp = tk.BooleanVar(value=True)
        self.include_system = tk.BooleanVar(value=False)
        
        ttk.Checkbutton(main_frame, text="Timestamps", 
                       variable=self.include_timestamp).grid(row=6, column=0, sticky=tk.W, pady=3)
        ttk.Checkbutton(main_frame, text="System messages", 
                       variable=self.include_system).grid(row=7, column=0, sticky=tk.W, pady=3)
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=8, column=0, pady=(20, 0))
        
        ttk.Button(button_frame, text="Export", 
                  command=self.export_report).grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="Cancel", 
                  command=self.destroy).grid(row=0, column=1, padx=5)
        
        # Show window after all widgets created
        self.deiconify()
        self.grab_set()
    
    def export_report(self):
        """Export the conversation in selected format"""
        if not self.app.current_chat_history:
            messagebox.showwarning("No Data", "No conversation to export!", parent=self)
            return
        
        format_type = self.export_format.get()
        
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            if format_type == "txt":
                filename = self.app.chat_history_dir / f"report_{timestamp}.txt"
                self._export_txt(filename)
            elif format_type == "pdf":
                filename = self.app.chat_history_dir / f"report_{timestamp}.pdf"
                self._export_pdf(filename)
            elif format_type == "docx":
                filename = self.app.chat_history_dir / f"report_{timestamp}.docx"
                self._export_docx(filename)
            
            messagebox.showinfo("Export Complete", 
                              f"Report exported successfully!\n\n{filename}",
                              parent=self)
            self.destroy()
        
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export: {e}", parent=self)
    
    def _export_txt(self, filename):
        """Export as plain text"""
        with open(filename, 'w', encoding='utf-8') as f:
            f.write("GARMIN CHAT CONVERSATION REPORT\n")
            f.write("=" * 60 + "\n\n")
            
            for msg in self.app.current_chat_history:
                if not self.include_system.get() and msg.get('type') == 'system':
                    continue
                
                sender = msg.get('sender', 'Unknown')
                text = msg.get('message', '')
                
                if self.include_timestamp.get():
                    timestamp = msg.get('timestamp', '')
                    f.write(f"[{timestamp}] {sender}:\n")
                else:
                    f.write(f"{sender}:\n")
                
                f.write(f"{text}\n\n")
                f.write("-" * 60 + "\n\n")
    
    def _export_pdf(self, filename):
        """Export as PDF"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch
            
            doc = SimpleDocTemplate(str(filename), pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Title
            title = Paragraph("HealthChat Conversation Report", styles['Title'])
            story.append(title)
            story.append(Spacer(1, 0.3*inch))
            
            # Messages
            for msg in self.app.current_chat_history:
                if not self.include_system.get() and msg.get('type') == 'system':
                    continue
                
                sender = msg.get('sender', 'Unknown')
                text = msg.get('message', '')
                
                if self.include_timestamp.get():
                    timestamp = msg.get('timestamp', '')
                    header = f"<b>[{timestamp}] {sender}:</b>"
                else:
                    header = f"<b>{sender}:</b>"
                
                story.append(Paragraph(header, styles['Normal']))
                story.append(Paragraph(text, styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
            
            doc.build(story)
        
        except ImportError:
            # Fallback to text if reportlab not installed
            messagebox.showwarning("PDF Export", 
                                 "PDF export requires 'reportlab' package.\nExporting as text instead.",
                                 parent=self)
            txt_filename = str(filename).replace('.pdf', '.txt')
            self._export_txt(Path(txt_filename))
    
    def _export_docx(self, filename):
        """Export as Word document"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading('HealthChat Conversation Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Messages
            for msg in self.app.current_chat_history:
                if not self.include_system.get() and msg.get('type') == 'system':
                    continue
                
                sender = msg.get('sender', 'Unknown')
                text = msg.get('message', '')
                
                if self.include_timestamp.get():
                    timestamp = msg.get('timestamp', '')
                    p = doc.add_paragraph()
                    p.add_run(f"[{timestamp}] {sender}:").bold = True
                else:
                    p = doc.add_paragraph()
                    p.add_run(f"{sender}:").bold = True
                
                doc.add_paragraph(text)
                doc.add_paragraph("_" * 60)
            
            doc.save(str(filename))
        
        except ImportError:
            # Fallback to text if python-docx not installed
            messagebox.showwarning("Word Export", 
                                 "Word export requires 'python-docx' package.\nExporting as text instead.",
                                 parent=self)
            txt_filename = str(filename).replace('.docx', '.txt')
            self._export_txt(Path(txt_filename))


class ChatHistoryViewer(tk.Toplevel):
    """Dialog for viewing saved chat histories"""
    
    def __init__(self, parent, app):
        super().__init__(parent)
        
        # Make transient FIRST
        self.transient(parent)
        
        self.title("Chat History")
        self.app = app
        
        # Calculate centered position BEFORE setting geometry
        width = 900
        height = 650
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (width // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (height // 2)
        self.geometry(f"{width}x{height}+{x}+{y}")
        
        # Withdraw immediately after geometry to prevent flash
        self.withdraw()
        
        # Set window icon (same as main window)
        try:
            # Get the correct base path for PyInstaller exe
            if getattr(sys, 'frozen', False):
                # Running as compiled executable
                base_path = Path(sys._MEIPASS)
            else:
                # Running as script
                base_path = Path(__file__).parent
            
            icon_path = base_path / "logo.ico"
            if icon_path.exists():
                self.iconbitmap(str(icon_path))
        except Exception as e:
            logger.debug(f"Could not load Chat History dialog icon: {e}")
        
        # Apply Fluent Design colors
        self.colors = app.colors
        self.configure(bg=self.colors['bg'])
        
        # Configure ttk styles
        style = ttk.Style()
        style.configure('History.TFrame', background=self.colors['bg'])
        style.configure('History.Card.TLabelframe',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       bordercolor=self.colors['border'],
                       borderwidth=1,
                       relief='flat')
        style.configure('History.Card.TLabelframe.Label',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 10, 'bold'))
        style.configure('History.TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'])
        style.configure('History.Title.TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 14, 'bold'))
        
        # Main container with two columns
        main_frame = ttk.Frame(self, padding="20", style='History.TFrame')
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(1, weight=1)
        
        # Title
        title = ttk.Label(main_frame, text="History", style='History.Title.TLabel')
        title.grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=(0, 15))
        
        # Left panel - Chat list
        left_frame = ttk.LabelFrame(main_frame, text="Saved Chats", padding="10", style='History.Card.TLabelframe')
        left_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 10))
        left_frame.columnconfigure(0, weight=1)
        left_frame.rowconfigure(0, weight=1)
        
        # Listbox with scrollbar
        list_frame = ttk.Frame(left_frame)
        list_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        list_frame.columnconfigure(0, weight=1)
        list_frame.rowconfigure(0, weight=1)
        
        self.chat_listbox = tk.Listbox(list_frame, 
                                       font=('Segoe UI', 9), 
                                       width=30,
                                       bg=self.colors['card_bg'],
                                       fg=self.colors['text'],
                                       selectbackground=self.colors['accent'],
                                       selectforeground='white',
                                       relief=tk.FLAT,
                                       borderwidth=0)
        self.chat_listbox.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.chat_listbox.bind('<<ListboxSelect>>', self.on_chat_select)
        
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.chat_listbox.yview)
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        self.chat_listbox.config(yscrollcommand=scrollbar.set)
        
        # Right panel - Chat viewer
        right_frame = ttk.LabelFrame(main_frame, text="Chat Content", padding="10", style='History.Card.TLabelframe')
        right_frame.grid(row=1, column=1, sticky=(tk.W, tk.E, tk.N, tk.S))
        right_frame.columnconfigure(0, weight=1)
        right_frame.rowconfigure(1, weight=1)
        
        # Chat info
        self.info_label = ttk.Label(right_frame, 
                                    text="Select a chat to view", 
                                    font=('Segoe UI', 9),
                                    background=self.colors['card_bg'],
                                    foreground=self.colors['text_secondary'])
        self.info_label.grid(row=0, column=0, sticky=tk.W, pady=(0, 5))
        
        # Chat content display
        self.chat_display = scrolledtext.ScrolledText(right_frame,
                                                      wrap=tk.WORD,
                                                      font=('Segoe UI', 9),
                                                      state=tk.DISABLED,
                                                      bg=self.colors['card_bg'],
                                                      fg=self.colors['text'],
                                                      relief=tk.FLAT,
                                                      borderwidth=0,
                                                      padx=10,
                                                      pady=10)
        self.chat_display.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure text tags
        self.chat_display.tag_configure('user', foreground='#2980b9', font=('Segoe UI', 9, 'bold'))
        self.chat_display.tag_configure('assistant', foreground='#27ae60', font=('Segoe UI', 9, 'bold'))
        self.chat_display.tag_configure('system', foreground='#e74c3c', font=('Segoe UI', 9, 'italic'))
        self.chat_display.tag_configure('timestamp', foreground='#95a5a6', font=('Segoe UI', 8))
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=2, column=0, columnspan=2, pady=(15, 0))
        
        ttk.Button(button_frame, text="Load Into Chat", command=self.load_into_current).grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="Rename", command=self.rename_chat).grid(row=0, column=1, padx=5)
        ttk.Button(button_frame, text="Delete", command=self.delete_chat).grid(row=0, column=2, padx=5)
        ttk.Button(button_frame, text="Open Folder", command=self.open_folder).grid(row=0, column=3, padx=5)
        ttk.Button(button_frame, text="Close", command=self.destroy).grid(row=0, column=4, padx=5)
        
        # Load chats
        self.load_chat_list()
        self.deiconify()
        self.grab_set()
    
    def load_chat_list(self):
        """Load list of saved chat files"""
        self.chat_listbox.delete(0, tk.END)
        self.chat_files = []
        
        try:
            # Get all chat JSON files, sorted newest first
            files = sorted(self.app.chat_history_dir.glob("chat_*.json"),
                          key=lambda f: f.stat().st_mtime,
                          reverse=True)
            
            for file in files:
                self.chat_files.append(file)
                
                # Try to load custom name from JSON
                try:
                    with open(file, 'r') as f:
                        data = json.load(f)
                        custom_name = data.get('custom_name')
                except:
                    custom_name = None
                
                # Format: chat_YYYYMMDD_HHMMSS.json or chat_YYYYMMDD_HHMMSS_CustomName.json
                try:
                    filename = file.stem
                    
                    # Extract custom name from filename if present
                    if '_' in filename and not custom_name:
                        parts = filename.split('_')
                        if len(parts) >= 4:  # chat_DATE_TIME_NAME format
                            custom_name = '_'.join(parts[3:])
                    
                    # Extract timestamp
                    timestamp = filename.replace('chat_', '').split('_')[0:2]
                    if len(timestamp) == 2:
                        date_part = timestamp[0]  # YYYYMMDD
                        time_part = timestamp[1]  # HHMMSS
                        
                        date_str = f"{date_part[:4]}-{date_part[4:6]}-{date_part[6:8]}"
                        time_str = f"{time_part[:2]}:{time_part[2:4]}:{time_part[4:6]}"
                        
                        if custom_name:
                            display = f"{custom_name} ({date_str} {time_str})"
                        else:
                            display = f"{date_str} {time_str}"
                    else:
                        display = file.name
                except:
                    display = file.name
                
                self.chat_listbox.insert(tk.END, display)
            
            if not files:
                self.info_label.config(text="No saved chats found")
        
        except Exception as e:
            logger.error(f"Error loading chat list: {e}")
            self.info_label.config(text="Error loading chat list")
    
    def on_chat_select(self, event):
        """Display selected chat"""
        selection = self.chat_listbox.curselection()
        if not selection:
            return
        
        index = selection[0]
        file = self.chat_files[index]
        
        try:
            with open(file, 'r') as f:
                data = json.load(f)
            
            # Update info
            saved_at = data.get('saved_at', '')
            messages = data.get('messages', [])
            
            try:
                dt = datetime.fromisoformat(saved_at)
                date_str = dt.strftime("%B %d, %Y at %I:%M %p")
            except:
                date_str = saved_at
            
            self.info_label.config(text=f"Saved: {date_str} | {len(messages)} messages")
            
            # Display messages
            self.chat_display.config(state=tk.NORMAL)
            self.chat_display.delete(1.0, tk.END)
            
            for msg in messages:
                timestamp = msg.get('timestamp', '')
                sender = msg.get('sender', 'Unknown')
                text = msg.get('message', '')
                msg_type = msg.get('type', 'user')
                
                # Format timestamp
                try:
                    ts = datetime.fromisoformat(timestamp)
                    ts_str = ts.strftime("%H:%M")
                except:
                    ts_str = timestamp[:5] if timestamp else ""
                
                self.chat_display.insert(tk.END, f"[{ts_str}] ", 'timestamp')
                self.chat_display.insert(tk.END, f"{sender}: ", msg_type)
                self.chat_display.insert(tk.END, f"{text}\n\n")
            
            self.chat_display.config(state=tk.DISABLED)
            self.chat_display.see(1.0)
        
        except Exception as e:
            logger.error(f"Error loading chat: {e}")
            messagebox.showerror("Error", f"Failed to load chat: {e}", parent=self)
    
    def load_into_current(self):
        """Load selected chat into current conversation"""
        selection = self.chat_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a chat to load", parent=self)
            return
        
        index = selection[0]
        file = self.chat_files[index]
        
        if messagebox.askyesno("Load Chat",
                              "This will replace your current conversation.\n\n"
                              "Load this chat into the main window?",
                              parent=self):
            try:
                with open(file, 'r') as f:
                    data = json.load(f)
                
                # Clear current chat
                self.app.chat_display.config(state=tk.NORMAL)
                self.app.chat_display.delete(1.0, tk.END)
                self.app.chat_display.config(state=tk.DISABLED)
                
                # Load messages
                messages = data.get('messages', [])
                self.app.current_chat_history = messages.copy()
                
                # Display in main window
                for msg in messages:
                    sender = msg.get('sender', 'Unknown')
                    text = msg.get('message', '')
                    msg_type = msg.get('type', 'user')
                    
                    self.app.add_message(sender, text, msg_type)
                
                messagebox.showinfo("Chat Loaded", f"Loaded {len(messages)} messages", parent=self)
                self.destroy()
            
            except Exception as e:
                logger.error(f"Error loading chat: {e}")
                messagebox.showerror("Error", f"Failed to load chat: {e}", parent=self)
    
    def delete_chat(self):
        """Delete selected chat file"""
        selection = self.chat_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a chat to delete", parent=self)
            return
        
        index = selection[0]
        file = self.chat_files[index]
        
        if messagebox.askyesno("Delete Chat",
                              f"Permanently delete this chat?\n\n{file.name}",
                              parent=self):
            try:
                file.unlink()
                
                # Clear display
                self.chat_display.config(state=tk.NORMAL)
                self.chat_display.delete(1.0, tk.END)
                self.chat_display.config(state=tk.DISABLED)
                self.info_label.config(text="Select a chat to view")
                
                # Reload list
                self.load_chat_list()
                
                logger.info(f"Deleted chat: {file.name}")
            
            except Exception as e:
                logger.error(f"Error deleting chat: {e}")
                messagebox.showerror("Error", f"Failed to delete chat: {e}", parent=self)
    
    def rename_chat(self):
        """Rename selected chat session"""
        selection = self.chat_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a chat to rename", parent=self)
            return
        
        index = selection[0]
        file = self.chat_files[index]
        
        # Load current name
        try:
            with open(file, 'r') as f:
                data = json.load(f)
                current_name = data.get('custom_name', '')
        except:
            current_name = ''
        
        # Prompt for new name
        dialog = tk.Toplevel(self)
        dialog.title("Rename Chat Session")
        dialog.geometry("500x180")
        dialog.transient(self)
        dialog.grab_set()
        
        # Set window icon
        try:
            if getattr(sys, 'frozen', False):
                base_path = Path(sys._MEIPASS)
            else:
                base_path = Path(__file__).parent
            
            icon_path = base_path / "logo.ico"
            if icon_path.exists():
                dialog.iconbitmap(str(icon_path))
        except Exception as e:
            logger.debug(f"Could not load Rename dialog icon: {e}")
        
        # Center on parent
        dialog.update_idletasks()
        x = self.winfo_x() + (self.winfo_width() // 2) - (dialog.winfo_width() // 2)
        y = self.winfo_y() + (self.winfo_height() // 2) - (dialog.winfo_height() // 2)
        dialog.geometry(f"+{x}+{y}")
        
        frame = ttk.Frame(dialog, padding="20")
        frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        dialog.columnconfigure(0, weight=1)
        dialog.rowconfigure(0, weight=1)
        frame.columnconfigure(1, weight=1)
        
        ttk.Label(frame, text="✏️ Rename Chat Session", 
                 font=('Segoe UI', 12, 'bold')).grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=(0, 15))
        
        ttk.Label(frame, text="New Name:", 
                 font=('Segoe UI', 10)).grid(row=1, column=0, sticky=tk.W, pady=5)
        
        name_var = tk.StringVar(value=current_name)
        name_entry = ttk.Entry(frame, textvariable=name_var, font=('Segoe UI', 10), width=40)
        name_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=5, padx=(10, 0))
        name_entry.focus()
        name_entry.select_range(0, tk.END)
        
        def save_rename():
            new_name = name_var.get().strip()
            try:
                # Load current data
                with open(file, 'r') as f:
                    data = json.load(f)
                
                # Update custom name
                data['custom_name'] = new_name if new_name else None
                
                # Save back to same file
                with open(file, 'w') as f:
                    json.dump(data, f, indent=2)
                
                # Optionally update filename
                if new_name:
                    # Extract timestamp from filename
                    filename_parts = file.stem.split('_')
                    if len(filename_parts) >= 3:
                        timestamp_date = filename_parts[1]
                        timestamp_time = filename_parts[2]
                        
                        # Sanitize name for filename
                        safe_name = "".join(c for c in new_name if c.isalnum() or c in (' ', '-', '_')).strip()
                        safe_name = safe_name.replace(' ', '_')
                        
                        new_filename = file.parent / f"chat_{timestamp_date}_{timestamp_time}_{safe_name}.json"
                        
                        # Rename file if new filename is different
                        if new_filename != file:
                            file.rename(new_filename)
                            logger.info(f"Renamed file: {file.name} -> {new_filename.name}")
                
                dialog.destroy()
                
                # Reload the chat list to show new name
                self.load_chat_list()
                
                # Reselect the same item (it may have moved due to sorting)
                for i, f in enumerate(self.chat_files):
                    if f.stem.startswith(f"chat_{timestamp_date}_{timestamp_time}"):
                        self.chat_listbox.selection_clear(0, tk.END)
                        self.chat_listbox.selection_set(i)
                        self.chat_listbox.see(i)
                        # Trigger the display update
                        self.on_chat_select(None)
                        break
                
                logger.info(f"Renamed chat session to: {new_name}")
                
            except Exception as e:
                logger.error(f"Error renaming chat: {e}")
                messagebox.showerror("Rename Error", f"Failed to rename chat: {e}", parent=dialog)
        
        def cancel():
            dialog.destroy()
        
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=2, column=0, columnspan=2, pady=(20, 0))
        
        ttk.Button(button_frame, text="Save Chat", command=save_rename, width=10).grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="Cancel", command=cancel, width=10).grid(row=0, column=1, padx=5)
        
        # Bind Enter key to save
        name_entry.bind('<Return>', lambda e: save_rename())
        dialog.bind('<Escape>', lambda e: cancel())
    
    def open_folder(self):
        """Open chat history folder in file explorer"""
        import subprocess
        import sys
        
        try:
            path = str(self.app.chat_history_dir)
            
            if sys.platform == 'win32':
                subprocess.run(['explorer', path])
            elif sys.platform == 'darwin':
                subprocess.run(['open', path])
            else:
                subprocess.run(['xdg-open', path])
        
        except Exception as e:
            logger.error(f"Error opening folder: {e}")
            messagebox.showerror("Error", f"Failed to open folder: {e}", parent=self)


def ensure_single_instance(mutex_name="HealthChatDesktop_SingleInstance_Mutex"):
    """Ensure only one instance of the application runs at a time on Windows."""
    if sys.platform == 'win32':
        try:
            import ctypes
            kernel32 = ctypes.windll.kernel32
            mutex = kernel32.CreateMutexW(None, False, mutex_name)
            last_error = kernel32.GetLastError()
            ERROR_ALREADY_EXISTS = 183
            if last_error == ERROR_ALREADY_EXISTS:
                return False, mutex
            return True, mutex
        except Exception as e:
            logger.warning(f"Failed to create single-instance mutex: {e}")
            return True, None
    return True, None


def main():
    """Main entry point"""
    import multiprocessing
    multiprocessing.freeze_support()

    # Single-instance enforcement
    is_single_instance, _mutex = ensure_single_instance()
    if not is_single_instance:
        print("Another instance of HealthChat Desktop is already running. Exiting.")
        root = tk.Tk()
        root.withdraw()
        messagebox.showinfo("HealthChat Desktop", "HealthChat Desktop körs redan!")
        root.destroy()
        sys.exit(0)

    print("\n" + "="*60)
    print("HealthChat - Desktop Application")
    print(f"Version {APP_VERSION}")
    print("="*60)
    print("\nStarting application...")
    print("="*60 + "\n")
    
    # Create hidden root window
    root = tk.Tk()
    root.withdraw()  # Hide main window initially
    
    # Show splash screen
    splash = SplashScreen(root)
    
    def load_app():
        """Load the main application"""
        try:
            splash.update_status("Initializing...")
            root.update()
            
            # Create main app
            splash.update_status("Loading configuration...")
            app = HealthChatApp(root)
            
            # Close splash and show main window
            splash.update_status("Ready!")
            root.after(500, lambda: (splash.close(), root.deiconify()))
            
        except Exception as e:
            splash.close()
            root.deiconify()
            messagebox.showerror("Startup Error", 
                               f"Failed to start application:\n\n{e}")
            root.destroy()
    
    # Load app after splash is shown
    root.after(100, load_app)
    root.mainloop()


if __name__ == "__main__":
    main()
